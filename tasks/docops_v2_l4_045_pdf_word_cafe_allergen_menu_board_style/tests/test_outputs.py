import json
import os
import re
import sys
import zipfile
from pathlib import Path
from datetime import date, datetime

import pdfplumber
from docx import Document
from docx.enum.section import WD_ORIENT
from pypdf import PdfReader

sys.path.insert(0, str(Path(__file__).parent))
from verifier_utils import *  # noqa: F401,F403

META_PATH = Path(os.environ.get('TASK_METADATA_PATH', '/tests/task_metadata.json'))
META = json.loads(META_PATH.read_text(encoding='utf-8'))
EXPECT = META['verifier_expectations']


def _vopt_norm_text(value):
    if value is None:
        return ''
    if isinstance(value, datetime):
        if value.hour or value.minute or value.second:
            return value.strftime('%Y-%m-%d %H:%M')
        return value.strftime('%Y-%m-%d')
    if isinstance(value, date):
        return value.strftime('%Y-%m-%d')
    if isinstance(value, float) and value.is_integer():
        value = int(value)
    text = str(value)
    text = text.replace('\u2013', '-').replace('\u2014', '-').replace('\u2212', '-')
    text = text.replace('\xa0', ' ')
    text = re.sub(r'\s+', ' ', text).strip()
    return text


def _vopt_norm_formula(value):
    return re.sub(r'\s+', '', _vopt_norm_text(value)).upper()


def _vopt_assert_formula(actual, expected, label='formula'):
    assert _vopt_norm_formula(actual) == _vopt_norm_formula(expected), f"{label}: expected {expected!r}, found {actual!r}"


def _vopt_norm_rows(rows):
    return [[_vopt_norm_text(cell) for cell in row] for row in rows]


def _vopt_rows_equal(actual, expected):
    return _vopt_norm_rows(actual) == _vopt_norm_rows(expected)


def _vopt_rows_text(rows):
    return '\n'.join('|'.join(_vopt_norm_text(cell) for cell in row) for row in rows)


def _vopt_norm_key(value):
    return re.sub(r'[^a-z0-9]+', '', _vopt_norm_text(value).lower())


def _vopt_header_has(header, required):
    actual = {_vopt_norm_text(cell).lower() for cell in header}
    missing = [cell for cell in required if _vopt_norm_text(cell).lower() not in actual]
    assert not missing, f"Missing expected header columns {missing!r}; actual={header!r}"


def _vopt_section_text(section_part):
    return '\n'.join(p.text for p in getattr(section_part, 'paragraphs', []) if p.text is not None)


def _vopt_header_text(doc):
    return '\n'.join(_vopt_section_text(section.header) for section in doc.sections)


def _vopt_footer_text(doc):
    return '\n'.join(_vopt_section_text(section.footer) for section in doc.sections)


def _vopt_table_ref(ws, expected_name=None):
    tables = ws.tables
    if expected_name:
        for name in tables.keys():
            if str(name).lower() == str(expected_name).lower():
                return tables[name].ref
    vals = list(tables.values())
    assert vals, f"No Excel native table found on sheet {ws.title!r}"
    return vals[0].ref


def _vopt_clean_area(value):
    text = _vopt_norm_text(value).replace("'", '').replace('$', '').replace(' ', '')
    if '!' in text:
        text = text.split('!', 1)[1]
    return text.upper()


def _vopt_assert_print_area(ws, expected):
    actual = _vopt_clean_area(ws.print_area)
    target = _vopt_clean_area(expected)
    assert target in actual or actual in target, f"{ws.title}: expected print area {expected!r}, found {ws.print_area!r}"


def _vopt_freeze_pane(value):
    return getattr(value, 'coordinate', value)


def _vopt_validated_cells(ws):
    cells = set()
    for dv in ws.data_validations.dataValidation:
        for rng in dv.cells.ranges:
            text = str(rng).replace('$', '')
            try:
                min_col, min_row, max_col, max_row = range_boundaries(text)
            except ValueError:
                cells.add(text)
                continue
            if max_row > 10000:
                max_row = max(min_row, 200)
            for row in range(min_row, max_row + 1):
                for col in range(min_col, max_col + 1):
                    cells.add(f"{get_column_letter(col)}{row}")
    return cells


class _VoptValidationRanges(list):
    def __init__(self, ws, ranges):
        super().__init__(ranges)
        self.ws = ws

    def __contains__(self, expected):
        expected = str(expected).replace('$', '')
        if super().__contains__(expected):
            return True
        try:
            min_col, min_row, max_col, max_row = range_boundaries(expected)
        except ValueError:
            return super().__contains__(expected)
        target = {f"{get_column_letter(col)}{row}" for row in range(min_row, max_row + 1) for col in range(min_col, max_col + 1)}
        return target.issubset(_vopt_validated_cells(self.ws))


def _vopt_dv_ranges(ws):
    ranges = []
    for dv in ws.data_validations.dataValidation:
        ranges.extend(str(rng).replace('$', '') for rng in dv.cells.ranges)
    return _VoptValidationRanges(ws, ranges)


def _vopt_ordered_subset(expected, actual):
    actual_iter = iter([_vopt_norm_text(item) for item in actual])
    for item in [_vopt_norm_text(value) for value in expected]:
        for candidate in actual_iter:
            if item == candidate:
                break
        else:
            return False
    return True



def _pdf_path():
    return Path(os.environ.get('PDF_OUTPUT_PATH', EXPECT['packet_output']))


def _doc_path():
    return Path(os.environ.get('DOCX_OUTPUT_PATH', EXPECT['document_output']))


def _pdf_texts(path):
    with pdfplumber.open(str(path)) as pdf:
        return [page.extract_text() or '' for page in pdf.pages]


def _norm_lower(value):
    return _vopt_norm_text(value).lower()


def _has_terms(text, terms):
    lowered = _norm_lower(text)
    return all(term.lower() in lowered for term in terms)


def _has_toc(path):
    with zipfile.ZipFile(path) as zf:
        xml = zf.read('word/document.xml').decode('utf-8', errors='ignore')
    return bool(re.search(r'TOC\s+(?:\\)?o|TOC\s*(?:&quot;|")', xml))


def _headings(doc):
    out = []
    for p in doc.paragraphs:
        style = p.style.name if p.style else ''
        if p.text.strip() and (style.startswith('Heading') or style == 'Title'):
            out.append(p.text.strip())
    return out


def _doc_text(doc):
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append('|'.join(cell.text.strip() for cell in row.cells))
    return '\n'.join(parts)


def _table_rows(table):
    return [[cell.text.strip() for cell in row.cells] for row in table.rows]


def _highlighted(doc, text):
    for p in doc.paragraphs:
        if _vopt_norm_text(text) in _vopt_norm_text(p.text):
            return any(run.font.highlight_color is not None for run in p.runs) or '<w:highlight' in p._p.xml
    return False


def _has_visual_highlight_xml(xml):
    if '<w:highlight' in xml:
        return True
    fills = re.findall(r'<w:shd\b[^>]*\bw:fill="([^"]+)"', xml)
    return any(fill.upper() not in {'', 'AUTO', 'FFFFFF', 'FFFFFF00', '00000000'} for fill in fills)


def _has_visual_highlighted_terms(doc, term_groups):
    for p in doc.paragraphs:
        text = _norm_lower(p.text)
        if not text:
            continue
        has_highlight = any(run.font.highlight_color is not None for run in p.runs) or _has_visual_highlight_xml(p._p.xml)
        if has_highlight and all(any(term.lower() in text for term in group) for group in term_groups):
            return True
    for table in doc.tables:
        for row in table.rows:
            row_text = _norm_lower(' '.join(cell.text for cell in row.cells))
            if not row_text:
                continue
            has_highlight = any(_has_visual_highlight_xml(cell._tc.xml) for cell in row.cells)
            if has_highlight and all(any(term.lower() in row_text for term in group) for group in term_groups):
                return True
    return False


def _assert_semantic_headings(doc):
    headings = [_norm_lower(item) for item in _headings(doc)]
    all_heading_text = ' '.join(headings)
    assert 'kitchen' in all_heading_text and 'allergen' in all_heading_text, f'Missing kitchen allergen heading: {_headings(doc)!r}'
    assert any('control' in item or 'matrix' in item or 'sheet' in item for item in headings), f'Missing allergen control heading: {_headings(doc)!r}'


def _assert_menu_item_semantics(text):
    lower = _norm_lower(text)
    item_terms = {
        'M-101': ['cardamom', 'oat'],
        'M-118': ['sesame', 'soy'],
        'M-126': ['hazelnut', 'tree nuts', 'milk', 'wheat'],
        'M-140': ['staff family meal', 'peanut'],
    }
    for item_id, terms in item_terms.items():
        assert item_id.lower() in lower, f'{item_id} missing from DOCX'
        missing = [term for term in terms if term not in lower]
        assert not missing, f'{item_id} missing expected semantic terms {missing!r}'
    internal_only = 'kitchen only' in lower or 'internal' in lower or 'staff family meal only' in lower
    public_exclusion = (
        'do not publish' in lower
        or 'not publish' in lower
        or 'not printed' in lower
        or 'not for public' in lower
        or 'off public' in lower
    )
    assert 'm-140' in lower and internal_only and public_exclusion, 'M-140 internal-only boundary is missing'
    assert 'control' in lower or 'cross-contact' in lower or 'handling' in lower, 'Kitchen control/cross-contact information is missing'


def _assert_high_risk_visual_highlight(doc):
    risk_terms = [
        ['m-118', 'm-126', 'm-140', 'sesame', 'tree nuts', 'peanut', 'hazelnut'],
        ['high-risk', 'high risk', 'high', 'elevated', 'control', 'handling'],
    ]
    assert _has_visual_highlighted_terms(doc, risk_terms), 'No visually highlighted high-risk control found'


def test_outputs_exist():
    assert _pdf_path().exists()
    assert _doc_path().exists()
    assert _pdf_path().suffix.lower() == '.pdf'
    assert _doc_path().suffix.lower() == '.docx'


def test_public_pdf_menu_board():
    path = _pdf_path()
    reader = PdfReader(str(path))
    assert len(reader.pages) >= 1
    texts = _pdf_texts(path)
    full = '\n'.join(texts)
    assert _has_terms(full, ['cafe', 'allergen']) or _has_terms(full, ['allergen', 'menu']), 'Missing public allergen menu title'
    for item_id in EXPECT['public_ids']:
        assert item_id in full, f'{item_id} missing from public PDF'
    assert 'sesame' in _norm_lower(full) and 'soy' in _norm_lower(full), 'Missing guest-facing sesame/soy allergen note'
    forbid_any(full, EXPECT['pdf_forbidden'], 'pdf')
    outline = [title for _level, title in flatten_outline(reader.outline)]
    assert outline, 'Public PDF must contain bookmarks'
    outline_text = '\n'.join(outline)
    bookmark_hits = sum(1 for item_id in EXPECT['public_ids'] if item_id in outline_text)
    assert bookmark_hits >= 2 or any(_has_terms(title, ['cafe', 'allergen']) for title in outline), f'Bookmarks do not identify the public allergen menu: {outline!r}'


def test_internal_docx_allergen_sheet():
    doc = Document(_doc_path())
    assert doc.sections[0].orientation == WD_ORIENT.LANDSCAPE
    assert _has_toc(_doc_path())
    _assert_semantic_headings(doc)
    text = _doc_text(doc)
    forbid_any(text, EXPECT['doc_forbidden'], 'docx')
    _assert_menu_item_semantics(text)
    assert doc.tables, 'Kitchen allergen sheet should include a structured matrix/table'
    _assert_high_risk_visual_highlight(doc)


def test_cross_output_public_internal_boundary():
    pdf_text = '\n'.join(_pdf_texts(_pdf_path()))
    doc_text = _doc_text(Document(_doc_path()))
    for item_id in EXPECT['public_ids']:
        assert item_id in pdf_text
        assert item_id in doc_text
    for item_id in EXPECT['internal_ids']:
        assert item_id not in pdf_text
        assert item_id in doc_text
