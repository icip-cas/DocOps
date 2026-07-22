import json
import os
import re
from pathlib import Path

from docx import Document

from verifier_utils import (
    docx_footer_text,
    docx_has_page_break,
    docx_has_toc_field,
    docx_header_text,
    docx_texts,
    forbid_any,
    normalize_text,
    run_preflight,
)

META_PATH = Path(os.environ.get("TASK_METADATA_PATH", "/tests/task_metadata.json"))
if not META_PATH.exists():
    META_PATH = Path(__file__).parent / "task_metadata.json"
META = json.loads(META_PATH.read_text(encoding="utf-8"))
EXPECT = META["verifier_expectations"]
INPUT_PATH = Path(os.environ.get("INPUT_PATH", META["input_path"]))
OUTPUT_PATH = Path(os.environ.get("OUTPUT_PATH", META["output_path"]))


def heading_order(doc):
    out = []
    first = next((p.text.strip() for p in doc.paragraphs if p.text.strip()), "")
    if first:
        out.append(first)
    for p in doc.paragraphs:
        style = p.style.name if p.style else ""
        text = p.text.strip()
        if normalize_text(text) in {"table of contents", "toc"}:
            continue
        if text and (style == "Title" or style.startswith("Heading")):
            if not out or normalize_text(text) != normalize_text(out[-1]):
                out.append(text)
    return out


def comparable_text(value):
    text = normalize_text(str(value))
    text = text.replace("≥", ">=").replace("≤", "<=")
    text = text.replace("°f", "f")
    text = text.replace(";", " and ")
    text = text.replace("four hours", "4 hours")
    text = re.sub(r"\s*,\s*and\s+", " and ", text)
    text = re.sub(r"[.。;；:：]+$", "", text)
    text = re.sub(r"^below\s+(\d+\s*f)$", r"< \1", text)
    text = re.sub(r"^more than\s+(\d+\s+hours)$", r"> \1", text)
    text = re.sub(r"^(\d+\s*f) or higher$", r">= \1", text)
    text = re.sub(r"^aqi\s+(\d+) or higher$", r"aqi >= \1", text)
    return " ".join(text.split())


def canonical_header(value):
    text = comparable_text(value)
    aliases = {
        "public role": "role",
        "public phone number": "public phone",
        "on-hand quantity": "on hand",
        "on hand quantity": "on hand",
        "on-hand qty": "on hand",
        "minimum quantity": "minimum",
        "minimum qty": "minimum",
        "gap (shortage)": "gap",
        "escalation trigger": "condition",
    }
    return aliases.get(text, text)


def row_key_matches(actual, expected):
    actual_norm = comparable_text(actual)
    expected_norm = comparable_text(expected)
    if actual_norm == expected_norm:
        return True
    if expected_norm in actual_norm or actual_norm in expected_norm:
        return True
    return (actual_norm, expected_norm) in {
        ("overnight low", "overnight cold"),
        ("air quality (aqi)", "air quality"),
    }


def value_matches(header, actual, expected):
    actual_norm = comparable_text(actual)
    expected_norm = comparable_text(expected)
    if actual_norm == expected_norm:
        return True
    header_norm = canonical_header(header)
    if header_norm == "trigger":
        return row_key_matches(actual, expected)
    if header_norm in {"condition", "public action"}:
        return expected_norm in actual_norm or actual_norm in expected_norm
    if header_norm == "threshold" and expected_norm == "aqi >= 151" and actual_norm == ">= 151":
        return True
    if header_norm == "reorder action":
        actual_nums = re.findall(r"\d+", actual_norm)
        expected_nums = re.findall(r"\d+", expected_norm)
        if actual_nums != expected_nums:
            return False
        action_groups = [["order"], ["move"], ["request"], ["add", "charge"], ["no reorder"]]
        for group in action_groups:
            if any(word in expected_norm for word in group):
                return any(word in actual_norm or word.rstrip("e") in actual_norm for word in group)
    return False


def has_page_break_before_flexible(doc, heading_text):
    prev = None
    wanted = comparable_text(heading_text)
    for p in doc.paragraphs:
        if comparable_text(p.text.strip()) == wanted:
            if prev is None:
                return docx_has_page_break(p)
            return docx_has_page_break(prev) or docx_has_page_break(p)
        prev = p
    raise AssertionError(f"Heading not found for page-break check: {heading_text}")


def contains_all(text, phrases):
    norm = comparable_text(text)
    return all(comparable_text(phrase) in norm for phrase in phrases)


def assert_public_content(text):
    checks = {
        "public operations summary": ["public release", "cooling", "charging", "wellness-check"],
        "public locations": ["North Pier Library", "East Yard School", "Mosaic Senior Center"],
        "accessibility support": ["ASL interpretation", "large-print intake forms", "step-free entry"],
    }
    missing = [label for label, phrases in checks.items() if not contains_all(text, phrases)]
    assert not missing, f"release playbook: missing public content groups: {missing}"


def table_with_headers(doc, headers):
    wanted = [canonical_header(h) for h in headers]
    for table in doc.tables:
        if not table.rows:
            continue
        actual = [canonical_header(cell.text) for cell in table.rows[0].cells]
        if all(h in actual for h in wanted):
            return table, {h: actual.index(canonical_header(h)) for h in headers}
    raise AssertionError(f"Missing table with headers: {headers}")


def row_by_key(table, key_col, key_value):
    for row in table.rows[1:]:
        if row_key_matches(row.cells[key_col].text, str(key_value)):
            return row
    raise AssertionError(f"Missing row keyed by {key_value!r}")


def assert_table_exact(doc, expected_rows, key_header):
    table, cols = table_with_headers(doc, expected_rows[0])
    assert len(table.rows) == len(expected_rows), f"{key_header}: wrong row count"
    assert len(table.columns) >= len(expected_rows[0]), f"{key_header}: wrong column count"
    key_col = cols[key_header]
    for expected in expected_rows[1:]:
        row = row_by_key(table, key_col, expected[expected_rows[0].index(key_header)])
        for header, expected_value in zip(expected_rows[0], expected):
            actual = row.cells[cols[header]].text.strip()
            assert value_matches(header, actual, expected_value), (
                f"{key_header} row {expected[0]!r} {header}: expected {expected_value!r}, found {actual!r}"
            )


def test_output_exists_and_is_docx():
    run_preflight(META, INPUT_PATH, OUTPUT_PATH)
    assert OUTPUT_PATH != INPUT_PATH


def test_title_sections_and_public_text():
    doc = Document(OUTPUT_PATH)
    headings = heading_order(doc)
    if headings and normalize_text(headings[0]) == normalize_text(EXPECT["title"]):
        section_headings = headings[1:]
    else:
        raise AssertionError(f"Document title missing or wrong: {headings[:2]!r}")
    assert [normalize_text(h) for h in section_headings] == [normalize_text(h) for h in EXPECT["heading_order"]]
    text = "\n".join(docx_texts(doc))
    assert_public_content(text)
    forbid_any(text, EXPECT["forbidden_phrases"], "release playbook")


def test_native_tables_and_computed_shortages():
    doc = Document(OUTPUT_PATH)
    assert_table_exact(doc, EXPECT["activation_rows"], "Trigger")
    assert_table_exact(doc, EXPECT["role_rows"], "Role")
    assert_table_exact(doc, EXPECT["supply_rows"], "Item")
    assert_table_exact(doc, EXPECT["communications_rows"], "Step")
    assert_table_exact(doc, EXPECT["contact_rows"], "Public Role")


def test_appendix_toc_header_footer():
    doc = Document(OUTPUT_PATH)
    assert docx_has_toc_field(doc), "Expected a real Word TOC field"
    for heading in EXPECT["page_break_before"]:
        assert has_page_break_before_flexible(doc, heading), f"Missing page break before {heading}"
    header = docx_header_text(doc)
    footer = docx_footer_text(doc)
    for phrase in EXPECT["header_contains"]:
        assert phrase in header, f"Missing header phrase: {phrase}"
    for phrase in EXPECT["footer_contains"]:
        assert phrase in footer, f"Missing footer phrase: {phrase}"
    assert "DRAFT" not in header.upper()


def test_source_artifact_was_not_modified():
    source = Document(INPUT_PATH)
    source_text = "\n".join(docx_texts(source))
    assert "DRAFT - Community Resilience Field Manual Working Copy" in source_text
    assert "parking lot" in normalize_text(source_text)
