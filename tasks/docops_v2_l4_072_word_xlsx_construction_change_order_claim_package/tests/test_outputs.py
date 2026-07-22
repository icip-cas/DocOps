import json
import os
import re
import sys
import zipfile
from pathlib import Path
from datetime import date, datetime

from docx import Document
from openpyxl import load_workbook
from openpyxl.utils.cell import get_column_letter, range_boundaries

sys.path.insert(0, str(Path(__file__).parent))
from verifier_utils import *  # noqa: F401,F403

META_PATH = Path(os.environ.get('TASK_METADATA_PATH', '/tests/task_metadata.json'))
META = json.loads(META_PATH.read_text(encoding='utf-8'))
EXPECT = META['verifier_expectations']


def _vopt_norm_text(value):
    if value is None:
        return ''
    if isinstance(value, datetime):
        if value.hour or value.minute or value.second:
            return value.strftime('%Y-%m-%d %H:%M')
        return value.strftime('%Y-%m-%d')
    if isinstance(value, date):
        return value.strftime('%Y-%m-%d')
    if isinstance(value, float) and value.is_integer():
        value = int(value)
    text = str(value)
    text = text.replace('\u2013', '-').replace('\u2014', '-').replace('\u2212', '-')
    text = text.replace('\xa0', ' ')
    text = re.sub(r'\s+', ' ', text).strip()
    return text


def _vopt_norm_formula(value):
    return re.sub(r'\s+', '', _vopt_norm_text(value)).upper()


def _vopt_norm_key(value):
    return re.sub(r'[^A-Z0-9]+', '', _vopt_norm_text(value).upper())


def _vopt_sheet_has_formula_tokens(ws, tokens, label):
    for row in ws.iter_rows():
        for cell in row:
            value = cell.value
            if isinstance(value, str) and value.startswith('='):
                norm = _vopt_norm_formula(value)
                if all(_vopt_norm_formula(token) in norm for token in tokens):
                    return
    raise AssertionError(f"{label}: no formula found with purpose tokens {tokens!r}")


def _vopt_assert_formula(actual, expected, label='formula'):
    assert _vopt_norm_formula(actual) == _vopt_norm_formula(expected), f"{label}: expected {expected!r}, found {actual!r}"


def _vopt_norm_rows(rows):
    return [[_vopt_norm_text(cell) for cell in row] for row in rows]


def _vopt_rows_equal(actual, expected):
    return _vopt_norm_rows(actual) == _vopt_norm_rows(expected)


def _vopt_rows_text(rows):
    return '\n'.join('|'.join(_vopt_norm_text(cell) for cell in row) for row in rows)


def _vopt_header_has(header, required):
    actual = {_vopt_norm_key(cell) for cell in header}
    missing = [
        cell for cell in required
        if not any(_vopt_norm_key(cell) in item or item in _vopt_norm_key(cell) for item in actual)
    ]
    assert not missing, f"Missing expected header columns {missing!r}; actual={header!r}"


def _vopt_section_text(section_part):
    return '\n'.join(p.text for p in getattr(section_part, 'paragraphs', []) if p.text is not None)


def _vopt_header_text(doc):
    return '\n'.join(_vopt_section_text(section.header) for section in doc.sections)


def _vopt_footer_text(doc):
    return '\n'.join(_vopt_section_text(section.footer) for section in doc.sections)


def _vopt_table_ref(ws, expected_name=None):
    tables = ws.tables
    if expected_name:
        for name in tables.keys():
            if str(name).lower() == str(expected_name).lower():
                return tables[name].ref
    vals = list(tables.values())
    assert vals, f"No Excel native table found on sheet {ws.title!r}"
    return vals[0].ref


def _vopt_table(ws, expected_name):
    for name in ws.tables.keys():
        if str(name).lower() == str(expected_name).lower():
            return ws.tables[name]
    raise AssertionError(f"No Excel native table named {expected_name!r} found on sheet {ws.title!r}")


def _vopt_table_has_size(ws, expected_name, min_rows, min_cols):
    table = _vopt_table(ws, expected_name)
    min_col, min_row, max_col, max_row = range_boundaries(table.ref)
    actual_rows = max_row - min_row + 1
    actual_cols = max_col - min_col + 1
    assert actual_rows >= min_rows and actual_cols >= min_cols, (
        f"{ws.title}.{expected_name}: expected at least {min_rows} rows x {min_cols} cols, "
        f"found {actual_rows} rows x {actual_cols} cols at {table.ref}"
    )
    return table


def _vopt_table_rows(ws, table_name):
    table = _vopt_table(ws, table_name)
    min_col, min_row, max_col, max_row = range_boundaries(table.ref)
    return [
        [ws.cell(row, col).value for col in range(min_col, max_col + 1)]
        for row in range(min_row, max_row + 1)
    ]


def _sheet(wb, name):
    assert name in wb.sheetnames, f"Missing required worksheet {name!r}; found {wb.sheetnames!r}"
    return wb[name]


def _vopt_clean_area(value):
    text = _vopt_norm_text(value).replace("'", '').replace('$', '').replace(' ', '')
    if '!' in text:
        text = text.split('!', 1)[1]
    return text.upper()


def _vopt_assert_print_area(ws, expected):
    actual = _vopt_clean_area(ws.print_area)
    target = _vopt_clean_area(expected)
    if target in actual or actual in target:
        return
    try:
        actual_min_col, actual_min_row, actual_max_col, actual_max_row = range_boundaries(actual)
        target_min_col, target_min_row, target_max_col, target_max_row = range_boundaries(target)
    except ValueError:
        raise AssertionError(f"{ws.title}: expected print area {expected!r}, found {ws.print_area!r}")
    assert (
        actual_min_col <= target_min_col
        and actual_min_row <= target_min_row
        and actual_max_col >= target_max_col
        and actual_max_row >= target_max_row
    ), f"{ws.title}: expected print area {expected!r}, found {ws.print_area!r}"


def _vopt_assert_hidden_strategy(text):
    norm = _vopt_norm_text(text).casefold()
    assert re.search(r'\b7\b.{0,20}\bdays\b', norm), 'hidden internal strategy: missing 7-day opening position'
    assert '18% markup' in norm, 'hidden internal strategy: missing 18% markup'
    assert 'settlement anchor' in norm, 'hidden internal strategy: missing settlement anchor'


def _vopt_freeze_pane(value):
    return getattr(value, 'coordinate', value)


def _vopt_validated_cells(ws):
    cells = set()
    for dv in ws.data_validations.dataValidation:
        for rng in dv.cells.ranges:
            text = str(rng).replace('$', '')
            try:
                min_col, min_row, max_col, max_row = range_boundaries(text)
            except ValueError:
                cells.add(text)
                continue
            if max_row > 10000:
                max_row = max(min_row, 200)
            for row in range(min_row, max_row + 1):
                for col in range(min_col, max_col + 1):
                    cells.add(f"{get_column_letter(col)}{row}")
    return cells


class _VoptValidationRanges(list):
    def __init__(self, ws, ranges):
        super().__init__(ranges)
        self.ws = ws

    def __contains__(self, expected):
        expected = str(expected).replace('$', '')
        if super().__contains__(expected):
            return True
        try:
            min_col, min_row, max_col, max_row = range_boundaries(expected)
        except ValueError:
            return super().__contains__(expected)
        target = {f"{get_column_letter(col)}{row}" for row in range(min_row, max_row + 1) for col in range(min_col, max_col + 1)}
        return target.issubset(_vopt_validated_cells(self.ws))


def _vopt_dv_ranges(ws):
    ranges = []
    for dv in ws.data_validations.dataValidation:
        ranges.extend(str(rng).replace('$', '') for rng in dv.cells.ranges)
    return _VoptValidationRanges(ws, ranges)



def _path(kind):
    env = {'docx': 'DOCX_OUTPUT_PATH', 'xlsx': 'XLSX_OUTPUT_PATH'}[kind]
    key = {'docx': 'docx_output', 'xlsx': 'xlsx_output'}[kind]
    return Path(os.environ.get(env, EXPECT[key]))


def _doc_text(doc):
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append('|'.join(cell.text.strip() for cell in row.cells))
    return '\n'.join(parts)


def _doc_heading_order(doc):
    out = []
    for p in doc.paragraphs:
        style = p.style.name if p.style else ''
        if p.text.strip() and (style.startswith('Heading') or style == 'Title'):
            out.append(p.text.strip())
    return out


def _norm_heading_label(value):
    text = str(value or '')
    text = text.replace('\u2013', '-').replace('\u2014', '-').replace('\xa0', ' ')
    text = re.sub(r'\s+', ' ', text).strip()
    text = re.sub(r'^(?:\d+|[IVXLCDM]+)[\.)]\s+', '', text, flags=re.I)
    text = re.sub(r'^[A-Z][\.)]\s+', '', text)
    text = re.sub(r'[:\-]+$', '', text).strip()
    return text.casefold()


def _heading_sequence_matches(actual, expected):
    actual_norm = [_norm_heading_label(item) for item in actual]
    expected_norm = [_norm_heading_label(item) for item in expected]
    pos = 0
    for item in actual_norm:
        if pos < len(expected_norm) and item == expected_norm[pos]:
            pos += 1
    return pos == len(expected_norm)


def _doc_table_rows(table):
    return [[cell.text.strip() for cell in row.cells] for row in table.rows]


def _has_toc_field(path):
    with zipfile.ZipFile(path) as zf:
        xml = zf.read('word/document.xml').decode('utf-8', errors='ignore')
    return bool(re.search(r'TOC\s+(?:\\)?o|TOC\s*(?:&quot;|")', xml))


def _highlighted_para(doc, text):
    for p in doc.paragraphs:
        if _vopt_norm_text(text) in _vopt_norm_text(p.text):
            return docx_para_has_highlight(p)
    raise AssertionError(f'Missing paragraph: {text}')


def _highlighted_critical_path_para(doc):
    for p in doc.paragraphs:
        norm = _vopt_norm_text(p.text)
        has_baseline = '2026-09-15' in norm or 'September 15, 2026' in norm
        has_revised = '2026-09-18' in norm or 'September 18, 2026' in norm
        if all(item in norm for item in ['OD-014', 'RFI-221', '3']) and has_baseline and has_revised:
            assert docx_para_has_highlight(p), f"Critical-path paragraph should be highlighted: {p.text!r}"
            return
    raise AssertionError('Missing highlighted critical-path paragraph connecting OD-014/RFI-221 to 3 days and 2026-09-18')


def _vopt_value_by_header(rows, key, header_name):
    header = [_vopt_norm_key(cell) for cell in rows[0]]
    target = _vopt_norm_key(header_name)
    assert target in header, f"Missing column {header_name!r}; header={rows[0]!r}"
    idx = header.index(target)
    for row in rows[1:]:
        if _vopt_norm_key(row[0]) == _vopt_norm_key(key):
            return row[idx]
    raise AssertionError(f"Missing row keyed by {key!r}")


def _vopt_has_any(value, choices):
    text = _vopt_norm_text(value).casefold()
    return any(choice.casefold() in text for choice in choices)


def _vopt_assert_core_entitlement_rows(rows):
    ids = [_vopt_norm_text(row[0]) for row in rows]
    assert ids == ['EV-001', 'EV-002', 'EV-003', 'EV-004', 'EV-005'], f"Unexpected entitlement IDs: {ids!r}"
    for evidence_id in ['EV-001', 'EV-002', 'EV-003']:
        row_text = _vopt_norm_text('|'.join(str(cell) for cell in rows[ids.index(evidence_id)]))
        assert 'compensable' in row_text.casefold(), f"{evidence_id}: must be compensable"
        row = rows[ids.index(evidence_id)]
        assert _vopt_has_any(row[6], ['Public'])
        assert _vopt_has_any(row[5], ['Yes', 'Include'])
    ev004 = rows[ids.index('EV-004')]
    ev004_text = _vopt_norm_text('|'.join(str(cell) for cell in ev004)).casefold()
    assert 'weather' in ev004_text and ('non-compensable' in ev004_text or 'exclude' in ev004_text), 'EV-004 must be weather/non-compensable'
    assert _vopt_has_any(ev004[5], ['No', 'Exclude']) and _vopt_has_any(ev004[6], ['Public'])
    ev005 = rows[ids.index('EV-005')]
    assert _vopt_has_any(ev005[4], ['Internal Only', 'Internal']) and _vopt_has_any(ev005[5], ['No', 'Exclude']) and _vopt_has_any(ev005[6], ['Internal'])


def _vopt_assert_core_cost_rows(rows):
    expected = {
        'LAB-014': (128, 95, 'compensable'),
        'MAT-014': (1, 18440, 'compensable'),
        'LAB-221': (72, 95, 'compensable'),
        'SUB-017': (1, 5250, 'compensable'),
        'EQ-003': (1, 5160, 'compensable'),
        'WTH-001': (1, None, 'excluded'),
    }
    by_code = {_vopt_norm_text(row[0]): row for row in rows}
    assert set(by_code) == set(expected), f"Unexpected cost codes: {sorted(by_code)!r}"
    compensable_total = 0
    for code, (qty, rate, treatment) in expected.items():
        row = by_code[code]
        assert row[4] == qty, f"{code}: unexpected quantity {row[4]!r}"
        if rate is not None:
            assert row[5] == rate, f"{code}: unexpected rate {row[5]!r}"
            compensable_total += qty * rate
        treatment_text = _vopt_norm_text(row[7]).casefold()
        if treatment == 'compensable':
            assert 'compensable' in treatment_text or 'include' in treatment_text, f"{code}: must be included/compensable"
        else:
            assert 'exclude' in treatment_text or 'weather' in treatment_text, f"{code}: weather row must be excluded"
    assert compensable_total == 47850


def _vopt_assert_core_schedule_rows(rows):
    text = _vopt_rows_text(rows)
    require_all(text, ['2026-09-15', '2026-09-18'], 'schedule impact')
    assert 'weather' in text.casefold() and '0' in text, 'schedule impact: weather row must carry 0 compensable days'
    assert any(str(row[4]) == '3' and _vopt_has_any(row[5], ['Compensable', 'Include', 'revised substantial completion']) for row in rows), (
        'schedule impact: missing 3 compensable calendar days'
    )


def _find_doc_table_with_header(doc, required):
    for table in doc.tables:
        rows = _doc_table_rows(table)
        if rows:
            header = {_vopt_norm_key(cell) for cell in rows[0]}
            if all(any(_vopt_norm_key(req) in item or item in _vopt_norm_key(req) for item in header) for req in required):
                return rows
    raise AssertionError(f'Missing DOCX table with columns {required!r}')


def _require_change_order_semantics(text):
    require_all(text, ['OD-014', 'RFI-221', 'bypass piping', 'valve relocation'], 'change order request')
    require_all(text, ['$52,635', '$47,850', '10%', '2026-09-15', '2026-09-18'], 'change order request')
    norm = _vopt_norm_text(text).casefold()
    assert 'weather' in norm and ('excluded' in norm or 'exclude' in norm), 'change order request: weather disruption must be excluded'
    assert '3' in norm and 'compensable' in norm and ('extension' in norm or 'calendar day' in norm), (
        'change order request: missing 3-day compensable time-extension semantics'
    )
    assert 'change_order_cost_schedule_workbook.xlsx' in text, 'change order request: missing workbook reconciliation reference'


def _row_values(ws, row, max_col):
    return [ws.cell(row, col).value for col in range(1, max_col + 1)]


def _dv_ranges(ws):
    return _vopt_dv_ranges(ws)


def test_outputs_exist():
    assert _path('docx').exists(), f"Missing DOCX output: {_path('docx')}"
    assert _path('xlsx').exists(), f"Missing XLSX output: {_path('xlsx')}"
    assert _path('docx').suffix.lower() == '.docx'
    assert _path('xlsx').suffix.lower() == '.xlsx'


def test_docx_change_order_structure_amounts_and_boundary():
    docx = _path('docx')
    doc = Document(docx)
    assert _heading_sequence_matches(_doc_heading_order(doc), EXPECT['doc_heading_order']), f"Unexpected heading order: {_doc_heading_order(doc)!r}"
    text = _doc_text(doc)
    _require_change_order_semantics(text)
    forbid_any(text, EXPECT['doc_forbidden_text'], 'change order request')
    assert EXPECT['doc_header_contains'] in _vopt_header_text(doc)
    assert EXPECT['doc_footer_contains'] in _vopt_footer_text(doc)
    assert _has_toc_field(docx)
    _highlighted_critical_path_para(doc)
    assert doc.tables, 'change order request: missing required evidence table'
    rows = _doc_table_rows(doc.tables[0])
    _vopt_header_has(rows[0], ['Evidence ID', 'Entitlement'])
    table_text = _vopt_rows_text(rows[1:])
    for evidence_id in ['EV-001', 'EV-002', 'EV-003', 'EV-004']:
        assert evidence_id in table_text
    assert 'EV-005' not in table_text and 'INT-CLAIM' not in table_text


def test_workbook_structure_controls_and_formulas():
    wb = load_workbook(_path('xlsx'), data_only=False)
    assert wb.sheetnames == EXPECT['sheet_order']
    min_table_rows = {
        'Entitlement Matrix': 6,
        'Cost Summary': 7,
        'Schedule Impact': 4,
        'Evidence Index': 6,
        'Public Backup': 2,
    }
    for sheet_name, (table_name, ref) in EXPECT['tables'].items():
        min_col, min_row, max_col, max_row = range_boundaries(ref)
        table = _vopt_table_has_size(_sheet(wb, sheet_name), table_name, min_table_rows[sheet_name], max_col - min_col + 1)
        _vopt_assert_print_area(wb[sheet_name], table.ref)
    for sheet in EXPECT['hidden_sheets']:
        assert wb[sheet].sheet_state in ('hidden', 'veryHidden')
    names = {dn.name for dn in wb.defined_names.values()}
    for name in EXPECT['defined_names']:
        assert name in names
    _vopt_sheet_has_formula_tokens(_sheet(wb, 'Cost Summary'), ['IF'], 'compensable-cost inclusion formula')
    _vopt_sheet_has_formula_tokens(_sheet(wb, 'Cost Summary'), ['SUM'], 'direct-cost total formula')
    _vopt_sheet_has_formula_tokens(_sheet(wb, 'Cost Summary'), ['0.10'], '10 percent markup formula')
    _vopt_sheet_has_formula_tokens(_sheet(wb, 'Schedule Impact'), ['SUM'], 'compensable-days total formula')
    for sheet_name, ranges in EXPECT['data_validation_ranges'].items():
        actual = _dv_ranges(wb[sheet_name])
        for expected in ranges:
            assert expected in actual
    assert len(wb['Cost Summary'].conditional_formatting) >= 1


def test_workbook_values_and_public_private_boundary():
    wb = load_workbook(_path('xlsx'), data_only=False)
    ent_rows = _vopt_table_rows(_sheet(wb, 'Entitlement Matrix'), 'entitlement_matrix')[1:]
    _vopt_assert_core_entitlement_rows(ent_rows)
    cost_rows = [row[:8] for row in _vopt_table_rows(_sheet(wb, 'Cost Summary'), 'cost_summary')[1:7]]
    _vopt_assert_core_cost_rows(cost_rows)
    sched_rows = _vopt_table_rows(_sheet(wb, 'Schedule Impact'), 'schedule_impact')[1:]
    _vopt_assert_core_schedule_rows(sched_rows)
    public_text = '\n'.join(str(cell.value) for sheet in ['Cost Summary', 'Schedule Impact', 'Public Backup'] for row in wb[sheet].iter_rows() for cell in row if cell.value)
    forbid_any(public_text, EXPECT['forbidden_public'], 'public workbook sheets')
    evidence_rows = _vopt_table_rows(_sheet(wb, 'Evidence Index'), 'evidence_index')[1:]
    internal_evidence = [row[0] for row in evidence_rows if _vopt_has_any(row[3], ['Internal'])]
    assert internal_evidence and set(internal_evidence).issubset({'EV-005', 'INT-CLAIM'})
    internal_text = '\n'.join(str(cell.value) for row in wb['Internal Strategy'].iter_rows() for cell in row if cell.value)
    _vopt_assert_hidden_strategy(internal_text)


def test_cross_output_change_order_consistency():
    doc_text = _doc_text(Document(_path('docx')))
    wb = load_workbook(_path('xlsx'), data_only=False)
    assert '$52,635' in doc_text
    _require_change_order_semantics(doc_text)
    _vopt_sheet_has_formula_tokens(_sheet(wb, 'Cost Summary'), ['SUM'], 'total request formula')
    _vopt_sheet_has_formula_tokens(_sheet(wb, 'Schedule Impact'), ['SUM'], 'schedule total formula')
    for evidence_id in ['EV-001', 'EV-002', 'EV-003', 'EV-004']:
        assert evidence_id in doc_text
    for forbidden in EXPECT['forbidden_public']:
        assert normalize_text(forbidden) not in normalize_text(doc_text)
