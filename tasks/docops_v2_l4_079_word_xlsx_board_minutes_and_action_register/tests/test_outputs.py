import json
import os
import re
import sys
import zipfile
from pathlib import Path
from datetime import date, datetime

from docx import Document
from openpyxl import load_workbook
from openpyxl.utils.cell import get_column_letter, range_boundaries

sys.path.insert(0, str(Path(__file__).parent))
from verifier_utils import *  # noqa: F401,F403

META_PATH = Path(os.environ.get('TASK_METADATA_PATH', '/tests/task_metadata.json'))
META = json.loads(META_PATH.read_text(encoding='utf-8'))
INPUT_PATH = Path(os.environ.get('INPUT_PATH', META['input_path']))
EXPECT = META['verifier_expectations']


def _vopt_norm_text(value):
    if value is None:
        return ''
    if isinstance(value, datetime):
        if value.hour or value.minute or value.second:
            return value.strftime('%Y-%m-%d %H:%M')
        return value.strftime('%Y-%m-%d')
    if isinstance(value, date):
        return value.strftime('%Y-%m-%d')
    if isinstance(value, float) and value.is_integer():
        value = int(value)
    text = str(value)
    text = text.replace('\u2013', '-').replace('\u2014', '-').replace('\u2212', '-')
    text = text.replace('\xa0', ' ')
    text = re.sub(r'\s+', ' ', text).strip()
    return text


def _vopt_norm_formula(value):
    return re.sub(r'\s+', '', _vopt_norm_text(value)).upper()


def _vopt_assert_formula(actual, expected, label='formula'):
    assert _vopt_norm_formula(actual) == _vopt_norm_formula(expected), f"{label}: expected {expected!r}, found {actual!r}"


def _vopt_norm_rows(rows):
    return [[_vopt_norm_text(cell) for cell in row] for row in rows]


def _vopt_rows_equal(actual, expected):
    return _vopt_norm_rows(actual) == _vopt_norm_rows(expected)


def _vopt_rows_text(rows):
    return '\n'.join('|'.join(_vopt_norm_text(cell) for cell in row) for row in rows)


def _vopt_header_has(header, required):
    actual = {_vopt_norm_text(cell).lower() for cell in header}
    missing = [cell for cell in required if _vopt_norm_text(cell).lower() not in actual]
    assert not missing, f"Missing expected header columns {missing!r}; actual={header!r}"


def _vopt_section_text(section_part):
    return '\n'.join(p.text for p in getattr(section_part, 'paragraphs', []) if p.text is not None)


def _vopt_header_text(doc):
    return '\n'.join(_vopt_section_text(section.header) for section in doc.sections)


def _vopt_footer_text(doc):
    return '\n'.join(_vopt_section_text(section.footer) for section in doc.sections)


def _vopt_table_ref(ws, expected_name=None):
    tables = ws.tables
    if expected_name:
        for name in tables.keys():
            if str(name).lower() == str(expected_name).lower():
                return tables[name].ref
    vals = list(tables.values())
    assert vals, f"No Excel native table found on sheet {ws.title!r}"
    return vals[0].ref


def _vopt_clean_area(value):
    text = _vopt_norm_text(value).replace("'", '').replace('$', '').replace(' ', '')
    if '!' in text:
        text = text.split('!', 1)[1]
    return text.upper()


def _vopt_assert_print_area(ws, expected):
    actual = _vopt_clean_area(ws.print_area)
    target = _vopt_clean_area(expected)
    assert target in actual or actual in target, f"{ws.title}: expected print area {expected!r}, found {ws.print_area!r}"


def _vopt_freeze_pane(value):
    return getattr(value, 'coordinate', value)


def _vopt_validated_cells(ws):
    cells = set()
    for dv in ws.data_validations.dataValidation:
        for rng in dv.cells.ranges:
            text = str(rng).replace('$', '')
            try:
                min_col, min_row, max_col, max_row = range_boundaries(text)
            except ValueError:
                cells.add(text)
                continue
            if max_row > 10000:
                max_row = max(min_row, 200)
            for row in range(min_row, max_row + 1):
                for col in range(min_col, max_col + 1):
                    cells.add(f"{get_column_letter(col)}{row}")
    return cells


class _VoptValidationRanges(list):
    def __init__(self, ws, ranges):
        super().__init__(ranges)
        self.ws = ws

    def __contains__(self, expected):
        expected = str(expected).replace('$', '')
        if super().__contains__(expected):
            return True
        try:
            min_col, min_row, max_col, max_row = range_boundaries(expected)
        except ValueError:
            return super().__contains__(expected)
        target = {f"{get_column_letter(col)}{row}" for row in range(min_row, max_row + 1) for col in range(min_col, max_col + 1)}
        return target.issubset(_vopt_validated_cells(self.ws))


def _vopt_dv_ranges(ws):
    ranges = []
    for dv in ws.data_validations.dataValidation:
        ranges.extend(str(rng).replace('$', '') for rng in dv.cells.ranges)
    return _VoptValidationRanges(ws, ranges)


def _vopt_ordered_subset(expected, actual):
    actual_iter = iter([_vopt_norm_text(item) for item in actual])
    for item in [_vopt_norm_text(value) for value in expected]:
        for candidate in actual_iter:
            if item == candidate:
                break
        else:
            return False
    return True



def _resolve_output(kind):
    if kind == 'minutes':
        return Path(os.environ.get('DOCX_OUTPUT_PATH', EXPECT['minutes_output']))
    if kind == 'register':
        return Path(os.environ.get('XLSX_OUTPUT_PATH', EXPECT['register_output']))
    raise KeyError(kind)


def _doc_text(doc):
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())
    for table in doc.tables:
        for row in table.rows:
            parts.append('|'.join(cell.text.strip() for cell in row.cells))
    return '\n'.join(parts)


def _heading_order(doc):
    out = []
    for p in doc.paragraphs:
        style = p.style.name if p.style else ''
        if p.text.strip() and (style.startswith('Heading') or style == 'Title'):
            out.append(p.text.strip())
    return out


def _table_rows(table):
    return [[cell.text.strip() for cell in row.cells] for row in table.rows]


def _has_toc_field(path):
    with zipfile.ZipFile(path) as zf:
        xml = zf.read('word/document.xml').decode('utf-8', errors='ignore')
    return bool(re.search(r'TOC\s+(?:\\)?o|TOC\s*(?:&quot;|")', xml))


def _has_break_before_appendix(path):
    with zipfile.ZipFile(path) as zf:
        xml = zf.read('word/document.xml').decode('utf-8', errors='ignore')
    idx = xml.lower().find('appendix a')
    if idx < 0:
        return False
    before = xml[max(0, idx - 1800):idx]
    return '<w:sectPr' in before or '<w:br' in before and 'w:type="page"' in before


def _para_by_text(doc, text):
    for p in doc.paragraphs:
        if _vopt_norm_text(text) in _vopt_norm_text(p.text):
            return p
    raise AssertionError(f'Paragraph not found: {text}')


def _para_highlighted(p):
    return any(run.font.highlight_color is not None for run in p.runs) or '<w:highlight' in p._p.xml


def _assert_minutes_semantics(text):
    require_all(text, ['Final', 'Posting package'], 'minutes')
    require_all(text, ['M-17', 'Approved as amended', 'M-18', 'Deferred'], 'minutes')
    require_all(text, ['Priya Shah', 'Remote', 'A-18', 'Watch', 'counsel'], 'minutes')


def _watch_item_paragraph(doc):
    for p in doc.paragraphs:
        norm = normalize_text(p.text)
        if 'a-18' in norm and 'watch' in norm and 'counsel' in norm:
            return p
    raise AssertionError('Expected an A-18 Watch item paragraph for counsel review')


def _assert_source_motion_check_formula(formula, row, label):
    norm = _vopt_norm_formula(formula)
    assert norm.startswith('='), f'{label}: expected a live formula, found {formula!r}'
    assert f'F{row}' in norm, f'{label}: formula must check source motion cell F{row}, found {formula!r}'
    assert 'OK' in norm, f'{label}: formula must return OK for valid rows, found {formula!r}'
    assert any(token in norm for token in ('CHECK', 'MISMATCH', 'ERR', 'ERROR', 'FALSE', 'INVALID')), (
        f'{label}: formula must have a non-OK branch for invalid rows, found {formula!r}'
    )
    assert any(token in norm for token in ('COUNTIF', 'MATCH', 'XLOOKUP', 'VLOOKUP', 'RULES!')), (
        f'{label}: formula must validate against a list/table of allowed values, found {formula!r}'
    )


def test_outputs_exist():
    minutes = _resolve_output('minutes')
    register = _resolve_output('register')
    assert minutes.exists(), f'Missing minutes output: {minutes}'
    assert register.exists(), f'Missing register output: {register}'
    assert minutes.suffix.lower() == '.docx'
    assert register.suffix.lower() == '.xlsx'
    assert minutes.stat().st_size > 0
    assert register.stat().st_size > 0


def test_minutes_structure_and_content():
    path = _resolve_output('minutes')
    doc = Document(path)
    assert _heading_order(doc) == EXPECT['heading_order'], f"Unexpected headings: {_heading_order(doc)!r}"
    text = _doc_text(doc)
    _assert_minutes_semantics(text)
    for forbidden in EXPECT['forbidden_text']:
        assert forbidden not in text, f'Forbidden text still present: {forbidden}'
    assert _has_toc_field(path), 'Expected real TOC field in DOCX XML'
    assert len(doc.sections) >= EXPECT['section_count_min'] or _has_break_before_appendix(path), (
        'Expected Appendix A to start on a new section or page'
    )
    assert EXPECT['footer_contains'] in _vopt_footer_text(doc)
    assert _para_highlighted(_watch_item_paragraph(doc)), 'Expected highlighted Watch item paragraph'


def test_minutes_tables():
    doc = Document(_resolve_output('minutes'))
    assert len(doc.tables) >= 3, f'Expected at least 3 tables, found {len(doc.tables)}'
    assert _vopt_rows_equal(_table_rows(doc.tables[0])[1:], EXPECT['attendance_rows'])
    assert _vopt_rows_equal(_table_rows(doc.tables[1])[1:], EXPECT['motion_rows'])
    action_rows = [row[:6] for row in _table_rows(doc.tables[2])[1:]]
    assert _vopt_rows_equal(action_rows, EXPECT['action_rows'])


def test_register_workbook_structure_and_values():
    wb = load_workbook(_resolve_output('register'), data_only=False)
    for sheet in EXPECT['register_sheet_order']:
        assert sheet in wb.sheetnames, f'Missing expected sheet: {sheet}'
    ws = wb['Action Register']
    actual_rows = [[ws.cell(r, c).value for c in range(1, 7)] for r in range(2, 5)]
    assert _vopt_rows_equal(actual_rows, EXPECT['action_rows']), f"Unexpected action rows: {actual_rows!r}"
    assert _vopt_table_ref(ws, 'board_action_register') == EXPECT['register_table_ref']
    for cell in EXPECT['register_formula_cells']:
        _assert_source_motion_check_formula(ws[cell].value, int(cell[1:]), cell)
    for sheet in EXPECT['hidden_sheets']:
        assert wb[sheet].sheet_state in ('hidden', 'veryHidden')
    names = {dn.name for dn in wb.defined_names.values()}
    for name in EXPECT['defined_names']:
        assert name in names, f'Missing defined name: {name}'
    actual_ranges = _vopt_dv_ranges(ws)
    for expected in EXPECT['data_validation_ranges']['Action Register']:
        assert expected in actual_ranges, f'Missing data validation range: {expected}'


def test_cross_output_action_consistency():
    doc = Document(_resolve_output('minutes'))
    word_rows = [row[:6] for row in _table_rows(doc.tables[2])[1:]]
    wb = load_workbook(_resolve_output('register'), data_only=False)
    ws = wb['Action Register']
    xlsx_rows = [[ws.cell(r, c).value for c in range(1, 7)] for r in range(2, 5)]
    assert word_rows == xlsx_rows, f'Word and Excel action rows diverge: {word_rows!r} vs {xlsx_rows!r}'
