import json
import os
import re
from datetime import date, datetime
from pathlib import Path

from docx import Document
from openpyxl import load_workbook
from openpyxl.utils.cell import get_column_letter, range_boundaries

from verifier_utils import *  # noqa: F401,F403

META_PATH = Path(os.environ.get("TASK_METADATA_PATH", "/tests/task_metadata.json"))
META = json.loads(META_PATH.read_text(encoding="utf-8"))
EXPECT = META["verifier_expectations"]


def _norm(value):
    if value is None:
        return ""
    if isinstance(value, datetime):
        return value.strftime("%Y-%m-%d")
    if isinstance(value, date):
        return value.strftime("%Y-%m-%d")
    if isinstance(value, float) and value.is_integer():
        value = int(value)
    return re.sub(r"\s+", " ", str(value).replace("\u2013", "-").replace("\u2014", "-").replace("\xa0", " ")).strip()


def _norm_formula(value):
    return re.sub(r"\s+", "", _norm(value)).upper()


def _norm_key(value):
    return re.sub(r"[^a-z0-9]+", "", _norm(value).casefold())


def _norm_heading_label(value):
    text = _norm(value)
    text = re.sub(r"^(?:\d+|[IVXLCDM]+)[\.)]\s+", "", text, flags=re.I)
    text = re.sub(r"^[A-Z][\.)]\s+", "", text)
    text = re.sub(r"[:\-]+$", "", text).strip()
    return text.casefold()


def _path(kind):
    env = {"docx": "DOCX_OUTPUT_PATH", "xlsx": "XLSX_OUTPUT_PATH"}[kind]
    key = {"docx": "docx_output", "xlsx": "xlsx_output"}[kind]
    return Path(os.environ.get(env, EXPECT[key]))


def _doc_text(doc):
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append("|".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts)


def _heading_order(doc):
    out = []
    first = next((p.text.strip() for p in doc.paragraphs if p.text.strip()), "")
    if first:
        out.append(first)
    for p in doc.paragraphs:
        style = p.style.name if p.style else ""
        if p.text.strip() and (style == "Title" or style.startswith("Heading")):
            if not out or _norm_heading_label(p.text) != _norm_heading_label(out[-1]):
                out.append(p.text.strip())
    return out


def _heading_sequence_matches(actual, expected):
    actual_norm = [_norm_heading_label(item) for item in actual]
    expected_norm = [_norm_heading_label(item) for item in expected]
    return actual_norm == expected_norm


def _assert_required_sheets(wb, sheets):
    missing = [sheet for sheet in sheets if sheet not in wb.sheetnames]
    assert not missing, f"Missing required workbook sheets: {missing}; found {wb.sheetnames}"


def _sheet(wb, expected_name):
    _assert_required_sheets(wb, [expected_name])
    return wb[expected_name]


def _table(ws, expected_name):
    for name in ws.tables.keys():
        if str(name).lower() == str(expected_name).lower():
            return ws.tables[name]
    raise AssertionError(f"Missing table {expected_name!r} on {ws.title!r}")


def _table_bounds(table):
    return range_boundaries(table.ref.replace("$", ""))


def _table_headers(ws, table):
    min_col, min_row, max_col, _ = _table_bounds(table)
    return [_norm(ws.cell(min_row, col).value) for col in range(min_col, max_col + 1)]


def _table_header_map(ws, table):
    return {_norm_key(header): idx for idx, header in enumerate(_table_headers(ws, table), start=_table_bounds(table)[0])}


def _require_headers(ws, table, expected_headers):
    actual = {_norm_key(header) for header in _table_headers(ws, table)}
    missing = [header for header in expected_headers if _norm_key(header) not in actual]
    assert not missing, f"{ws.title}: table {table.name} missing headers {missing}; found {_table_headers(ws, table)}"


def _table_has_size(table, min_data_rows, min_cols):
    min_col, min_row, max_col, max_row = _table_bounds(table)
    assert max_col - min_col + 1 >= min_cols, f"{table.name}: expected at least {min_cols} columns"
    assert max_row - min_row >= min_data_rows, f"{table.name}: expected at least {min_data_rows} data rows"


def _table_rows(ws, table):
    min_col, min_row, _, max_row = _table_bounds(table)
    header_map = _table_header_map(ws, table)
    rows = []
    for row_idx in range(min_row + 1, max_row + 1):
        row = {}
        for key, col_idx in header_map.items():
            row[key] = ws.cell(row_idx, col_idx)
        if any(cell.value is not None for cell in row.values()):
            rows.append(row)
    return rows


def _cell_by_header(row, header):
    key = _norm_key(header)
    assert key in row, f"Missing header {header!r}"
    return row[key]


def _value_matches(actual, expected):
    actual_text = _norm(actual)
    expected_text = _norm(expected)
    if actual_text.casefold() == expected_text.casefold():
        return True
    if len(expected_text) >= 12 and expected_text.casefold() in actual_text.casefold():
        return True
    return False


FIELD_ALIASES = {
    "manuscript": {"manuscript", "manuscript title", "title"},
    "mapped citations": {"mapped citations", "mapped citation count", "mapped count"},
    "fixed cross-references": {
        "fixed cross-references",
        "fixed cross-reference count",
        "fixed cross-ref count",
        "fixed cross reference count",
    },
    "final references": {"final references", "final reference count", "reference count"},
    "private notes excluded": {"private notes excluded", "private notes excluded count", "private comments excluded"},
    "submission ready": {"submission ready", "submission-ready status", "ready status"},
}


def _key_value_matches(actual, expected):
    if _value_matches(actual, expected):
        return True
    expected_key = _norm_key(expected)
    actual_key = _norm_key(actual)
    aliases = {_norm_key(item) for item in FIELD_ALIASES.get(_norm(expected).casefold(), set())}
    if expected_key in aliases and actual_key in aliases:
        return True
    return False


def _find_sheet_row(rows, key_header, key_value):
    for row in rows:
        if _key_value_matches(_cell_by_header(row, key_header).value, key_value):
            return row
    raise AssertionError(f"Could not find row with {key_header!r}={key_value!r}")


def _assert_row_contains(rows, headers, expected_row, key_header=None):
    key_header = key_header or headers[0]
    key_value = expected_row[headers.index(key_header)]
    row = _find_sheet_row(rows, key_header, key_value)
    for header, expected_value in zip(headers, expected_row):
        actual = _cell_by_header(row, header).value
        assert _value_matches(actual, expected_value), f"{header}: expected {expected_value!r}, found {actual!r}"
    return row


def _range_covered(ws, expected):
    min_col, min_row, max_col, max_row = range_boundaries(expected.replace("$", ""))
    target = {f"{get_column_letter(col)}{row}" for row in range(min_row, max_row + 1) for col in range(min_col, max_col + 1)}
    cells = set()
    for dv in ws.data_validations.dataValidation:
        for rng in dv.cells.ranges:
            min_c, min_r, max_c, max_r = range_boundaries(str(rng).replace("$", ""))
            cells.update(f"{get_column_letter(col)}{row}" for row in range(min_r, max_r + 1) for col in range(min_c, max_c + 1))
    return target.issubset(cells)


def _clean_area(value):
    if isinstance(value, (list, tuple)):
        value = ",".join(str(item) for item in value)
    text = _norm(value).replace("'", "").replace("$", "").replace(" ", "")
    if "!" in text:
        text = text.split("!", 1)[1]
    return text.upper()


def _sheet_text(wb, sheets):
    _assert_required_sheets(wb, sheets)
    parts = []
    for sheet in sheets:
        for row in wb[sheet].iter_rows():
            for cell in row:
                if cell.value is not None:
                    parts.append(str(cell.value))
    return "\n".join(parts)


def _defined_name_names(wb):
    try:
        values = wb.defined_names.values()
    except AttributeError:
        values = getattr(wb.defined_names, "definedName", [])
    return {dn.name for dn in values if getattr(dn, "name", None)}


def _formula_with_linked_cell(wb, value):
    text = _norm(value)
    linked = ""
    match = re.match(r"^=\s*'([^']+)'!\$?([A-Z]+)\$?(\d+)\s*$", text, flags=re.I)
    if not match:
        match = re.match(r"^=\s*([A-Za-z0-9_ ]+)!\$?([A-Z]+)\$?(\d+)\s*$", text, flags=re.I)
    if match and match.group(1) in wb.sheetnames:
        linked = _norm(wb[match.group(1)][f"{match.group(2)}{match.group(3)}"].value)
    return _norm_formula(f"{text} {linked}")


def _count_formula_ok(wb, value, sheet_name, table_name, status_or_count):
    formula = _norm_formula(value)
    assert formula.startswith("="), f"{sheet_name} summary value must be a formula"
    expanded = _formula_with_linked_cell(wb, value)
    refs_target = sheet_name.replace(" ", "").upper() in expanded.replace("'", "") or table_name.upper() in expanded
    assert refs_target, f"formula must reference {sheet_name} or {table_name}"
    has_count_logic = any(token in expanded for token in ("COUNTIF", "COUNTA", "COUNT(", "ROWS(", "SUBTOTAL("))
    if status_or_count:
        assert status_or_count.upper() in expanded or has_count_logic, f"formula must count {status_or_count}"
    else:
        assert has_count_logic, "reference-count formula must use count logic"


def _submission_ready_formula_ok(value):
    formula = _norm_formula(value)
    assert formula.startswith("="), "submission-ready control must be a formula"
    for token in ("READY", "HOLD"):
        assert token in formula, f"submission-ready formula missing {token}"
    assert "IF" in formula or "IFS" in formula, "submission-ready formula must be conditional"
    assert "AND" in formula or formula.count("=") >= 4, "submission-ready formula must combine all readiness checks"
    assert any(token in formula for token in ("B4=5", "MAPPEDCITATIONCOUNT=5", "MAPPEDCITATIONS")), "submission-ready formula must check five mapped citations"
    assert any(token in formula for token in ("B5=3", "FIXEDCROSSREFCOUNT=3", "FIXEDCROSS")), "submission-ready formula must check three cross-reference fixes"
    assert any(token in formula for token in ("B6=5", "FINALREFERENCE", "REFERENCELIST")), "submission-ready formula must check five final references"
    assert any(token in formula for token in ("B7=3", "PRIVATE", "EXCLUDED")), "submission-ready formula must check three excluded private notes"


def _status_validation_ok(ws, table_name, status_header):
    table = _table(ws, table_name)
    min_col, min_row, _, max_row = _table_bounds(table)
    col_idx = _table_header_map(ws, table)[_norm_key(status_header)]
    target = f"{get_column_letter(col_idx)}{min_row + 1}:{get_column_letter(col_idx)}{max_row}"
    assert _range_covered(ws, target), f"{ws.title}: missing validation over {target}"


def _summary_rows(wb):
    ws = _sheet(wb, "Submission Summary")
    table = _table(ws, "tblSubmissionSummary")
    return _table_rows(ws, table)


def _summary_value_cell(wb, field):
    return _cell_by_header(_find_sheet_row(_summary_rows(wb), "Field", field), "Value")


def _row_text(row):
    return " ".join(_norm(cell.value) for cell in row.values() if cell.value is not None)


def _assert_status_pass_or_fixed(value, expected):
    actual = _norm(value).casefold()
    assert actual in {_norm(expected).casefold(), "pass", "passed", "ok", "complete", "completed"}, f"unexpected status {value!r}"


def _assert_summary_values(wb):
    rows = _summary_rows(wb)
    manuscript = _find_sheet_row(rows, "Field", "Manuscript")
    assert _value_matches(_cell_by_header(manuscript, "Value").value, "Hybrid Clinic Follow-up Study")
    citation_style = _find_sheet_row(rows, "Field", "Citation style")
    style_value = _norm(_cell_by_header(citation_style, "Value").value).casefold()
    assert "numeric" in style_value and ("vancouver" in style_value or "bracket" in style_value)
    for field in ("Mapped citations", "Fixed cross-references", "Final references", "Submission ready"):
        row = _find_sheet_row(rows, "Field", field)
        value = _cell_by_header(row, "Value").value
        assert isinstance(value, str) and value.startswith("="), f"{field} must use a formula"
    private = _find_sheet_row(rows, "Field", "Private notes excluded")
    private_value = _cell_by_header(private, "Value").value
    if isinstance(private_value, str) and private_value.startswith("="):
        assert "PRIVATE" in _formula_with_linked_cell(wb, private_value)
    else:
        assert _norm(private_value) == "3", "private-notes-excluded count must be 3"


def _assert_crossref_values(wb):
    rows = _table_rows(_sheet(wb, "Cross-Reference Fixes"), _table(_sheet(wb, "Cross-Reference Fixes"), "tblCrossRefFixes"))
    requirements = [
        ("Cohort flow figure", "Figure 3", "Figure 1"),
        ("Visit completion table", "Table 2", "Table 1"),
        ("Supplementary Figure S1", "Supplementary Figure S1", "Removed"),
    ]
    for item, original, final in requirements:
        row = _find_sheet_row(rows, "Item", item)
        assert _value_matches(_cell_by_header(row, "Original Reference").value, original)
        final_value = _norm(_cell_by_header(row, "Final Reference").value)
        if final == "Removed":
            assert "removed" in final_value.casefold() or final_value in {"", "(removed)"}
        else:
            assert _value_matches(final_value, final)
        assert _norm(_cell_by_header(row, "Reason").value), f"{item} must document a reason"
        _assert_status_pass_or_fixed(_cell_by_header(row, "Status").value, "Fixed")


def _assert_qa_values(wb):
    rows = _table_rows(_sheet(wb, "Style QA"), _table(_sheet(wb, "Style QA"), "tblStyleQA"))
    combined = "\n".join(_row_text(row).casefold() for row in rows)
    checks = [
        ("citation", ["numeric", "first", "appearance"]),
        ("reference list", ["reference", "vancouver"]),
        ("figure/table numbering", ["figure", "table"]),
        ("supplement callout", ["supplement", "removed"]),
        ("private comments", ["private", "excluded"]),
    ]
    for label, tokens in checks:
        assert all(token in combined for token in tokens), f"Style QA missing {label} check"
    for row in rows:
        _assert_status_pass_or_fixed(_cell_by_header(row, "Status").value, "Pass")


def _assert_private_values(wb):
    rows = _table_rows(_sheet(wb, "Private Notes"), _table(_sheet(wb, "Private Notes"), "tblPrivateNotes"))
    assert len(rows) == 3, "Private Notes must retain exactly three excluded notes"
    combined = "\n".join(_row_text(row).casefold() for row in rows)
    for tokens in (("internal", "author", "note"), ("reviewer", "sensitive"), ("sample", "strength")):
        assert all(token in combined for token in tokens), f"Private Notes missing {' '.join(tokens)}"
    for row in rows:
        treatment = _norm(_cell_by_header(row, "Treatment").value).casefold()
        assert "exclude" in treatment or "excluded" in treatment


def test_outputs_exist():
    assert _path("docx").exists()
    assert _path("xlsx").exists()
    assert _path("docx").suffix.lower() == ".docx"
    assert _path("xlsx").suffix.lower() == ".xlsx"


def test_docx_citations_crossrefs_and_privacy():
    doc = Document(_path("docx"))
    assert _heading_sequence_matches(_heading_order(doc), EXPECT["doc_heading_order"])
    text = _doc_text(doc)
    non_reference_required = [item for item in EXPECT["doc_required"] if not re.match(r"^\d+\.\s+", item)]
    require_all(text, non_reference_required, "revised manuscript")
    for ref in EXPECT["reference_rows"][1:]:
        assert ref[2] in text, f"revised manuscript missing reference: {ref[2]}"
    forbid_any(text, EXPECT["doc_forbidden"], "revised manuscript")
    positions = [text.index(f"[{i}]") for i in range(1, 6)]
    assert positions == sorted(positions), "numeric citations are not in first-appearance order"


def test_workbook_structure_formulas_controls():
    wb = load_workbook(_path("xlsx"), data_only=False)
    assert wb.sheetnames == EXPECT["sheet_order"]
    for sheet, (table_name, _ref) in EXPECT["tables"].items():
        ws = _sheet(wb, sheet)
        table = _table(ws, table_name)
        expected_headers = EXPECT[f"{_norm_key(sheet).replace('crossreferencefixes', 'crossref').replace('privatenotes', 'private').replace('submissionsummary', 'summary').replace('citationmap', 'citation').replace('referencelist', 'reference').replace('styleqa', 'qa')}_rows"][0]
        _require_headers(ws, table, expected_headers)
        min_rows = len(EXPECT[f"{_norm_key(sheet).replace('crossreferencefixes', 'crossref').replace('privatenotes', 'private').replace('submissionsummary', 'summary').replace('citationmap', 'citation').replace('referencelist', 'reference').replace('styleqa', 'qa')}_rows"]) - 1
        if sheet == "Style QA":
            min_rows = 5
        _table_has_size(table, min_rows, len(expected_headers))
    for sheet in EXPECT["hidden_sheets"]:
        assert _sheet(wb, sheet).sheet_state in ("hidden", "veryHidden")
    names = _defined_name_names(wb)
    for name in EXPECT["defined_names"]:
        assert name in names
    _count_formula_ok(wb, _summary_value_cell(wb, "Mapped citations").value, "Citation Map", "tblCitationMap", "Mapped")
    _count_formula_ok(wb, _summary_value_cell(wb, "Fixed cross-references").value, "Cross-Reference Fixes", "tblCrossRefFixes", "Fixed")
    _count_formula_ok(wb, _summary_value_cell(wb, "Final references").value, "Reference List", "tblReferenceList", "")
    _submission_ready_formula_ok(_summary_value_cell(wb, "Submission ready").value)
    _status_validation_ok(_sheet(wb, "Citation Map"), "tblCitationMap", "Status")
    _status_validation_ok(_sheet(wb, "Cross-Reference Fixes"), "tblCrossRefFixes", "Status")
    _status_validation_ok(_sheet(wb, "Style QA"), "tblStyleQA", "Status")
    for sheet, area in EXPECT["print_areas"].items():
        actual = _clean_area(_sheet(wb, sheet).print_area)
        target = _clean_area(area)
        if sheet == "Style QA" and actual == "A1:C6":
            continue
        assert target in actual or actual in target


def test_workbook_values_and_privacy():
    wb = load_workbook(_path("xlsx"), data_only=False)
    _assert_required_sheets(wb, ["Submission Summary", "Citation Map", "Cross-Reference Fixes", "Reference List", "Style QA", "Private Notes"])
    _assert_summary_values(wb)
    for sheet, table_name, expected_rows, key_header in [
        ("Citation Map", "tblCitationMap", EXPECT["citation_rows"], "Citation No."),
        ("Reference List", "tblReferenceList", EXPECT["reference_rows"], "Citation No."),
    ]:
        ws = _sheet(wb, sheet)
        table = _table(ws, table_name)
        headers = expected_rows[0]
        rows = _table_rows(ws, table)
        for expected_row in expected_rows[1:]:
            _assert_row_contains(rows, headers, expected_row, key_header=key_header)
    _assert_crossref_values(wb)
    _assert_qa_values(wb)
    _assert_private_values(wb)
    private_count = _summary_value_cell(wb, "Private notes excluded").value
    if isinstance(private_count, str) and private_count.startswith("="):
        formula = _formula_with_linked_cell(wb, private_count)
        assert "PRIVATE" in formula and any(token in formula for token in ("COUNT", "COUNTA", "ROWS", "SUBTOTAL")), "private-notes-excluded formula must count private excluded notes"
    else:
        assert _norm(private_count) == "3", "private-notes-excluded count must be 3"
    public = _sheet_text(wb, ["Submission Summary", "Citation Map", "Cross-Reference Fixes", "Reference List", "Style QA"])
    forbid_any(
        public,
        ["do not mention reviewer 2", "reviewer-sensitive", "sample is weaker than planned", "sample-strength caveat"],
        "public workbook sheets",
    )


def test_cross_output_consistency():
    doc_text = _doc_text(Document(_path("docx")))
    wb = load_workbook(_path("xlsx"), data_only=False)
    _assert_required_sheets(wb, ["Reference List"])
    for anchor in ["[1]", "[2]", "[3]", "[4]", "[5]", "Figure 1", "Table 1"]:
        assert anchor in doc_text
    for ref in EXPECT["reference_rows"][1:]:
        assert ref[2] in doc_text
        assert ref[2] in _sheet_text(wb, ["Reference List"])
    _submission_ready_formula_ok(_summary_value_cell(wb, "Submission ready").value)
