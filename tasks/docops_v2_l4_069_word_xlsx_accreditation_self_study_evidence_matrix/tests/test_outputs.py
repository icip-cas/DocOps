import json
import os
import re
import sys
import zipfile
from pathlib import Path
from datetime import date, datetime

from docx import Document
from openpyxl import load_workbook
from openpyxl.utils.cell import get_column_letter, range_boundaries

sys.path.insert(0, str(Path(__file__).parent))
from verifier_utils import *  # noqa: F401,F403

META_PATH = Path(os.environ.get('TASK_METADATA_PATH', '/tests/task_metadata.json'))
META = json.loads(META_PATH.read_text(encoding='utf-8'))
EXPECT = META['verifier_expectations']


def _vopt_norm_text(value):
    if value is None:
        return ''
    if isinstance(value, datetime):
        if value.hour or value.minute or value.second:
            return value.strftime('%Y-%m-%d %H:%M')
        return value.strftime('%Y-%m-%d')
    if isinstance(value, date):
        return value.strftime('%Y-%m-%d')
    if isinstance(value, float) and value.is_integer():
        value = int(value)
    text = str(value)
    text = text.replace('\u2013', '-').replace('\u2014', '-').replace('\u2212', '-')
    text = text.replace('\xa0', ' ')
    text = re.sub(r'\s+', ' ', text).strip()
    return text


def _vopt_norm_formula(value):
    return re.sub(r'\s+', '', _vopt_norm_text(value)).upper()


def _vopt_norm_key(value):
    return re.sub(r'[^A-Z0-9]+', '', _vopt_norm_text(value).upper())


def _vopt_sheet_has_formula_tokens(ws, tokens, label):
    for row in ws.iter_rows():
        for cell in row:
            value = cell.value
            if isinstance(value, str) and value.startswith('='):
                norm = _vopt_norm_formula(value)
                if all(_vopt_norm_formula(token) in norm for token in tokens):
                    return
    raise AssertionError(f"{label}: no formula found with purpose tokens {tokens!r}")


def _vopt_norm_rows(rows):
    return [[_vopt_norm_text(cell) for cell in row] for row in rows]


def _vopt_rows_equal(actual, expected):
    return _vopt_norm_rows(actual) == _vopt_norm_rows(expected)


def _vopt_rows_text(rows):
    return '\n'.join('|'.join(_vopt_norm_text(cell) for cell in row) for row in rows)


def _vopt_header_has(header, required):
    actual = {_vopt_norm_key(cell) for cell in header}
    missing = [
        cell
        for cell in required
        if not any(_vopt_norm_key(cell) == item or _vopt_norm_key(cell) in item or item in _vopt_norm_key(cell) for item in actual)
    ]
    assert not missing, f"Missing expected header columns {missing!r}; actual={header!r}"


def _vopt_section_text(section_part):
    return '\n'.join(p.text for p in getattr(section_part, 'paragraphs', []) if p.text is not None)


def _vopt_header_text(doc):
    return '\n'.join(_vopt_section_text(section.header) for section in doc.sections)


def _vopt_footer_text(doc):
    return '\n'.join(_vopt_section_text(section.footer) for section in doc.sections)


def _vopt_table_ref(ws, expected_name=None):
    tables = ws.tables
    if expected_name:
        for name in tables.keys():
            if str(name).lower() == str(expected_name).lower():
                return tables[name].ref
    vals = list(tables.values())
    assert vals, f"No Excel native table found on sheet {ws.title!r}"
    return vals[0].ref


def _vopt_table(ws, expected_name):
    for name in ws.tables.keys():
        if str(name).lower() == str(expected_name).lower():
            return ws.tables[name]
    raise AssertionError(f"No Excel native table named {expected_name!r} found on sheet {ws.title!r}")


def _vopt_table_has_size(ws, expected_name, min_rows, min_cols):
    table = _vopt_table(ws, expected_name)
    min_col, min_row, max_col, max_row = range_boundaries(table.ref)
    actual_rows = max_row - min_row + 1
    actual_cols = max_col - min_col + 1
    assert actual_rows >= min_rows and actual_cols >= min_cols, (
        f"{ws.title}.{expected_name}: expected at least {min_rows} rows x {min_cols} cols, "
        f"found {actual_rows} rows x {actual_cols} cols at {table.ref}"
    )
    return table


def _vopt_table_rows(ws, table_name):
    table = _vopt_table(ws, table_name)
    min_col, min_row, max_col, max_row = range_boundaries(table.ref)
    return [
        [ws.cell(row, col).value for col in range(min_col, max_col + 1)]
        for row in range(min_row, max_row + 1)
    ]


def _vopt_find_header_index(header, accepted_names):
    accepted = {_vopt_norm_key(name) for name in accepted_names}
    for idx, value in enumerate(header):
        if _vopt_norm_key(value) in accepted:
            return idx
    raise AssertionError(f"Missing expected column {accepted_names!r}; header={header!r}")


def _vopt_clean_area(value):
    text = _vopt_norm_text(value).replace("'", '').replace('$', '').replace(' ', '')
    if '!' in text:
        text = text.split('!', 1)[1]
    return text.upper()


def _vopt_assert_print_area(ws, expected):
    actual = _vopt_clean_area(ws.print_area)
    target = _vopt_clean_area(expected)
    assert target in actual or actual in target, f"{ws.title}: expected print area {expected!r}, found {ws.print_area!r}"


def _vopt_freeze_pane(value):
    return getattr(value, 'coordinate', value)


def _vopt_validated_cells(ws):
    cells = set()
    for dv in ws.data_validations.dataValidation:
        for rng in dv.cells.ranges:
            text = str(rng).replace('$', '')
            try:
                min_col, min_row, max_col, max_row = range_boundaries(text)
            except ValueError:
                cells.add(text)
                continue
            if max_row > 10000:
                max_row = max(min_row, 200)
            for row in range(min_row, max_row + 1):
                for col in range(min_col, max_col + 1):
                    cells.add(f"{get_column_letter(col)}{row}")
    return cells


class _VoptValidationRanges(list):
    def __init__(self, ws, ranges):
        super().__init__(ranges)
        self.ws = ws

    def __contains__(self, expected):
        expected = str(expected).replace('$', '')
        if super().__contains__(expected):
            return True
        try:
            min_col, min_row, max_col, max_row = range_boundaries(expected)
        except ValueError:
            return super().__contains__(expected)
        target = {f"{get_column_letter(col)}{row}" for row in range(min_row, max_row + 1) for col in range(min_col, max_col + 1)}
        return target.issubset(_vopt_validated_cells(self.ws))


def _vopt_dv_ranges(ws):
    ranges = []
    for dv in ws.data_validations.dataValidation:
        ranges.extend(str(rng).replace('$', '') for rng in dv.cells.ranges)
    return _VoptValidationRanges(ws, ranges)



def _path(kind):
    env = {'docx': 'DOCX_OUTPUT_PATH', 'xlsx': 'XLSX_OUTPUT_PATH'}[kind]
    key = {'docx': 'docx_output', 'xlsx': 'xlsx_output'}[kind]
    return Path(os.environ.get(env, EXPECT[key]))


def _doc_text(doc):
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append('|'.join(cell.text.strip() for cell in row.cells))
    return '\n'.join(parts)


def _doc_heading_order(doc):
    out = []
    for p in doc.paragraphs:
        style = p.style.name if p.style else ''
        text = p.text.strip()
        if not text:
            continue
        if not out and _norm_heading_label(text) == _norm_heading_label(EXPECT['doc_heading_order'][0]):
            out.append(text)
            continue
        if style.startswith('Heading') or style == 'Title':
            out.append(text)
    return out


def _norm_heading_label(value):
    text = str(value or '')
    text = text.replace('\u2013', '-').replace('\u2014', '-').replace('\xa0', ' ')
    text = re.sub(r'\s+', ' ', text).strip()
    text = re.sub(r'^(?:\d+|[IVXLCDM]+)[\.)]\s+', '', text, flags=re.I)
    text = re.sub(r'^[A-Z][\.)]\s+', '', text)
    text = re.sub(r'[:\-]+$', '', text).strip()
    return text.casefold()


def _heading_sequence_matches(actual, expected):
    actual_norm = [
        _norm_heading_label(item)
        for item in actual
        if _norm_heading_label(item) not in {'table of contents', 'contents'}
    ]
    expected_norm = [_norm_heading_label(item) for item in expected]
    pos = 0
    for item in actual_norm:
        if pos < len(expected_norm) and item == expected_norm[pos]:
            pos += 1
    return pos == len(expected_norm)


def _doc_table_rows(table):
    return [[cell.text.strip() for cell in row.cells] for row in table.rows]


def _has_toc_field(path):
    with zipfile.ZipFile(path) as zf:
        xml = zf.read('word/document.xml').decode('utf-8', errors='ignore')
    return bool(re.search(r'TOC\s+(?:\\)?o|TOC\s*(?:&quot;|")', xml))


def _highlighted_para(doc, text):
    for p in doc.paragraphs:
        if _vopt_norm_text(text) in _vopt_norm_text(p.text):
            return docx_para_has_highlight(p)
    raise AssertionError(f'Missing paragraph: {text}')


def _highlighted_std2_action_para(doc):
    candidates = []
    for p in doc.paragraphs:
        norm = _vopt_norm_text(p.text)
        if (
            'STD-2' in norm
            and re.search(r'72(?:\.0)?\s*%', norm)
            and re.search(r'75(?:\.0)?\s*%', norm)
            and 'two' in norm.lower()
            and 'commitment' in norm.lower()
        ):
            candidates.append(p.text)
            if docx_para_has_highlight(p):
                return
    if candidates:
        raise AssertionError(f"STD-2 action paragraph should be highlighted; candidates={candidates!r}")
    raise AssertionError('Missing highlighted STD-2 action paragraph with threshold and two public commitments')


def _find_doc_table_with_header(doc, required):
    for table in doc.tables:
        rows = _doc_table_rows(table)
        if rows:
            header = {_vopt_norm_key(cell) for cell in rows[0]}
            if all(any(_vopt_norm_key(req) in item or item in _vopt_norm_key(req) for item in header) for req in required):
                return rows
    raise AssertionError(f'Missing DOCX table with columns {required!r}')


def _require_self_study_semantics(text):
    require_all(text, ['STD-1', 'STD-2', 'STD-3', 'STD-4'], 'self-study response')
    assert re.search(r'72(?:\.0)?\s*%', text), 'self-study response: missing 72% or 72.0% achievement'
    assert re.search(r'75(?:\.0)?\s*%', text), 'self-study response: missing 75% threshold'
    require_all(text, ['EV-101', 'EV-102', 'EV-201', 'EV-203', 'EV-301', 'EV-401'], 'self-study response')
    require_all(text, ['rubric calibration', 'evidence upload'], 'self-study response')


def _evidence_set(value):
    return {item for item in re.split(r'[^A-Z0-9-]+', _vopt_norm_text(value).upper()) if item.startswith('EV-')}


def _cell_has_number_or_link(value, number):
    text = _vopt_norm_text(value)
    if isinstance(value, (int, float)):
        return abs(float(value) - float(number)) < 0.11 or abs(float(value) * 100 - float(number)) < 0.11
    if str(text).startswith('='):
        return True
    return str(number) in text or str(float(number)) in text


def _rows_by_first_col(rows):
    return {_vopt_norm_text(row[0]): row for row in rows if row and _vopt_norm_text(row[0])}


def _status_value_or_formula(value, expected):
    text = _vopt_norm_text(value)
    if _vopt_norm_key(text) == _vopt_norm_key(expected):
        return True
    return text.startswith('=')


def _table_dicts(rows):
    header, data = rows[0], rows[1:]
    return [
        {_vopt_norm_key(header[idx]): row[idx] for idx in range(min(len(header), len(row)))}
        for row in data
    ]


def _sheet(wb, name):
    assert name in wb.sheetnames, f"Missing required worksheet {name!r}; found {wb.sheetnames!r}"
    return wb[name]


def _row_values(ws, row, max_col):
    return [ws.cell(row, col).value for col in range(1, max_col + 1)]


def _dv_ranges(ws):
    return _vopt_dv_ranges(ws)


def test_outputs_exist():
    assert _path('docx').exists(), f"Missing DOCX output: {_path('docx')}"
    assert _path('xlsx').exists(), f"Missing XLSX output: {_path('xlsx')}"
    assert _path('docx').suffix.lower() == '.docx'
    assert _path('xlsx').suffix.lower() == '.xlsx'


def test_docx_self_study_structure_privacy_and_actions():
    docx = _path('docx')
    doc = Document(docx)
    assert _heading_sequence_matches(_doc_heading_order(doc), EXPECT['doc_heading_order']), f"Unexpected heading order: {_doc_heading_order(doc)!r}"
    text = _doc_text(doc)
    _require_self_study_semantics(text)
    forbid_any(text, EXPECT['doc_forbidden_text'], 'self-study response')
    assert EXPECT['doc_header_contains'] in _vopt_header_text(doc)
    assert EXPECT['doc_footer_contains'] in _vopt_footer_text(doc)
    assert _has_toc_field(docx)
    _highlighted_std2_action_para(doc)
    standards_table = _find_doc_table_with_header(doc, ['Standard', 'Status', 'PLO'])
    _vopt_header_has(standards_table[0], ['Standard ID', 'Status', 'PLO'])
    standards_text = _vopt_rows_text(standards_table[1:])
    for std in ['STD-1', 'STD-2', 'STD-3', 'STD-4']:
        assert std in standards_text
    assert 'Action Required' in standards_text
    action_table = _find_doc_table_with_header(doc, ['Action', 'Owner'])
    action_text = _vopt_rows_text(action_table[1:])
    require_all(action_text, ['ACT-201', 'ACT-202', 'rubric calibration', 'evidence upload'], 'DOCX action table')


def test_workbook_structure_controls_and_formulas():
    wb = load_workbook(_path('xlsx'), data_only=False)
    assert wb.sheetnames == EXPECT['sheet_order']
    for sheet_name, (table_name, ref) in EXPECT['tables'].items():
        min_col, min_row, max_col, max_row = range_boundaries(ref)
        _vopt_table_has_size(wb[sheet_name], table_name, max_row - min_row + 1, max_col - min_col + 1)
    for sheet in EXPECT['hidden_sheets']:
        assert wb[sheet].sheet_state in ('hidden', 'veryHidden')
    names = {dn.name for dn in wb.defined_names.values()}
    for name in EXPECT['defined_names']:
        assert name in names
    for sheet_name, print_area in EXPECT['print_areas'].items():
        if sheet_name == 'Evidence Index':
            continue
        _vopt_assert_print_area(wb[sheet_name], print_area)
    _vopt_sheet_has_formula_tokens(wb['Standards Matrix'], ['IF', '>=', 'Met', 'Action Required'], 'Standards Matrix status')
    _vopt_sheet_has_formula_tokens(wb['Standards Matrix'], ['IF', 'Ready', 'Action Plan Required'], 'Standards Matrix readiness')
    _vopt_sheet_has_formula_tokens(wb['Assessment Summary'], ['*100'], 'Assessment Summary aggregate percentage')
    _vopt_sheet_has_formula_tokens(wb['Assessment Summary'], ['IF', 'Met', 'Action Required'], 'Assessment Summary status')
    _vopt_sheet_has_formula_tokens(wb['Evidence Index'], ['IF', 'Publishable'], 'Evidence Index citeability')
    _vopt_sheet_has_formula_tokens(wb['Action Tracker'], ['IF', 'G'], 'Action Tracker completion')
    for sheet_name, ranges in EXPECT['data_validation_ranges'].items():
        actual = _dv_ranges(wb[sheet_name])
        for expected in ranges:
            assert expected in actual
    assert len(wb['Standards Matrix'].conditional_formatting) >= 1


def test_workbook_public_private_boundary_and_values():
    wb = load_workbook(_path('xlsx'), data_only=False)
    matrix_rows = _vopt_table_rows(_sheet(wb, 'Standards Matrix'), 'standards_matrix')[1:]
    matrix_by_std = _rows_by_first_col(matrix_rows)
    for expected in EXPECT['standards_rows']:
        std, _name, plo, evidence, threshold, achievement = expected[:6]
        assert std in matrix_by_std, f'Missing Standards Matrix row for {std}'
        row_text = _vopt_rows_text([matrix_by_std[std]])
        assert plo in row_text, f'{std}: missing PLO {plo}'
        assert _evidence_set(evidence).issubset(_evidence_set(row_text)), f'{std}: missing publishable evidence IDs'
        if std == 'STD-2':
            assert any(_cell_has_number_or_link(cell, threshold) for cell in matrix_by_std[std]), f'{std}: missing 75 threshold'
            assert any(_cell_has_number_or_link(cell, achievement) for cell in matrix_by_std[std]), f'{std}: missing 72 achievement'
    public_rows = _vopt_table_rows(_sheet(wb, 'Public Appendix'), 'public_appendix')[1:]
    public_by_std = _rows_by_first_col(public_rows)
    for std, status, evidence, summary in EXPECT['public_appendix_rows']:
        assert std in public_by_std, f'Missing Public Appendix row for {std}'
        row = public_by_std[std]
        assert _status_value_or_formula(row[1], status), f'{std}: wrong public status {row[1]!r}'
        assert _evidence_set(evidence).issubset(_evidence_set(_vopt_rows_text([row]))), f'{std}: missing public evidence IDs'
        row_text = _vopt_rows_text([row]).lower()
        if std == 'STD-2':
            require_all(row_text, ['72', '75', 'action', 'rubric', 'evidence upload'], f'{std} public summary')
    action_rows = _vopt_table_rows(_sheet(wb, 'Action Tracker'), 'action_tracker')[1:]
    action_text = _vopt_rows_text(action_rows)
    require_all(action_text, ['ACT-201', 'ACT-202', 'ACT-401', 'STD-2', 'STD-4'], 'Action Tracker')
    require_all(action_text, ['rubric calibration', 'evidence upload'], 'Action Tracker')
    assert 'Yes' in action_text and 'No' in action_text
    public_text = '\n'.join(str(cell.value) for sheet in ['Standards Matrix', 'Assessment Summary', 'Action Tracker', 'Public Appendix'] for row in wb[sheet].iter_rows() for cell in row if cell.value)
    forbid_any(public_text, EXPECT['forbidden_private'], 'public-facing workbook sheets')
    evidence_rows = _vopt_table_rows(_sheet(wb, 'Evidence Index'), 'evidence_index')[1:]
    assert [row[0] for row in evidence_rows if row[3] == 'Internal Only'] == ['EV-202', 'EV-402']
    protected_text = '\n'.join(str(cell.value) for row in wb['Raw Protected Notes'].iter_rows() for cell in row if cell.value)
    require_all(protected_text, ['Jordan Ellis', 'Mina Patel', 'Morgan Lee', 'EV-202', 'EV-402'], 'hidden protected notes')


def test_cross_output_standard_status_consistency():
    doc_text = _doc_text(Document(_path('docx')))
    wb = load_workbook(_path('xlsx'), data_only=False)
    matrix_rows = _vopt_table_rows(_sheet(wb, 'Standards Matrix'), 'standards_matrix')[1:]
    matrix_text = _vopt_rows_text(matrix_rows)
    require_all(matrix_text, ['STD-1', 'STD-2', 'STD-3', 'STD-4', 'Action Required'], 'Standards Matrix')
    assert 'STD-2' in doc_text and 'action' in doc_text.lower()
    public_rows = _vopt_table_rows(_sheet(wb, 'Public Appendix'), 'public_appendix')[1:]
    public_statuses = [row[1] for row in public_rows]
    for actual, expected in zip(public_statuses, ['Met', 'Action Required', 'Met', 'Met']):
        assert _status_value_or_formula(actual, expected)
    for private in EXPECT['forbidden_private']:
        assert private not in doc_text
