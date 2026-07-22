import json
import os
import re
import sys
import zipfile
from pathlib import Path
from datetime import date, datetime

from docx import Document
from openpyxl import load_workbook
from openpyxl.utils.cell import get_column_letter, range_boundaries

sys.path.insert(0, str(Path(__file__).parent))
from verifier_utils import *  # noqa: F401,F403

META_PATH = Path(os.environ.get('TASK_METADATA_PATH', '/tests/task_metadata.json'))
META = json.loads(META_PATH.read_text(encoding='utf-8'))
EXPECT = META['verifier_expectations']


def _vopt_norm_text(value):
    if value is None:
        return ''
    if isinstance(value, datetime):
        if value.hour or value.minute or value.second:
            return value.strftime('%Y-%m-%d %H:%M')
        return value.strftime('%Y-%m-%d')
    if isinstance(value, date):
        return value.strftime('%Y-%m-%d')
    if isinstance(value, float) and value.is_integer():
        value = int(value)
    text = str(value)
    text = text.replace('\u2013', '-').replace('\u2014', '-').replace('\u2212', '-')
    text = text.replace('\xa0', ' ')
    text = re.sub(r'\s+', ' ', text).strip()
    return text


def _vopt_norm_formula(value):
    return re.sub(r'\s+', '', _vopt_norm_text(value)).upper()


def _vopt_assert_formula(actual, expected, label='formula'):
    assert _vopt_norm_formula(actual) == _vopt_norm_formula(expected), f"{label}: expected {expected!r}, found {actual!r}"


def _vopt_norm_rows(rows):
    return [[_vopt_norm_text(cell) for cell in row] for row in rows]


def _vopt_rows_equal(actual, expected):
    return _vopt_norm_rows(actual) == _vopt_norm_rows(expected)


def _vopt_rows_text(rows):
    return '\n'.join('|'.join(_vopt_norm_text(cell) for cell in row) for row in rows)


def _vopt_header_has(header, required):
    actual = {_vopt_norm_text(cell).lower() for cell in header}
    missing = [cell for cell in required if _vopt_norm_text(cell).lower() not in actual]
    assert not missing, f"Missing expected header columns {missing!r}; actual={header!r}"


def _vopt_section_text(section_part):
    return '\n'.join(p.text for p in getattr(section_part, 'paragraphs', []) if p.text is not None)


def _vopt_header_text(doc):
    return '\n'.join(_vopt_section_text(section.header) for section in doc.sections)


def _vopt_footer_text(doc):
    return '\n'.join(_vopt_section_text(section.footer) for section in doc.sections)


def _vopt_table_ref(ws, expected_name=None):
    tables = ws.tables
    if expected_name:
        for name in tables.keys():
            if str(name).lower() == str(expected_name).lower():
                return tables[name].ref
    vals = list(tables.values())
    assert vals, f"No Excel native table found on sheet {ws.title!r}"
    return vals[0].ref


def _vopt_clean_area(value):
    text = _vopt_norm_text(value).replace("'", '').replace('$', '').replace(' ', '')
    if '!' in text:
        text = text.split('!', 1)[1]
    return text.upper()


def _vopt_assert_print_area(ws, expected):
    actual = _vopt_clean_area(ws.print_area)
    target = _vopt_clean_area(expected)
    assert target in actual or actual in target, f"{ws.title}: expected print area {expected!r}, found {ws.print_area!r}"


def _vopt_freeze_pane(value):
    return getattr(value, 'coordinate', value)


def _vopt_validated_cells(ws):
    cells = set()
    for dv in ws.data_validations.dataValidation:
        for rng in dv.cells.ranges:
            text = str(rng).replace('$', '')
            try:
                min_col, min_row, max_col, max_row = range_boundaries(text)
            except ValueError:
                cells.add(text)
                continue
            if max_row > 10000:
                max_row = max(min_row, 200)
            for row in range(min_row, max_row + 1):
                for col in range(min_col, max_col + 1):
                    cells.add(f"{get_column_letter(col)}{row}")
    return cells


class _VoptValidationRanges(list):
    def __init__(self, ws, ranges):
        super().__init__(ranges)
        self.ws = ws

    def __contains__(self, expected):
        expected = str(expected).replace('$', '')
        if super().__contains__(expected):
            return True
        try:
            min_col, min_row, max_col, max_row = range_boundaries(expected)
        except ValueError:
            return super().__contains__(expected)
        target = {f"{get_column_letter(col)}{row}" for row in range(min_row, max_row + 1) for col in range(min_col, max_col + 1)}
        return target.issubset(_vopt_validated_cells(self.ws))


def _vopt_dv_ranges(ws):
    ranges = []
    for dv in ws.data_validations.dataValidation:
        ranges.extend(str(rng).replace('$', '') for rng in dv.cells.ranges)
    return _VoptValidationRanges(ws, ranges)



def _path(kind):
    env = {'docx': 'DOCX_OUTPUT_PATH', 'xlsx': 'XLSX_OUTPUT_PATH'}[kind]
    key = {'docx': 'docx_output', 'xlsx': 'xlsx_output'}[kind]
    return Path(os.environ.get(env, EXPECT[key]))


def _doc_text(doc):
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append('|'.join(cell.text.strip() for cell in row.cells))
    return '\n'.join(parts)


def _doc_heading_order(doc):
    out = []
    for p in doc.paragraphs:
        style = p.style.name if p.style else ''
        if p.text.strip() and (style.startswith('Heading') or style == 'Title'):
            out.append(p.text.strip())
    return out


def _doc_table_rows(table):
    return [[cell.text.strip() for cell in row.cells] for row in table.rows]


def _has_toc_field(path):
    with zipfile.ZipFile(path) as zf:
        xml = zf.read('word/document.xml').decode('utf-8', errors='ignore')
    return bool(re.search(r'TOC\s+(?:\\)?o|TOC\s*(?:&quot;|")', xml))


def _highlighted_para(doc, text):
    for p in doc.paragraphs:
        if _vopt_norm_text(text) in _vopt_norm_text(p.text):
            return docx_para_has_highlight(p)
    raise AssertionError(f'Missing paragraph: {text}')


def _row_values(ws, row, max_col):
    return [ws.cell(row, col).value for col in range(1, max_col + 1)]


def _dv_ranges(ws):
    return _vopt_dv_ranges(ws)


def test_outputs_exist():
    assert _path('docx').exists(), f"Missing DOCX output: {_path('docx')}"
    assert _path('xlsx').exists(), f"Missing XLSX output: {_path('xlsx')}"
    assert _path('docx').suffix.lower() == '.docx'
    assert _path('xlsx').suffix.lower() == '.xlsx'


def test_docx_reimbursement_request_structure_amounts_and_boundary():
    docx = _path('docx')
    doc = Document(docx)
    assert _doc_heading_order(doc) == EXPECT['doc_heading_order']
    text = _doc_text(doc)
    require_all(text, EXPECT['doc_required_paragraphs'], 'reimbursement request')
    forbid_any(text, EXPECT['doc_forbidden_text'], 'reimbursement request')
    assert EXPECT['doc_header_contains'] in _vopt_header_text(doc)
    assert EXPECT['doc_footer_contains'] in _vopt_footer_text(doc)
    assert _has_toc_field(docx)
    assert _highlighted_para(doc, EXPECT['doc_highlight_text'])
    rows = _doc_table_rows(doc.tables[0])
    _vopt_header_has(rows[0], ['Cost ID', 'Amount'])
    assert [row[0] for row in rows[1:]] == [row[0] for row in EXPECT['eligible_rows']]


def test_workbook_structure_controls_and_formulas():
    wb = load_workbook(_path('xlsx'), data_only=False)
    assert wb.sheetnames == EXPECT['sheet_order']
    for sheet_name, (table_name, ref) in EXPECT['tables'].items():
        assert _vopt_table_ref(wb[sheet_name], table_name) == ref
    for sheet in EXPECT['hidden_sheets']:
        assert wb[sheet].sheet_state in ('hidden', 'veryHidden')
    names = {dn.name for dn in wb.defined_names.values()}
    for name in EXPECT['defined_names']:
        assert name in names
    for sheet_name, print_area in EXPECT['print_areas'].items():
        _vopt_assert_print_area(wb[sheet_name], print_area)
    for ref, formula in EXPECT['formula_cells'].items():
        sheet, cell = ref.split('!')
        _vopt_assert_formula(wb[sheet][cell].value, formula, ref)
    for sheet_name, ranges in EXPECT['data_validation_ranges'].items():
        actual = _dv_ranges(wb[sheet_name])
        for expected in ranges:
            assert expected in actual
    assert len(wb['Cost Detail'].conditional_formatting) >= 1


def test_workbook_values_public_private_boundary_and_match():
    wb = load_workbook(_path('xlsx'), data_only=False)
    cost_rows = [_row_values(wb['Cost Detail'], row, 7) for row in range(2, 10)]
    assert _vopt_rows_equal(cost_rows, EXPECT['all_cost_rows'])
    match_rows = [_row_values(wb['Match Ledger'], row, 6) for row in range(2, 5)]
    assert _vopt_rows_equal(match_rows, EXPECT['match_rows'])
    procurement_rows = [_row_values(wb['Procurement Review'], row, 5) for row in range(2, 6)]
    assert _vopt_rows_equal(procurement_rows, EXPECT['procurement_rows'])
    public_text = '\n'.join(str(cell.value) for sheet in ['Reimbursement Summary', 'Match Ledger', 'Procurement Review', 'Public Narrative'] for row in wb[sheet].iter_rows() for cell in row if cell.value)
    forbid_any(public_text, EXPECT['forbidden_public'], 'public workbook sheets')
    assert [row[0] for row in cost_rows if row[5] == 'Internal Only' and row[6] == 'Internal'] == ['CASH-999']
    internal_text = '\n'.join(str(cell.value) for row in wb['Internal Strategy'].iter_rows() for cell in row if cell.value)
    require_all(internal_text, ['CASH-999', 'cash-flow gap', 'advance draw', 'temporary payroll float'], 'hidden internal strategy')
    evidence_rows = [_row_values(wb['Evidence Index'], row, 5) for row in range(2, 11)]
    assert _vopt_rows_equal(evidence_rows, EXPECT['evidence_rows'])


def test_cross_output_reimbursement_consistency():
    doc_text = _doc_text(Document(_path('docx')))
    wb = load_workbook(_path('xlsx'), data_only=False)
    assert '$43,912.50' in doc_text
    assert '$58,550.00' in doc_text
    assert '$14,637.50' in doc_text
    assert wb['Reimbursement Summary']['B3'].value == '=B2*0.75'
    assert wb['Reimbursement Summary']['B4'].value == '=B2-B3'
    for cost_id in [row[0] for row in EXPECT['eligible_rows']]:
        assert cost_id in doc_text
    for forbidden in EXPECT['forbidden_public']:
        assert normalize_text(forbidden) not in normalize_text(doc_text)
