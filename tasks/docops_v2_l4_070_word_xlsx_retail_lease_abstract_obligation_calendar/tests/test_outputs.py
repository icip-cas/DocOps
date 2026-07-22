import json
import os
import re
import sys
import zipfile
from pathlib import Path
from datetime import date, datetime

from docx import Document
from openpyxl import load_workbook
from openpyxl.utils.cell import get_column_letter, range_boundaries

sys.path.insert(0, str(Path(__file__).parent))
from verifier_utils import *  # noqa: F401,F403

META_PATH = Path(os.environ.get('TASK_METADATA_PATH', '/tests/task_metadata.json'))
META = json.loads(META_PATH.read_text(encoding='utf-8'))
EXPECT = META['verifier_expectations']


def _vopt_norm_text(value):
    if value is None:
        return ''
    if isinstance(value, datetime):
        if value.hour or value.minute or value.second:
            return value.strftime('%Y-%m-%d %H:%M')
        return value.strftime('%Y-%m-%d')
    if isinstance(value, date):
        return value.strftime('%Y-%m-%d')
    if isinstance(value, float) and value.is_integer():
        value = int(value)
    text = str(value)
    text = text.replace('\u2013', '-').replace('\u2014', '-').replace('\u2212', '-')
    text = text.replace('\xa0', ' ')
    text = re.sub(r'\s+', ' ', text).strip()
    return text


def _vopt_norm_formula(value):
    return re.sub(r'\s+', '', _vopt_norm_text(value)).upper()


def _vopt_norm_key(value):
    return re.sub(r'[^A-Z0-9]+', '', _vopt_norm_text(value).upper())


def _vopt_sheet_has_formula_tokens(ws, tokens, label):
    for row in ws.iter_rows():
        for cell in row:
            value = cell.value
            if isinstance(value, str) and value.startswith('='):
                norm = _vopt_norm_formula(value)
                if all(_vopt_norm_formula(token) in norm for token in tokens):
                    return
    raise AssertionError(f"{label}: no formula found with purpose tokens {tokens!r}")


def _vopt_norm_rows(rows):
    return [[_vopt_norm_text(cell) for cell in row] for row in rows]


def _vopt_rows_equal(actual, expected):
    return _vopt_norm_rows(actual) == _vopt_norm_rows(expected)


def _vopt_rows_text(rows):
    return '\n'.join('|'.join(_vopt_norm_text(cell) for cell in row) for row in rows)


def _vopt_header_has(header, required):
    actual = {_vopt_norm_key(cell) for cell in header}
    missing = [
        cell
        for cell in required
        if not any(_vopt_norm_key(cell) == item or _vopt_norm_key(cell) in item or item in _vopt_norm_key(cell) for item in actual)
    ]
    assert not missing, f"Missing expected header columns {missing!r}; actual={header!r}"


def _vopt_section_text(section_part):
    return '\n'.join(p.text for p in getattr(section_part, 'paragraphs', []) if p.text is not None)


def _vopt_header_text(doc):
    return '\n'.join(_vopt_section_text(section.header) for section in doc.sections)


def _vopt_footer_text(doc):
    return '\n'.join(_vopt_section_text(section.footer) for section in doc.sections)


def _vopt_table_ref(ws, expected_name=None):
    tables = ws.tables
    if expected_name:
        for name in tables.keys():
            if str(name).lower() == str(expected_name).lower():
                return tables[name].ref
    vals = list(tables.values())
    assert vals, f"No Excel native table found on sheet {ws.title!r}"
    return vals[0].ref


def _vopt_table(ws, expected_name):
    for name in ws.tables.keys():
        if str(name).lower() == str(expected_name).lower():
            return ws.tables[name]
    raise AssertionError(f"No Excel native table named {expected_name!r} found on sheet {ws.title!r}")


def _vopt_table_has_size(ws, expected_name, min_rows, min_cols):
    table = _vopt_table(ws, expected_name)
    min_col, min_row, max_col, max_row = range_boundaries(table.ref)
    actual_rows = max_row - min_row + 1
    actual_cols = max_col - min_col + 1
    assert actual_rows >= min_rows and actual_cols >= min_cols, (
        f"{ws.title}.{expected_name}: expected at least {min_rows} rows x {min_cols} cols, "
        f"found {actual_rows} rows x {actual_cols} cols at {table.ref}"
    )
    return table


def _vopt_table_rows(ws, table_name):
    table = _vopt_table(ws, table_name)
    min_col, min_row, max_col, max_row = range_boundaries(table.ref)
    return [
        [ws.cell(row, col).value for col in range(min_col, max_col + 1)]
        for row in range(min_row, max_row + 1)
    ]


def _sheet(wb, name):
    assert name in wb.sheetnames, f"Missing required worksheet {name!r}; found {wb.sheetnames!r}"
    return wb[name]


def _vopt_clean_area(value):
    text = _vopt_norm_text(value).replace("'", '').replace('$', '').replace(' ', '')
    if '!' in text:
        text = text.split('!', 1)[1]
    return text.upper()


def _vopt_assert_print_area(ws, expected):
    actual = _vopt_clean_area(ws.print_area)
    target = _vopt_clean_area(expected)
    assert target in actual or actual in target, f"{ws.title}: expected print area {expected!r}, found {ws.print_area!r}"


def _vopt_freeze_pane(value):
    return getattr(value, 'coordinate', value)


def _vopt_validated_cells(ws):
    cells = set()
    for dv in ws.data_validations.dataValidation:
        for rng in dv.cells.ranges:
            text = str(rng).replace('$', '')
            try:
                min_col, min_row, max_col, max_row = range_boundaries(text)
            except ValueError:
                cells.add(text)
                continue
            if max_row > 10000:
                max_row = max(min_row, 200)
            for row in range(min_row, max_row + 1):
                for col in range(min_col, max_col + 1):
                    cells.add(f"{get_column_letter(col)}{row}")
    return cells


class _VoptValidationRanges(list):
    def __init__(self, ws, ranges):
        super().__init__(ranges)
        self.ws = ws

    def __contains__(self, expected):
        expected = str(expected).replace('$', '')
        if super().__contains__(expected):
            return True
        try:
            min_col, min_row, max_col, max_row = range_boundaries(expected)
        except ValueError:
            return super().__contains__(expected)
        target = {f"{get_column_letter(col)}{row}" for row in range(min_row, max_row + 1) for col in range(min_col, max_col + 1)}
        return target.issubset(_vopt_validated_cells(self.ws))


def _vopt_dv_ranges(ws):
    ranges = []
    for dv in ws.data_validations.dataValidation:
        ranges.extend(str(rng).replace('$', '') for rng in dv.cells.ranges)
    return _VoptValidationRanges(ws, ranges)



def _path(kind):
    env = {'docx': 'DOCX_OUTPUT_PATH', 'xlsx': 'XLSX_OUTPUT_PATH'}[kind]
    key = {'docx': 'docx_output', 'xlsx': 'xlsx_output'}[kind]
    return Path(os.environ.get(env, EXPECT[key]))


def _doc_text(doc):
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append('|'.join(cell.text.strip() for cell in row.cells))
    return '\n'.join(parts)


def _doc_heading_order(doc):
    out = []
    for p in doc.paragraphs:
        style = p.style.name if p.style else ''
        text = p.text.strip()
        if not text:
            continue
        if not out and _norm_heading_label(text) == _norm_heading_label(EXPECT['doc_heading_order'][0]):
            out.append(text)
            continue
        if style.startswith('Heading') or style == 'Title':
            out.append(text)
    return out


def _norm_heading_label(value):
    text = str(value or '')
    text = text.replace('\u2013', '-').replace('\u2014', '-').replace('\xa0', ' ')
    text = re.sub(r'\s+', ' ', text).strip()
    text = re.sub(r'^(?:\d+|[IVXLCDM]+)[\.)]\s+', '', text, flags=re.I)
    text = re.sub(r'^[A-Z][\.)]\s+', '', text)
    text = re.sub(r'[:\-]+$', '', text).strip()
    return text.casefold()


def _heading_sequence_matches(actual, expected):
    actual_norm = [
        _norm_heading_label(item)
        for item in actual
        if _norm_heading_label(item) not in {'table of contents', 'contents'}
    ]
    expected_norm = [_norm_heading_label(item) for item in expected]
    pos = 0
    for item in actual_norm:
        if pos < len(expected_norm) and item == expected_norm[pos]:
            pos += 1
    return pos == len(expected_norm)


def _doc_table_rows(table):
    return [[cell.text.strip() for cell in row.cells] for row in table.rows]


def _has_toc_field(path):
    with zipfile.ZipFile(path) as zf:
        xml = zf.read('word/document.xml').decode('utf-8', errors='ignore')
    return bool(re.search(r'TOC\s+(?:\\)?o|TOC\s*(?:&quot;|")', xml))


def _highlighted_para(doc, text):
    for p in doc.paragraphs:
        if _vopt_norm_text(text) in _vopt_norm_text(p.text):
            return docx_para_has_highlight(p)
    raise AssertionError(f'Missing paragraph: {text}')


def _highlighted_cam_para(doc):
    candidates = []
    for p in doc.paragraphs:
        norm = _vopt_norm_text(p.text)
        if '$2,800' in norm and '2026-08-15' in norm and 'CAM' in norm:
            candidates.append(p.text)
            if docx_para_has_highlight(p):
                return
    if candidates:
        raise AssertionError(f"CAM paragraph should be highlighted; candidates={candidates!r}")
    raise AssertionError('Missing highlighted CAM paragraph with $2,800 over-cap amount and 2026-08-15 objection deadline')


def _find_doc_table_with_header(doc, required):
    for table in doc.tables:
        rows = _doc_table_rows(table)
        if rows:
            header = {_vopt_norm_key(cell) for cell in rows[0]}
            if all(any(_vopt_norm_key(req) in item or item in _vopt_norm_key(req) for item in header) for req in required):
                return rows
    raise AssertionError(f'Missing DOCX table with columns {required!r}')


def _require_lease_memo_semantics(text):
    require_all(text, ['North Pier Bikes', 'Unit 4B', '2028-07-01', '2028-09-30', '$2,800', '2026-08-15'], 'lease abstract memo')
    require_all(text, ['OPT-001', 'OPT-002', 'CAM-001', 'INS-001', 'RPT-001'], 'lease abstract memo')
    require_all(text, ['$18,750', '$19,500', '$20,250'], 'lease abstract memo')


def _rows_by_first_col(rows):
    return {_vopt_norm_text(row[0]): row for row in rows if row and _vopt_norm_text(row[0])}


def _money_or_number_present(text, value):
    norm = _vopt_norm_text(text).replace(',', '')
    return str(value).replace(',', '') in norm


def _row_text_contains_any(row, choices):
    text = _vopt_rows_text([row]).lower()
    return any(str(choice).lower() in text for choice in choices)


def _row_values(ws, row, max_col):
    return [ws.cell(row, col).value for col in range(1, max_col + 1)]


def _dv_ranges(ws):
    return _vopt_dv_ranges(ws)


def test_outputs_exist():
    assert _path('docx').exists(), f"Missing DOCX output: {_path('docx')}"
    assert _path('xlsx').exists(), f"Missing XLSX output: {_path('xlsx')}"
    assert _path('docx').suffix.lower() == '.docx'
    assert _path('xlsx').suffix.lower() == '.xlsx'


def test_docx_lease_abstract_structure_dates_and_boundary():
    docx = _path('docx')
    doc = Document(docx)
    assert _heading_sequence_matches(_doc_heading_order(doc), EXPECT['doc_heading_order']), f"Unexpected heading order: {_doc_heading_order(doc)!r}"
    text = _doc_text(doc)
    _require_lease_memo_semantics(text)
    forbid_any(text, EXPECT['doc_forbidden_text'], 'lease abstract memo')
    assert EXPECT['doc_header_contains'] in _vopt_header_text(doc)
    assert EXPECT['doc_footer_contains'] in _vopt_footer_text(doc)
    assert _has_toc_field(docx)
    _highlighted_cam_para(doc)
    rows = _find_doc_table_with_header(doc, ['Obligation', 'Due Date', 'Status'])
    _vopt_header_has(rows[0], ['Obligation', 'Due Date', 'Status'])
    table_text = _vopt_rows_text(rows[1:])
    for row in EXPECT['calendar_rows']:
        assert row[0] in table_text and row[3] in table_text


def test_workbook_structure_controls_formulas_and_dates():
    wb = load_workbook(_path('xlsx'), data_only=False)
    assert wb.sheetnames == EXPECT['sheet_order']
    for sheet_name, (table_name, ref) in EXPECT['tables'].items():
        min_col, min_row, max_col, max_row = range_boundaries(ref)
        min_cols = max_col - min_col + 1
        if sheet_name == 'CAM Reconciliation':
            min_cols = 5
        elif sheet_name == 'Rent Schedule':
            min_cols = 3
        elif sheet_name == 'Public Summary':
            min_cols = 2
        _vopt_table_has_size(_sheet(wb, sheet_name), table_name, max_row - min_row + 1, min_cols)
    for sheet in EXPECT['hidden_sheets']:
        assert wb[sheet].sheet_state in ('hidden', 'veryHidden')
    names = {dn.name for dn in wb.defined_names.values()}
    for name in EXPECT['defined_names']:
        assert name in names
    for sheet_name, print_area in EXPECT['print_areas'].items():
        _vopt_assert_print_area(wb[sheet_name], print_area)
    _vopt_sheet_has_formula_tokens(_sheet(wb, 'CAM Reconciliation'), ['1.05'], 'CAM cap formula')
    _vopt_sheet_has_formula_tokens(_sheet(wb, 'CAM Reconciliation'), ['-', 'C'], 'CAM over-cap formula')
    _vopt_sheet_has_formula_tokens(_sheet(wb, 'CAM Reconciliation'), ['IF', 'Dispute', 'OK'], 'CAM dispute status formula')
    _vopt_sheet_has_formula_tokens(_sheet(wb, 'Rent Schedule'), ['*12'], 'annualized rent formula')
    for sheet_name, ranges in EXPECT['data_validation_ranges'].items():
        actual = _dv_ranges(wb[sheet_name])
        for expected in ranges:
            assert expected in actual
    assert len(wb['Obligation Calendar'].conditional_formatting) >= 1
    assert len(wb['CAM Reconciliation'].conditional_formatting) >= 1


def test_workbook_public_private_boundary_and_values():
    wb = load_workbook(_path('xlsx'), data_only=False)
    abstract_rows = _vopt_table_rows(_sheet(wb, 'Lease Abstract'), 'lease_abstract')[1:]
    abstract_by_id = _rows_by_first_col(abstract_rows)
    for expected in EXPECT['abstract_rows']:
        item_id = expected[0]
        assert item_id in abstract_by_id, f'Missing Lease Abstract row {item_id}'
    abstract_text = _vopt_rows_text(abstract_rows)
    require_all(abstract_text, ['2028-07-01', '2028-09-30', '$18,750', '$19,500', '$20,250', '$2,800', '2026-08-15'], 'Lease Abstract')
    calendar_rows = _vopt_table_rows(_sheet(wb, 'Obligation Calendar'), 'obligation_calendar')[1:]
    calendar_by_id = _rows_by_first_col(calendar_rows)
    for expected in EXPECT['calendar_rows']:
        obligation_id, _obligation, category, due_date = expected[:4]
        assert obligation_id in calendar_by_id, f'Missing Obligation Calendar row {obligation_id}'
        row_text = _vopt_rows_text([calendar_by_id[obligation_id]])
        assert due_date in row_text, f'{obligation_id}: missing due date {due_date}'
        assert 'Public' in row_text, f'{obligation_id}: missing public boundary'
    public_text = '\n'.join(str(cell.value) for sheet in ['Lease Abstract', 'Obligation Calendar', 'CAM Reconciliation', 'Rent Schedule', 'Public Summary'] for row in wb[sheet].iter_rows() for cell in row if cell.value)
    forbid_any(public_text, EXPECT['forbidden_public'], 'public workbook sheets')
    internal_text = '\n'.join(str(cell.value) for row in wb['Internal Strategy'].iter_rows() for cell in row if cell.value)
    require_all(internal_text, ['INT-900', '$22,000', 'option leverage'], 'hidden internal strategy')
    cam_text = _vopt_rows_text(_vopt_table_rows(_sheet(wb, 'CAM Reconciliation'), 'cam_reconciliation'))
    require_all(cam_text, ['42000', '46900', 'Dispute'], 'CAM Reconciliation')
    assert _money_or_number_present(cam_text, 2800) or 'C3-D2' in _vopt_norm_formula(cam_text) or 'C3-D' in _vopt_norm_formula(cam_text)
    rent_text = _vopt_rows_text(_vopt_table_rows(_sheet(wb, 'Rent Schedule'), 'rent_schedule'))
    require_all(rent_text, ['2026', '18750', '2027', '19500', '2028', '20250'], 'Rent Schedule')


def test_cross_output_calendar_consistency():
    doc_text = _doc_text(Document(_path('docx')))
    wb = load_workbook(_path('xlsx'), data_only=False)
    calendar_ids = [row[0] for row in _vopt_table_rows(_sheet(wb, 'Obligation Calendar'), 'obligation_calendar')[1:]]
    assert calendar_ids == [row[0] for row in EXPECT['calendar_rows']]
    for row in EXPECT['calendar_rows']:
        assert row[0] in doc_text
        assert row[3] in doc_text
    assert '2028-10-31' not in doc_text
    assert '2028-10-31' not in '\n'.join(str(cell.value) for sheet in ['Lease Abstract', 'Obligation Calendar', 'Public Summary'] for row in _sheet(wb, sheet).iter_rows() for cell in row if cell.value)
