import json
import os
import re
import sys
import zipfile
from pathlib import Path
from datetime import date, datetime

from docx import Document
from openpyxl import load_workbook
from openpyxl.utils.cell import get_column_letter, range_boundaries

sys.path.insert(0, str(Path(__file__).parent))
from verifier_utils import *  # noqa: F401,F403

META_PATH = Path(os.environ.get('TASK_METADATA_PATH', '/tests/task_metadata.json'))
META = json.loads(META_PATH.read_text(encoding='utf-8'))
EXPECT = META['verifier_expectations']


def _vopt_norm_text(value):
    if value is None:
        return ''
    if isinstance(value, datetime):
        if value.hour or value.minute or value.second:
            return value.strftime('%Y-%m-%d %H:%M')
        return value.strftime('%Y-%m-%d')
    if isinstance(value, date):
        return value.strftime('%Y-%m-%d')
    if isinstance(value, float) and value.is_integer():
        value = int(value)
    text = str(value)
    text = text.replace('\u2013', '-').replace('\u2014', '-').replace('\u2212', '-')
    text = text.replace('\xa0', ' ')
    text = re.sub(r'\s+', ' ', text).strip()
    return text


def _vopt_norm_formula(value):
    return re.sub(r'\s+', '', _vopt_norm_text(value)).upper()


def _vopt_assert_formula(actual, expected, label='formula'):
    assert _vopt_norm_formula(actual) == _vopt_norm_formula(expected), f"{label}: expected {expected!r}, found {actual!r}"


def _vopt_norm_rows(rows):
    return [[_vopt_norm_text(cell) for cell in row] for row in rows]


def _vopt_rows_equal(actual, expected):
    return _vopt_norm_rows(actual) == _vopt_norm_rows(expected)


def _vopt_norm_key(value):
    return re.sub(r'[^a-z0-9]+', '', _vopt_norm_text(value).lower())


def _vopt_rows_text(rows):
    return '\n'.join('|'.join(_vopt_norm_text(cell) for cell in row) for row in rows)


def _vopt_header_has(header, required):
    actual = {_vopt_norm_text(cell).lower() for cell in header}
    missing = [cell for cell in required if _vopt_norm_text(cell).lower() not in actual]
    assert not missing, f"Missing expected header columns {missing!r}; actual={header!r}"


def _vopt_section_text(section_part):
    return '\n'.join(p.text for p in getattr(section_part, 'paragraphs', []) if p.text is not None)


def _vopt_header_text(doc):
    return '\n'.join(_vopt_section_text(section.header) for section in doc.sections)


def _vopt_footer_text(doc):
    return '\n'.join(_vopt_section_text(section.footer) for section in doc.sections)


def _vopt_table_ref(ws, expected_name=None):
    tables = ws.tables
    if expected_name:
        for name in tables.keys():
            if str(name).lower() == str(expected_name).lower():
                return tables[name].ref
    vals = list(tables.values())
    assert vals, f"No Excel native table found on sheet {ws.title!r}"
    return vals[0].ref


def _vopt_clean_area(value):
    text = _vopt_norm_text(value).replace("'", '').replace('$', '').replace(' ', '')
    if '!' in text:
        text = text.split('!', 1)[1]
    return text.upper()


def _vopt_assert_print_area(ws, expected):
    actual = _vopt_clean_area(ws.print_area)
    target = _vopt_clean_area(expected)
    assert target in actual or actual in target, f"{ws.title}: expected print area {expected!r}, found {ws.print_area!r}"


def _vopt_freeze_pane(value):
    return getattr(value, 'coordinate', value)


def _vopt_validated_cells(ws):
    cells = set()
    for dv in ws.data_validations.dataValidation:
        for rng in dv.cells.ranges:
            text = str(rng).replace('$', '')
            try:
                min_col, min_row, max_col, max_row = range_boundaries(text)
            except ValueError:
                cells.add(text)
                continue
            if max_row > 10000:
                max_row = max(min_row, 200)
            for row in range(min_row, max_row + 1):
                for col in range(min_col, max_col + 1):
                    cells.add(f"{get_column_letter(col)}{row}")
    return cells


class _VoptValidationRanges(list):
    def __init__(self, ws, ranges):
        super().__init__(ranges)
        self.ws = ws

    def __contains__(self, expected):
        expected = str(expected).replace('$', '')
        if super().__contains__(expected):
            return True
        try:
            min_col, min_row, max_col, max_row = range_boundaries(expected)
        except ValueError:
            return super().__contains__(expected)
        target = {f"{get_column_letter(col)}{row}" for row in range(min_row, max_row + 1) for col in range(min_col, max_col + 1)}
        return target.issubset(_vopt_validated_cells(self.ws))


def _vopt_dv_ranges(ws):
    ranges = []
    for dv in ws.data_validations.dataValidation:
        ranges.extend(str(rng).replace('$', '') for rng in dv.cells.ranges)
    return _VoptValidationRanges(ws, ranges)


def _vopt_ordered_subset(expected, actual):
    actual_iter = iter([_vopt_norm_text(item) for item in actual])
    for item in [_vopt_norm_text(value) for value in expected]:
        for candidate in actual_iter:
            if item == candidate:
                break
        else:
            return False
    return True



def _doc_path():
    return Path(os.environ.get('DOCX_OUTPUT_PATH', EXPECT['document_output']))


def _xlsx_path():
    return Path(os.environ.get('XLSX_OUTPUT_PATH', EXPECT['workbook_output']))


def _headings(doc):
    out = []
    for p in doc.paragraphs:
        style = p.style.name if p.style else ''
        if p.text.strip() and (style.startswith('Heading') or style == 'Title'):
            out.append(p.text.strip())
    return out


def _doc_text(doc):
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append('|'.join(cell.text.strip() for cell in row.cells))
    return '\n'.join(parts)


def _table_rows(table):
    return [[cell.text.strip() for cell in row.cells] for row in table.rows]


def _has_toc(path):
    with zipfile.ZipFile(path) as zf:
        xml = zf.read('word/document.xml').decode('utf-8', errors='ignore')
    return bool(re.search(r'TOC\s+(?:\\)?o|TOC\s*(?:&quot;|")', xml))


def _highlighted(doc, text):
    for p in doc.paragraphs:
        if _vopt_norm_text(text) in _vopt_norm_text(p.text):
            return any(run.font.highlight_color is not None for run in p.runs) or '<w:highlight' in p._p.xml
    return False


def _has_highlighted_terms(doc, terms):
    for p in doc.paragraphs:
        text = _vopt_norm_text(p.text).lower()
        if not text:
            continue
        has_highlight = any(run.font.highlight_color is not None for run in p.runs) or '<w:highlight' in p._p.xml
        if has_highlight and all(term.lower() in text for term in terms):
            return True
    return False


def _has_highlighted_format_exception(doc):
    for p in doc.paragraphs:
        text = _vopt_norm_text(p.text).lower()
        if not text:
            continue
        has_highlight = any(run.font.highlight_color is not None for run in p.runs) or '<w:highlight' in p._p.xml
        if has_highlight and 'format' in text and (
            'exception' in text or 'margin' in text or '0.7' in text or '0.75' in text or 'excluded' in text
        ):
            return True
    return False


def _assert_word_semantics(text):
    norm = _vopt_norm_text(text).lower()
    assert 'workbook' in norm, 'docx: missing workbook reference'
    assert any(term in norm for term in ['foundation', 'funder', 'format']), 'docx: missing foundation/funder formatting context'
    for line_id in EXPECT['public_line_ids']:
        assert line_id in text, f'docx: missing public line {line_id}'


def _assert_heading_sequence(doc):
    actual = [_vopt_norm_key(item) for item in _headings(doc)]
    expected = [_vopt_norm_key(item) for item in EXPECT['heading_order']]
    pos = 0
    for item in actual:
        if pos < len(expected) and (item == expected[pos] or item.endswith(expected[pos])):
            pos += 1
    assert pos == len(expected), f"Missing or out-of-order required headings: {_headings(doc)!r}"


def _assert_public_budget_table(doc):
    expected = EXPECT['public_budget_rows']
    for table in doc.tables:
        rows = _vopt_norm_rows(_table_rows(table))
        if not rows:
            continue
        body = rows[1:]
        row_texts = [' | '.join(row) for row in body]
        if all(any(exp[0] in row for row in row_texts) for exp in expected):
            for exp in expected:
                row = next(row for row in row_texts if exp[0] in row)
                for token in exp:
                    assert _vopt_norm_text(token) in row, f'Budget row {exp[0]} missing {token!r}'
            assert not any('P-04' in row or 'Gala catering' in row for row in row_texts), 'Internal budget line leaked to public table'
            return
    raise AssertionError('Missing public budget table with P-01/P-02/P-03')


def _row_record_from_headers(ws, row):
    header = [_vopt_norm_text(ws.cell(2, col).value) for col in range(1, ws.max_column + 1)]
    return {_vopt_norm_key(header[col - 1]): ws.cell(row, col).value for col in range(1, ws.max_column + 1)}


def _assert_formula_contains(cell, tokens, context):
    formula = _vopt_norm_formula(cell.value)
    assert formula.startswith('='), f'{context}: expected live formula, found {cell.value!r}'
    missing = [token for token in tokens if token.upper() not in formula]
    assert not missing, f'{context}: formula missing {missing!r}; found {cell.value!r}'


def test_outputs_exist():
    assert _doc_path().exists()
    assert _xlsx_path().exists()
    assert _doc_path().suffix.lower() == '.docx'
    assert _xlsx_path().suffix.lower() == '.xlsx'


def test_word_format_structure_and_public_boundary():
    doc = Document(_doc_path())
    section = doc.sections[0]
    assert int(section.top_margin) == EXPECT['margins']['top']
    assert int(section.bottom_margin) == EXPECT['margins']['bottom']
    assert int(section.left_margin) == EXPECT['margins']['left']
    assert int(section.right_margin) == EXPECT['margins']['right']
    _assert_heading_sequence(doc)
    text = _doc_text(doc)
    _assert_word_semantics(text)
    forbid_any(text, EXPECT['forbidden_text'], 'docx')
    assert _has_toc(_doc_path()), 'Expected real TOC field'
    assert EXPECT['header_contains'] in _vopt_header_text(doc)
    assert EXPECT['footer_contains'] in _vopt_footer_text(doc)
    _assert_public_budget_table(doc)
    assert _has_highlighted_format_exception(doc)


def test_workbook_format_controls():
    wb = load_workbook(_xlsx_path(), data_only=False)
    for sheet in EXPECT['sheet_order']:
        assert sheet in wb.sheetnames, f'Missing sheet: {sheet}'
    assert wb.sheetnames[:2] == EXPECT['sheet_order']
    ws = wb['Budget Summary']
    expected_by_id = {row[0]: row for row in EXPECT['workbook_rows']}
    for r in range(3, 7):
        record = _row_record_from_headers(ws, r)
        required = ['lineid', 'category', 'amount', 'fundingclass', 'narrativesection', 'publish']
        missing = [key for key in required if key not in record]
        assert not missing, f'Budget Summary row {r}: missing columns {missing!r}; actual={list(record.keys())!r}'
        line_id = _vopt_norm_text(record['lineid'])
        assert line_id in expected_by_id, f'Unexpected budget line: {line_id}'
        expected = expected_by_id[line_id]
        checks = {
            'category': expected[1],
            'amount': expected[2],
            'fundingclass': expected[3],
            'narrativesection': expected[4],
            'publish': expected[5],
        }
        for key, expected_value in checks.items():
            assert _vopt_norm_text(record[key]) == _vopt_norm_text(expected_value), (
                f'{line_id}: expected {key}={expected_value!r}, found {record[key]!r}'
            )
    assert _vopt_table_ref(ws, 'grant_budget_summary') == EXPECT['table_ref']
    _vopt_assert_print_area(ws, EXPECT['print_area'])
    assert ws.freeze_panes == EXPECT['freeze_panes']
    for row in range(3, 7):
        _assert_formula_contains(ws[f'G{row}'], [f'F{row}', 'INNARRATIVE', 'INTERNAL'], f'G{row}')
        _assert_formula_contains(ws[f'H{row}'], [f'C{row}', 'PRINT', 'CHECK'], f'H{row}')
    actual_ranges = _vopt_dv_ranges(ws)
    for rng in EXPECT['data_validation_ranges']:
        assert rng in actual_ranges, f'Missing data validation range: {rng}'
    assert len(ws.conditional_formatting) >= 1
    for sheet in EXPECT['hidden_sheets']:
        assert wb[sheet].sheet_state in ('hidden', 'veryHidden')
    names = {dn.name for dn in wb.defined_names.values()}
    for name in EXPECT['defined_names']:
        assert name in names


def test_cross_output_public_budget_lines():
    doc_text = _doc_text(Document(_doc_path()))
    wb = load_workbook(_xlsx_path(), data_only=False)
    assert 'Budget Summary' in wb.sheetnames, f"Missing sheet: Budget Summary; actual={wb.sheetnames!r}"
    ws = wb['Budget Summary']
    public_ids = [ws.cell(r, 1).value for r in range(3, 7) if ws.cell(r, 6).value == 'Yes']
    internal_ids = [ws.cell(r, 1).value for r in range(3, 7) if ws.cell(r, 6).value == 'No']
    assert public_ids == EXPECT['public_line_ids']
    assert internal_ids == EXPECT['internal_line_ids']
    for line_id in public_ids:
        assert line_id in doc_text, f'{line_id} missing from Word narrative'
    for line_id in internal_ids:
        assert line_id not in doc_text, f'{line_id} leaked to Word narrative'
