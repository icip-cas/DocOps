import json
import os
import re
import sys
import zipfile
from pathlib import Path
from datetime import date, datetime

from docx import Document
from openpyxl import load_workbook
from openpyxl.utils.cell import get_column_letter, range_boundaries

sys.path.insert(0, str(Path(__file__).parent))
from verifier_utils import *  # noqa: F401,F403

META_PATH = Path(os.environ.get('TASK_METADATA_PATH', '/tests/task_metadata.json'))
META = json.loads(META_PATH.read_text(encoding='utf-8'))
EXPECT = META['verifier_expectations']


def _vopt_norm_text(value):
    if value is None:
        return ''
    if isinstance(value, datetime):
        if value.hour or value.minute or value.second:
            return value.strftime('%Y-%m-%d %H:%M')
        return value.strftime('%Y-%m-%d')
    if isinstance(value, date):
        return value.strftime('%Y-%m-%d')
    if isinstance(value, float) and value.is_integer():
        value = int(value)
    text = str(value)
    text = text.replace('\u2013', '-').replace('\u2014', '-').replace('\u2212', '-')
    text = text.replace('\xa0', ' ')
    text = re.sub(r'\s+', ' ', text).strip()
    return text


def _vopt_norm_formula(value):
    return re.sub(r'\s+', '', _vopt_norm_text(value)).upper()


def _vopt_assert_formula(actual, expected, label='formula'):
    assert _vopt_norm_formula(actual) == _vopt_norm_formula(expected), f"{label}: expected {expected!r}, found {actual!r}"


def _vopt_norm_rows(rows):
    return [[_vopt_norm_text(cell) for cell in row] for row in rows]


def _vopt_rows_equal(actual, expected):
    return _vopt_norm_rows(actual) == _vopt_norm_rows(expected)


def _vopt_norm_key(value):
    return re.sub(r'[^a-z0-9]+', '', _vopt_norm_text(value).lower())


def _vopt_rows_text(rows):
    return '\n'.join('|'.join(_vopt_norm_text(cell) for cell in row) for row in rows)


def _vopt_header_has(header, required):
    actual = {_vopt_norm_text(cell).lower() for cell in header}
    missing = [cell for cell in required if _vopt_norm_text(cell).lower() not in actual]
    assert not missing, f"Missing expected header columns {missing!r}; actual={header!r}"


def _vopt_section_text(section_part):
    return '\n'.join(p.text for p in getattr(section_part, 'paragraphs', []) if p.text is not None)


def _vopt_header_text(doc):
    return '\n'.join(_vopt_section_text(section.header) for section in doc.sections)


def _vopt_footer_text(doc):
    return '\n'.join(_vopt_section_text(section.footer) for section in doc.sections)


def _vopt_table_ref(ws, expected_name=None):
    tables = ws.tables
    if expected_name:
        for name in tables.keys():
            if str(name).lower() == str(expected_name).lower():
                return tables[name].ref
    vals = list(tables.values())
    assert vals, f"No Excel native table found on sheet {ws.title!r}"
    return vals[0].ref


def _vopt_clean_area(value):
    text = _vopt_norm_text(value).replace("'", '').replace('$', '').replace(' ', '')
    if '!' in text:
        text = text.split('!', 1)[1]
    return text.upper()


def _vopt_assert_print_area(ws, expected):
    actual = _vopt_clean_area(ws.print_area)
    target = _vopt_clean_area(expected)
    assert target in actual or actual in target, f"{ws.title}: expected print area {expected!r}, found {ws.print_area!r}"


def _vopt_freeze_pane(value):
    return getattr(value, 'coordinate', value)


def _vopt_validated_cells(ws):
    cells = set()
    for dv in ws.data_validations.dataValidation:
        for rng in dv.cells.ranges:
            text = str(rng).replace('$', '')
            try:
                min_col, min_row, max_col, max_row = range_boundaries(text)
            except ValueError:
                cells.add(text)
                continue
            if max_row > 10000:
                max_row = max(min_row, 200)
            for row in range(min_row, max_row + 1):
                for col in range(min_col, max_col + 1):
                    cells.add(f"{get_column_letter(col)}{row}")
    return cells


class _VoptValidationRanges(list):
    def __init__(self, ws, ranges):
        super().__init__(ranges)
        self.ws = ws

    def __contains__(self, expected):
        expected = str(expected).replace('$', '')
        if super().__contains__(expected):
            return True
        try:
            min_col, min_row, max_col, max_row = range_boundaries(expected)
        except ValueError:
            return super().__contains__(expected)
        target = {f"{get_column_letter(col)}{row}" for row in range(min_row, max_row + 1) for col in range(min_col, max_col + 1)}
        return target.issubset(_vopt_validated_cells(self.ws))


def _vopt_dv_ranges(ws):
    ranges = []
    for dv in ws.data_validations.dataValidation:
        ranges.extend(str(rng).replace('$', '') for rng in dv.cells.ranges)
    return _VoptValidationRanges(ws, ranges)


def _vopt_ordered_subset(expected, actual):
    actual_iter = iter([_vopt_norm_text(item) for item in actual])
    for item in [_vopt_norm_text(value) for value in expected]:
        for candidate in actual_iter:
            if item == candidate:
                break
        else:
            return False
    return True



def _resolve_output(kind):
    if kind == 'response':
        return Path(os.environ.get('DOCX_OUTPUT_PATH', EXPECT['response_output']))
    if kind == 'tracker':
        return Path(os.environ.get('XLSX_OUTPUT_PATH', EXPECT['tracker_output']))
    raise KeyError(kind)


def _doc_heading_order(doc):
    out = []
    for p in doc.paragraphs:
        style = p.style.name if p.style else ''
        if p.text.strip() and (style.startswith('Heading') or style == 'Title'):
            out.append(p.text.strip())
    return out


def _doc_text(doc):
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append('|'.join(cell.text.strip() for cell in row.cells))
    return '\n'.join(parts)


def _table_rows(table):
    return [[cell.text.strip() for cell in row.cells] for row in table.rows]


def _date_variants(value):
    text = _vopt_norm_text(value)
    variants = {text}
    try:
        dt = datetime.strptime(text, '%Y-%m-%d')
    except ValueError:
        return variants
    variants.update({
        dt.strftime('%Y-%m-%d'),
        dt.strftime('%Y/%m/%d'),
        dt.strftime('%d %B %Y'),
        dt.strftime('%B %d, %Y'),
        f"{dt.strftime('%B')} {dt.day}, {dt.year}",
        dt.strftime('%d %b %Y'),
        dt.strftime('%b %d, %Y'),
        f"{dt.strftime('%b')} {dt.day}, {dt.year}",
    })
    return variants


def _row_record_from_headers(ws, row):
    header = [_vopt_norm_text(ws.cell(1, col).value) for col in range(1, ws.max_column + 1)]
    return {_vopt_norm_key(header[col - 1]): ws.cell(row, col).value for col in range(1, ws.max_column + 1)}


def _record_value(record, aliases):
    for alias in aliases:
        key = _vopt_norm_key(alias)
        if key in record:
            return record[key]
    return None


def _require_value(record, aliases, context):
    value = _record_value(record, aliases)
    assert value is not None, f'{context}: missing columns {aliases!r}; actual={list(record.keys())!r}'
    return value


def _assert_formula_contains(cell, required_tokens, context):
    formula = _vopt_norm_formula(cell.value)
    assert formula.startswith('='), f'{context}: expected live formula, found {cell.value!r}'
    missing = [token for token in required_tokens if token.upper() not in formula]
    assert not missing, f'{context}: formula missing {missing!r}; found {cell.value!r}'


def _assert_live_formula(cell, context):
    formula = _vopt_norm_formula(cell.value)
    assert formula.startswith('='), f'{context}: expected live formula, found {cell.value!r}'


def _decision_kind(value):
    text = normalize_text(_vopt_norm_text(value))
    public_tokens = ('yes', 'publish', 'publishable', 'external letter', 'in letter', 'include in letter')
    internal_tokens = ('no', 'internal', 'internal only', 'internal tracker', 'internal tracker only')
    if any(token in text for token in public_tokens) and 'internal' not in text:
        return 'public'
    if any(token in text for token in internal_tokens):
        return 'internal'
    return text


def _assert_response_table(doc):
    public_rows = EXPECT['response_rows']
    all_rows = []
    for table in doc.tables:
        for row in _table_rows(table)[1:]:
            row_text = _vopt_norm_text(' | '.join(row))
            all_rows.append(row_text)
    for expected in public_rows:
        oid, finding, capa, evidence, due = expected
        matches = [row for row in all_rows if oid in row]
        assert matches, f'Missing response-table row for {oid}'
        row = matches[0]
        for eid in [part.strip() for part in evidence.split(';')]:
            assert eid in row, f'{oid}: missing evidence {eid}'
        assert any(variant in row for variant in _date_variants(due)), f'{oid}: missing due date {due}'
        finding_tokens = [token for token in re.findall(r'[A-Za-z]{5,}', finding) if token.lower() != 'protocol']
        assert any(token.lower() in row.lower() for token in finding_tokens[:4]), (
            f'{oid}: row does not describe expected finding/action context'
        )
    for internal_id in EXPECT['internal_observation_ids']:
        assert not any(internal_id in row for row in all_rows), f'{internal_id} must not appear in response table'


def _highlighted_urgent_count(doc):
    count = 0
    for p in doc.paragraphs:
        text = _vopt_norm_text(p.text).lower()
        if not text:
            continue
        has_highlight = any(run.font.highlight_color is not None for run in p.runs) or '<w:highlight' in p._p.xml
        if has_highlight and ('urgent' in text or 'obs-01' in text or 'obs-02' in text):
            count += 1
    return count


def _has_toc_field(path):
    with zipfile.ZipFile(path) as zf:
        xml = zf.read('word/document.xml').decode('utf-8', errors='ignore')
    return bool(re.search(r'TOC\s+(?:\\)?o|TOC\s*(?:&quot;|")', xml))


def _highlighted_para(doc, text):
    for p in doc.paragraphs:
        if _vopt_norm_text(text) in _vopt_norm_text(p.text):
            return any(run.font.highlight_color is not None for run in p.runs) or '<w:highlight' in p._p.xml
    raise AssertionError(f'Paragraph not found: {text}')


def _row_values(ws, row, max_col):
    return [ws.cell(row, col).value for col in range(1, max_col + 1)]


def _row_highlighted(ws, row):
    for col in range(1, ws.max_column + 1):
        rgb = cell_fill_rgb(ws.cell(row, col))
        if rgb and rgb not in ('FFFFFF', '000000', '00000000'):
            return True
    return False


def test_outputs_exist():
    response = _resolve_output('response')
    tracker = _resolve_output('tracker')
    assert response.exists(), f'Missing response output: {response}'
    assert tracker.exists(), f'Missing tracker output: {tracker}'
    assert response.suffix.lower() == '.docx'
    assert tracker.suffix.lower() == '.xlsx'


def test_response_letter_structure_and_public_boundary():
    response = _resolve_output('response')
    doc = Document(response)
    headings_text = '\n'.join(_doc_heading_order(doc)).lower()
    for required in ['response', 'observation', 'evidence', 'commitment']:
        assert required in headings_text, f'Missing public section heading containing {required!r}; headings={_doc_heading_order(doc)!r}'
    text = _doc_text(doc)
    for oid in EXPECT['public_observation_ids']:
        assert oid in text, f'{oid} missing from response letter'
    for evidence_id in ['E-101A', 'E-101B', 'E-102A', 'E-104A']:
        assert evidence_id in text, f'{evidence_id} missing from response letter'
    forbid_any(text, ['Manual TOC placeholder', 'Draft Notes - Remove', 'CMT-03', 'E-103A', 'Internal sponsor note:'], 'response letter')
    assert _has_toc_field(response), 'Expected real TOC field'
    header_footer = _vopt_header_text(doc) + '\n' + _vopt_footer_text(doc)
    assert 'Draft' not in header_footer, 'Draft header/footer must be removed from final response letter'
    _assert_response_table(doc)
    assert _highlighted_urgent_count(doc) >= 2, 'Expected at least two highlighted urgent commitment paragraphs'


def test_tracker_workbook_structure_values_and_controls():
    wb = load_workbook(_resolve_output('tracker'), data_only=False)
    for sheet in EXPECT['tracker_sheet_order']:
        assert sheet in wb.sheetnames, f'Missing sheet: {sheet}'
    ws = wb['Commitment Tracker']
    expected_by_commitment = {row[0]: row for row in EXPECT['tracker_rows']}
    for row_idx in range(2, 6):
        record = _row_record_from_headers(ws, row_idx)
        commitment_id = _vopt_norm_text(_require_value(record, ['Commitment ID'], f'Commitment Tracker row {row_idx}'))
        assert commitment_id in expected_by_commitment, f'Unexpected commitment row: {commitment_id}'
        expected = expected_by_commitment[commitment_id]
        checks = [
            (['Observation ID'], expected[1], 'observation ID'),
            (['CAPA ID'], expected[2], 'CAPA ID'),
            (['Owner'], expected[4], 'owner'),
            (['Evidence IDs'], expected[8], 'evidence IDs'),
        ]
        for aliases, expected_value, label in checks:
            actual = _require_value(record, aliases, f'{commitment_id} {label}')
            assert _vopt_norm_text(actual) == _vopt_norm_text(expected_value), (
                f'{commitment_id}: expected {label}={expected_value!r}, found {actual!r}'
            )
        actual_status = _vopt_norm_text(_require_value(record, ['Status', 'Commitment Status'], f'{commitment_id} status'))
        expected_status = 'Open' if commitment_id == 'CMT-03' else expected[7]
        assert actual_status == _vopt_norm_text(expected_status), (
            f'{commitment_id}: expected status={expected_status!r}, found {actual_status!r}'
        )
        due = _require_value(record, ['Due Date'], f'{commitment_id} due date')
        assert _vopt_norm_text(due) == _vopt_norm_text(expected[6]), (
            f'{commitment_id}: expected due date {expected[6]!r}, found {due!r}'
        )
        decision = _decision_kind(_require_value(record, ['Publish in Letter', 'Publish Decision', 'Publish Boundary', 'Letter Inclusion'], f'{commitment_id} publish decision'))
        expected_decision = 'public' if expected[9] == 'Yes' else 'internal'
        assert decision == expected_decision, f'{commitment_id}: expected {expected_decision} publish decision, found {decision!r}'
        description = ' '.join(_vopt_norm_text(value) for value in record.values() if value is not None)
        assert len(description) > 20, f'{commitment_id}: tracker row should contain descriptive commitment/finding text'
    assert _vopt_table_ref(ws, 'site_commitment_tracker') == EXPECT['tracker_table_ref']
    ev = wb['Evidence Review']
    expected_evidence = {row[0]: row for row in EXPECT['evidence_rows']}
    for row_idx in range(2, 7):
        record = _row_record_from_headers(ev, row_idx)
        evidence_id = _vopt_norm_text(_require_value(record, ['Evidence ID'], f'Evidence Review row {row_idx}'))
        assert evidence_id in expected_evidence, f'Unexpected evidence row: {evidence_id}'
        expected = expected_evidence[evidence_id]
        checks = [
            (['Observation ID'], expected[1], 'observation ID'),
            (['Boundary', 'Evidence Boundary'], expected[3], 'boundary'),
            (['Status'], expected[4], 'status'),
        ]
        for aliases, expected_value, label in checks:
            actual = _require_value(record, aliases, f'{evidence_id} {label}')
            actual_text = _vopt_norm_text(actual)
            expected_text = _vopt_norm_text(expected_value)
            if evidence_id == 'E-102A' and label == 'status':
                assert 'Pending' in actual_text and 'signature' in actual_text, (
                    f'{evidence_id}: expected pending signature status, found {actual!r}'
                )
            else:
                assert actual_text == expected_text, (
                    f'{evidence_id}: expected {label}={expected_value!r}, found {actual!r}'
                )
        description = _record_value(record, ['Evidence Description', 'Evidence'])
        if description is not None:
            assert _vopt_norm_text(description), f'{evidence_id}: missing evidence description'
    assert _vopt_table_ref(ev, 'inspection_evidence_review') == EXPECT['evidence_table_ref']
    for row_idx in range(2, 6):
        _assert_live_formula(ws[f'K{row_idx}'], f'Commitment Tracker!K{row_idx}')
        _assert_live_formula(ws[f'L{row_idx}'], f'Commitment Tracker!L{row_idx}')
    for row_idx in range(2, 7):
        _assert_live_formula(ev[f'F{row_idx}'], f'Evidence Review!F{row_idx}')
    for sheet in EXPECT['hidden_sheets']:
        assert wb[sheet].sheet_state in ('hidden', 'veryHidden'), f'{sheet} should be hidden'
    names = {dn.name for dn in wb.defined_names.values()}
    for name in EXPECT['defined_names']:
        assert name in names, f'Missing defined name: {name}'
    actual_ranges = _vopt_dv_ranges(ws)
    for expected in EXPECT['data_validation_ranges']['Commitment Tracker']:
        assert expected in actual_ranges, f'Missing data validation range: {expected}'
    for row in EXPECT['highlight_rows']:
        assert _row_highlighted(ws, row), f'Expected highlighted tracker row {row}'


def test_cross_output_observation_consistency():
    doc = Document(_resolve_output('response'))
    response_text = _doc_text(doc)
    wb = load_workbook(_resolve_output('tracker'), data_only=False)
    ws = wb['Commitment Tracker']
    public_ids = []
    internal_ids = []
    for row_idx in range(2, 6):
        record = _row_record_from_headers(ws, row_idx)
        observation_id = _vopt_norm_text(_require_value(record, ['Observation ID'], f'Commitment Tracker row {row_idx} observation ID'))
        decision = _decision_kind(_require_value(record, ['Publish in Letter', 'Publish Decision', 'Publish Boundary', 'Letter Inclusion'], f'Commitment Tracker row {row_idx} publish decision'))
        if decision == 'public':
            public_ids.append(observation_id)
        elif decision == 'internal':
            internal_ids.append(observation_id)
    assert public_ids == EXPECT['public_observation_ids'], f"Unexpected public IDs: {public_ids!r}"
    assert internal_ids == EXPECT['internal_observation_ids'], f"Unexpected internal IDs: {internal_ids!r}"
    for oid in public_ids:
        assert oid in response_text, f'{oid} missing from response letter'
    for oid in internal_ids:
        assert oid not in response_text, f'{oid} must remain only in the internal tracker'
    for forbidden in ['OBS-03', 'CMT-03', 'E-103A', 'patient initials']:
        assert forbidden.lower() not in response_text.lower(), f'External response leaked internal material: {forbidden}'
