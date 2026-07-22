import json
import os
import re
from datetime import date, datetime
from pathlib import Path

from docx import Document
from openpyxl import load_workbook
from openpyxl.utils.cell import get_column_letter, range_boundaries

from verifier_utils import *  # noqa: F401,F403

META_PATH = Path(os.environ.get("TASK_METADATA_PATH", "/tests/task_metadata.json"))
META = json.loads(META_PATH.read_text(encoding="utf-8"))
EXPECT = META["verifier_expectations"]
PUBLIC_FORBIDDEN = EXPECT["forbidden_public"] + [
    "fire Luis",
    "tell manager",
    "credit issue",
    "extra medical deduction",
    "delay the IWO",
    "hold IWO",
    "ignore support",
    "HR risk note",
]


def _norm(value):
    if value is None:
        return ""
    if isinstance(value, datetime):
        return value.strftime("%Y-%m-%d")
    if isinstance(value, date):
        return value.strftime("%Y-%m-%d")
    if isinstance(value, float) and value.is_integer():
        value = int(value)
    text = str(value).replace("\u2013", "-").replace("\u2014", "-").replace("\xa0", " ")
    return re.sub(r"\s+", " ", text).strip()


def _norm_formula(value):
    return re.sub(r"\s+", "", _norm(value)).upper()


def _norm_key(value):
    return re.sub(r"[^A-Z0-9]+", "", _norm(value).upper())


def _rows(ws, start, end, max_col):
    return [[_norm(ws.cell(r, c).value) for c in range(1, max_col + 1)] for r in range(start, end + 1)]


def _expected(rows):
    return [[_norm(c) for c in row] for row in rows]


def _doc_text(doc):
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append("|".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts)


def _heading_order(doc):
    out = []
    expected = {_norm_heading_label(item) for item in EXPECT.get("doc_heading_order", [])}
    for p in doc.paragraphs:
        style = p.style.name if p.style else ""
        text = p.text.strip()
        if not text:
            continue
        if style == "Title" or style.startswith("Heading") or _norm_heading_label(text) in expected:
            out.append(text)
    return out


def _norm_heading_label(value):
    text = str(value or "")
    text = text.replace("\u2013", "-").replace("\u2014", "-").replace("\xa0", " ")
    text = re.sub(r"\s+", " ", text).strip()
    text = re.sub(r"^(?:\d+|[IVXLCDM]+)[\.)]\s+", "", text, flags=re.I)
    text = re.sub(r"^[A-Z][\.)]\s+", "", text)
    text = re.sub(r"[:\-]+$", "", text).strip()
    return text.casefold()


def _heading_sequence_matches(actual, expected):
    actual_norm = [_norm_heading_label(item) for item in actual]
    expected_norm = [_norm_heading_label(item) for item in expected]
    pos = 0
    for item in actual_norm:
        if pos < len(expected_norm) and item == expected_norm[pos]:
            pos += 1
    return pos == len(expected_norm)


def _header_footer_text(doc):
    headers, footers = [], []
    for section in doc.sections:
        headers.extend(p.text for p in section.header.paragraphs)
        footers.extend(p.text for p in section.footer.paragraphs)
    return "\n".join(headers), "\n".join(footers)


def _table_rows(table):
    return [[cell.text.strip() for cell in row.cells] for row in table.rows]


def _rows_text(rows):
    return "\n".join("|".join(_norm(cell) for cell in row) for row in rows)


def _row_text(row):
    return " ".join(_norm(cell) for cell in row if _norm(cell))


def _find_row_containing(rows, tokens, label):
    for row in rows:
        text = _row_text(row).casefold()
        if all(_norm(token).casefold() in text for token in tokens):
            return row
    raise AssertionError(f"{label}: no row contains required tokens {tokens!r}; rows={rows!r}")


def _has_row_containing(rows, tokens):
    try:
        _find_row_containing(rows, tokens, "rows")
        return True
    except AssertionError:
        return False


def _require_rows(rows, row_specs, label):
    for tokens in row_specs:
        _find_row_containing(rows, tokens, label)


def _find_doc_table_with_header(doc, required):
    for table in doc.tables:
        rows = _table_rows(table)
        if rows:
            header = {_norm_key(cell) for cell in rows[0]}
            if all(any(_norm_key(req) in item or item in _norm_key(req) for item in header) for req in required):
                return rows
    raise AssertionError(f"Missing DOCX table with columns {required!r}")


def _has_decision_table(doc):
    for table in doc.tables:
        rows = _table_rows(table)
        if not rows:
            continue
        text = normalize_text("\n".join(" | ".join(row) for row in rows))
        has_cs = "cs-22-1189" in text and "$730.00" in text and any(term in text for term in ["active", "withhold", "withholding"])
        has_cg = "cg-7741" in text and "$0.00" in text and any(term in text for term in ["answer", "no withholding", "withhold"])
        has_decision_context = any(term in text for term in ["decision", "reason", "priority", "cap", "support"])
        if has_cs and has_cg and has_decision_context:
            return True
    return False


def _require_memo_semantics(text):
    require_all(text, [
        "Luis Romero",
        "E-4187",
        "biweekly",
        "2026-07-10",
        "$2,515.20",
        "CS-22-1189",
        "$730.00",
        "55%",
        "$1,383.36",
        "CG-7741",
        "$0.00",
        "401(k)",
        "medical",
        "gym",
        "2026-07-08",
    ], "setup memo")
    norm = normalize_text(text)
    assert "620" in norm and "current" in norm and "110" in norm and "arrears" in norm, (
        "setup memo: missing current-support and arrears breakdown"
    )
    assert "child support" in norm and "priority" in norm, "setup memo: missing child-support priority semantics"
    assert "consumer garnishment" in norm and ("25%" in text or "ordinary" in norm), (
        "setup memo: missing ordinary consumer garnishment cap context"
    )


def _path(kind):
    env = {"docx": "DOCX_OUTPUT_PATH", "xlsx": "XLSX_OUTPUT_PATH"}[kind]
    key = {"docx": "docx_output", "xlsx": "xlsx_output"}[kind]
    return Path(os.environ.get(env, EXPECT[key]))


def _sheet(wb, name):
    assert name in wb.sheetnames, f"Missing required worksheet {name!r}; found {wb.sheetnames!r}"
    return wb[name]


def _table(ws, expected_name):
    for name in ws.tables.keys():
        if str(name).lower() == str(expected_name).lower():
            return ws.tables[name]
    raise AssertionError(f"Missing table {expected_name!r} on {ws.title!r}; found {list(ws.tables.keys())!r}")


def _table_has_size(ws, expected_name, min_rows, min_cols):
    table = _table(ws, expected_name)
    min_col, min_row, max_col, max_row = range_boundaries(table.ref)
    actual_rows = max_row - min_row + 1
    actual_cols = max_col - min_col + 1
    assert actual_rows >= min_rows and actual_cols >= min_cols, (
        f"{ws.title}.{expected_name}: expected at least {min_rows} rows x {min_cols} cols, "
        f"found {actual_rows} rows x {actual_cols} cols at {table.ref}"
    )
    return table


def _table_data_rows(ws, table_name):
    table = _table(ws, table_name)
    min_col, min_row, max_col, max_row = range_boundaries(table.ref)
    return [
        [ws.cell(row, col).value for col in range(min_col, max_col + 1)]
        for row in range(min_row + 1, max_row + 1)
    ]


def _sheet_has_formula_tokens(ws, tokens, label):
    for row in ws.iter_rows():
        for cell in row:
            value = cell.value
            if isinstance(value, str) and value.startswith("="):
                norm = _norm_formula(value)
                if all(_norm_formula(token) in norm for token in tokens):
                    return
    raise AssertionError(f"{label}: no formula found with purpose tokens {tokens!r}")


def _sheet_has_formula_option(ws, token_options, label):
    for tokens in token_options:
        try:
            _sheet_has_formula_tokens(ws, tokens, label)
            return
        except AssertionError:
            continue
    raise AssertionError(f"{label}: no formula found with any purpose tokens {token_options!r}")


def _clean_area(value):
    text = _norm(value).replace("'", "").replace("$", "").replace(" ", "")
    if "!" in text:
        text = text.split("!", 1)[1]
    return text.upper()


def _assert_print_area(ws, expected):
    actual = _clean_area(ws.print_area)
    target = _clean_area(expected)
    if target in actual or actual in target:
        return
    try:
        actual_min_col, actual_min_row, actual_max_col, actual_max_row = range_boundaries(actual)
        target_min_col, target_min_row, target_max_col, target_max_row = range_boundaries(target)
    except ValueError:
        raise AssertionError(f"{ws.title}: expected print area {expected!r}, found {ws.print_area!r}")
    assert (
        actual_min_col <= target_min_col
        and actual_min_row <= target_min_row
        and actual_max_col >= target_max_col
        and actual_max_row >= target_max_row
    ), f"{ws.title}: expected print area {expected!r}, found {ws.print_area!r}"


def _range_covered(ws, expected):
    expected = expected.replace("$", "")
    min_col, min_row, max_col, max_row = range_boundaries(expected)
    target = {f"{get_column_letter(col)}{row}" for row in range(min_row, max_row + 1) for col in range(min_col, max_col + 1)}
    cells = set()
    for dv in ws.data_validations.dataValidation:
        for rng in dv.cells.ranges:
            min_c, min_r, max_c, max_r = range_boundaries(str(rng).replace("$", ""))
            cells.update(f"{get_column_letter(col)}{row}" for row in range(min_r, max_r + 1) for col in range(min_c, max_c + 1))
    return target.issubset(cells)


def _sheet_text(wb, sheets):
    parts = []
    for sheet in sheets:
        ws = _sheet(wb, sheet)
        for row in ws.iter_rows():
            for cell in row:
                if cell.value is not None:
                    parts.append(str(cell.value))
    return "\n".join(parts)


def test_outputs_exist():
    assert _path("docx").exists()
    assert _path("xlsx").exists()
    assert _path("docx").suffix.lower() == ".docx"
    assert _path("xlsx").suffix.lower() == ".xlsx"


def test_memo_structure_content_and_privacy():
    doc = Document(_path("docx"))
    assert _heading_sequence_matches(_heading_order(doc), EXPECT["doc_heading_order"])
    text = _doc_text(doc)
    _require_memo_semantics(text)
    forbid_any(text, PUBLIC_FORBIDDEN, "setup memo")
    header, footer = _header_footer_text(doc)
    assert EXPECT["doc_header"] in header or EXPECT["doc_header"] in text
    assert EXPECT["doc_footer"] in footer
    decision = _find_doc_table_with_header(doc, ["Order", "Decision", "Amount", "Reason"])
    decision_text = normalize_text(_rows_text(decision))
    assert "cs-22-1189" in decision_text and "$730.00" in decision_text and "active" in decision_text
    assert "cg-7741" in decision_text and "$0.00" in decision_text and "answer" in decision_text


def test_workbook_structure_formulas_and_controls():
    wb = load_workbook(_path("xlsx"), data_only=False)
    assert wb.sheetnames == EXPECT["sheet_order"]
    min_table_rows = {
        "Appeal Summary": 11,
        "Disposable Earnings": 11,
        "Orders Priority": 3,
        "Paycheck Simulation": 7,
        "Court Response Log": 4,
        "Remittance Schedule": 3,
        "Privacy Review": 4,
        "Raw Payroll": 2,
        "Rules": 5,
    }
    for sheet, (table_name, ref) in EXPECT["tables"].items():
        min_col, min_row, max_col, max_row = range_boundaries(ref)
        _table_has_size(_sheet(wb, sheet), table_name, min_table_rows[sheet], max_col - min_col + 1)
    for sheet in EXPECT["hidden_sheets"]:
        assert _sheet(wb, sheet).sheet_state in ("hidden", "veryHidden")
    names = {dn.name for dn in wb.defined_names.values()}
    for name in EXPECT["defined_names"]:
        assert name in names
    _sheet_has_formula_tokens(_sheet(wb, "Appeal Summary"), ["Disposable Earnings", "B7"], "required deductions summary formula")
    _sheet_has_formula_tokens(_sheet(wb, "Appeal Summary"), ["DisposableEarnings"], "disposable earnings summary formula")
    _sheet_has_formula_tokens(_sheet(wb, "Appeal Summary"), ["ChildSupportWithholding"], "child-support summary formula")
    _sheet_has_formula_tokens(_sheet(wb, "Appeal Summary"), ["ConsumerWithholding"], "consumer summary formula")
    _sheet_has_formula_option(_sheet(wb, "Appeal Summary"), [
        ["IF", "Ready"],
        ["IF", "Hold"],
        ["IF", "Yes"],
        ["IF", "No"],
        ["IF", "ChildSupportWithholding"],
    ], "setup readiness formula")
    _sheet_has_formula_tokens(_sheet(wb, "Disposable Earnings"), ["SUM", "B3"], "disposable earnings formula")
    _sheet_has_formula_tokens(_sheet(wb, "Orders Priority"), ["DisposableEarnings"], "child-support cap formula")
    _sheet_has_formula_tokens(_sheet(wb, "Orders Priority"), ["MIN"], "approved child-support withholding formula")
    _sheet_has_formula_tokens(_sheet(wb, "Orders Priority"), ["MIN"], "consumer garnishment cap formula")
    _sheet_has_formula_tokens(_sheet(wb, "Orders Priority"), ["MAX", "0"], "consumer approved withholding formula")
    _sheet_has_formula_tokens(_sheet(wb, "Paycheck Simulation"), ["ChildSupportWithholding"], "paycheck child-support formula")
    _sheet_has_formula_tokens(_sheet(wb, "Paycheck Simulation"), ["ConsumerWithholding"], "paycheck consumer formula")
    _sheet_has_formula_tokens(_sheet(wb, "Paycheck Simulation"), ["SUM"], "net after garnishment formula")
    _sheet_has_formula_tokens(_sheet(wb, "Remittance Schedule"), ["ChildSupportWithholding"], "child-support remittance formula")
    _sheet_has_formula_tokens(_sheet(wb, "Remittance Schedule"), ["ConsumerWithholding"], "consumer remittance formula")
    for sheet, ranges in EXPECT["data_validation"].items():
        for rng in ranges:
            assert _range_covered(_sheet(wb, sheet), rng), f"{sheet}: missing validation over {rng}"
    for sheet, area in EXPECT["print_areas"].items():
        ws = _sheet(wb, sheet)
        table_name = EXPECT["tables"][sheet][0]
        _assert_print_area(ws, _table(ws, table_name).ref)


def test_workbook_values_and_privacy():
    wb = load_workbook(_path("xlsx"), data_only=False)
    summary = _table_data_rows(_sheet(wb, "Appeal Summary"), "tblSetupSummary")
    _require_rows(summary, [
        ["Employee", "Luis Romero"],
        ["Employee ID", "E-4187"],
        ["Pay frequency", "Biweekly"],
        ["First affected pay date", "2026-07-10"],
        ["Gross wages"],
        ["Legally required deductions"],
        ["Disposable earnings"],
        ["Child", "Support", "Withholding"],
        ["Consumer garnishment withholding"],
        ["Court answer due", "2026-07-08"],
        ["Setup", "Ready"],
    ], "Appeal Summary")

    disposable = _table_data_rows(_sheet(wb, "Disposable Earnings"), "tblDisposable")
    _require_rows(disposable, [
        ["Gross wages"],
        ["Federal", "Tax"],
        ["State", "Tax"],
        ["Social Security"],
        ["Medicare"],
        ["Disposable earnings"],
        ["Voluntary", "401"],
        ["Medical", "Premium"],
        ["Gym"],
    ], "Disposable Earnings")

    orders = _table_data_rows(_sheet(wb, "Orders Priority"), "tblOrdersPriority")
    _require_rows(orders, [
        ["CS-22-1189", "Child support", "1", "730", "55%", "Active"],
        ["CG-7741", "Consumer", "2", "410", "Answer"],
    ], "Orders Priority")
    orders_text = normalize_text(_rows_text(orders))
    assert "support" in orders_text and "priority" in orders_text

    simulation = _table_data_rows(_sheet(wb, "Paycheck Simulation"), "tblPaycheckSimulation")
    _require_rows(simulation, [
        ["Gross wages"],
        ["Disposable earnings"],
        ["Child support"],
            ["CG-7741"],
            ["Net"],
    ], "Paycheck Simulation")
    assert (
        _has_row_containing(simulation, ["Required deductions"])
        or all(_has_row_containing(simulation, tokens) for tokens in [
            ["Federal", "Tax"],
            ["State", "Tax"],
            ["Social Security"],
            ["Medicare"],
        ])
    ), "Paycheck Simulation: missing required-deduction lines"

    response = _table_data_rows(_sheet(wb, "Court Response Log"), "tblCourtResponse")
    _require_rows(response, [
        ["Court", "2026-07-08", "CG-7741"],
        ["Employee notice", "Luis Romero"],
        ["CS-22-1189"],
    ], "Court Response Log")

    remittance = _table_data_rows(_sheet(wb, "Remittance Schedule"), "tblRemittance")
    _require_rows(remittance, [
        ["2026-07-10", "CS-22-1189", "California State Disbursement Unit"],
        ["2026-07-10", "CG-7741", "County Court Clerk"],
    ], "Remittance Schedule")

    privacy = _table_data_rows(_sheet(wb, "Privacy Review"), "tblPrivacyReview")
    _require_rows(privacy, [
        ["Internal", "Pass"],
        ["Manager", "Pass"],
        ["Court", "Pass"],
    ], "Privacy Review")

    raw_payroll = _table_data_rows(_sheet(wb, "Raw Payroll"), "tblRawPayroll")
    _require_rows(raw_payroll, [
        ["2026-07-10", "3200", "320", "120", "198.4", "46.4", "160", "145", "20"],
    ], "Raw Payroll")

    rules = _table_data_rows(_sheet(wb, "Rules"), "tblRules")
    _require_rows(rules, [
        ["Federal minimum wage", "7.25"],
        ["Ordinary", "Cap"],
    ], "Rules")
    assert (
        _has_row_containing(rules, ["Child", "Cap"])
        or _has_row_containing(rules, ["Support", "Cap"])
        or _has_row_containing(rules, ["CS", "Cap"])
    ), "Rules: missing child-support cap rule"
    public = _sheet_text(wb, ["Appeal Summary", "Disposable Earnings", "Orders Priority", "Paycheck Simulation", "Court Response Log", "Remittance Schedule", "Privacy Review"])
    forbid_any(public, PUBLIC_FORBIDDEN, "public workbook sheets")


def test_cross_output_consistency_and_decision_logic():
    doc_text = _doc_text(Document(_path("docx")))
    wb = load_workbook(_path("xlsx"), data_only=False)
    for anchor in ["Luis Romero", "E-4187", "2026-07-10", "CS-22-1189", "CG-7741", "$730.00", "$0.00", "2026-07-08"]:
        assert anchor in doc_text
    public_text = _sheet_text(wb, ["Appeal Summary", "Disposable Earnings", "Orders Priority", "Paycheck Simulation", "Court Response Log", "Remittance Schedule", "Privacy Review"])
    for anchor in ["Luis Romero", "E-4187", "2026-07-10", "CS-22-1189", "CG-7741", "730", "2026-07-08"]:
        assert anchor in public_text
    orders = _table_data_rows(_sheet(wb, "Orders Priority"), "tblOrdersPriority")
    _find_row_containing(orders, ["CS-22-1189", "Active"], "Orders Priority")
    _find_row_containing(orders, ["CG-7741", "Answer"], "Orders Priority")
    _sheet_has_formula_tokens(_sheet(wb, "Disposable Earnings"), ["SUM", "B3"], "disposable earnings formula")
    _sheet_has_formula_tokens(_sheet(wb, "Orders Priority"), ["MIN"], "approved child-support withholding formula")
    _sheet_has_formula_tokens(_sheet(wb, "Orders Priority"), ["MAX", "0"], "consumer approved withholding formula")
    for forbidden in PUBLIC_FORBIDDEN:
        assert normalize_text(forbidden) not in normalize_text(doc_text)
