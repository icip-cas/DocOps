import json
import os
import re
from datetime import date, datetime
from pathlib import Path

from docx import Document
from openpyxl import load_workbook
from openpyxl.utils.cell import get_column_letter, range_boundaries
from pypdf import PdfReader

from verifier_utils import *  # noqa: F401,F403

META_PATH = Path(os.environ.get("TASK_METADATA_PATH", "/tests/task_metadata.json"))
META = json.loads(META_PATH.read_text(encoding="utf-8"))
EXPECT = META["verifier_expectations"]


def _norm(value):
    if value is None:
        return ""
    if isinstance(value, datetime):
        return value.strftime("%Y-%m-%d")
    if isinstance(value, date):
        return value.strftime("%Y-%m-%d")
    if isinstance(value, float) and value.is_integer():
        value = int(value)
    return re.sub(r"\s+", " ", str(value).replace("\u2013", "-").replace("\u2014", "-").replace("\xa0", " ")).strip()


def _norm_formula(value):
    return re.sub(r"\s+", "", _norm(value)).upper()


def _norm_key(value):
    return re.sub(r"[^a-z0-9]+", "", _norm(value).casefold())


def _norm_heading_label(value):
    text = _norm(value)
    text = re.sub(r"^(?:\d+|[IVXLCDM]+)[\.)]\s+", "", text, flags=re.I)
    text = re.sub(r"^[A-Z][\.)]\s+", "", text)
    text = re.sub(r"[:\-]+$", "", text).strip()
    return text.casefold()


def _path(kind):
    env = {"docx": "DOCX_OUTPUT_PATH", "pdf": "PDF_OUTPUT_PATH", "xlsx": "XLSX_OUTPUT_PATH"}[kind]
    key = {"docx": "docx_output", "pdf": "pdf_output", "xlsx": "xlsx_output"}[kind]
    return Path(os.environ.get(env, EXPECT[key]))


def _doc_text(doc):
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append("|".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts)


def _heading_order(doc):
    out = []
    first = next((p.text.strip() for p in doc.paragraphs if p.text.strip()), "")
    if first:
        out.append(first)
    for p in doc.paragraphs:
        style = p.style.name if p.style else ""
        if p.text.strip() and (style == "Title" or style.startswith("Heading")):
            if not out or _norm_heading_label(p.text) != _norm_heading_label(out[-1]):
                out.append(p.text.strip())
    return out


def _heading_sequence_matches(actual, expected):
    actual_norm = [_norm_heading_label(item) for item in actual]
    expected_norm = [_norm_heading_label(item) for item in expected]
    pos = 0
    for item in actual_norm:
        if pos < len(expected_norm) and item == expected_norm[pos]:
            pos += 1
    return pos == len(expected_norm)


def _pdf_text(path):
    reader = PdfReader(str(path))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _seconds(ts):
    h, m, rest = ts.split(":")
    s, ms = rest.split(".")
    return int(h) * 3600 + int(m) * 60 + int(s) + int(ms) / 1000


def _forbid_private_and_raw_labels(text, label):
    private_phrases = [p for p in EXPECT["forbidden"] if not re.fullmatch(r"[A-Z]+:", p)]
    forbid_any(text, private_phrases, label)
    hits = re.findall(r"(?m)^\s*(HOST|MAYA|FIELD|LISTENER|PRODUCER):", text)
    assert not hits, f"{label}: raw speaker labels still present: {hits}"


def _require_caption_cues(text, label):
    norm = normalize_text(text).replace("->", " - ").replace("→", " - ")
    starts = []
    ends = []
    for _, start, end, caption, _ in EXPECT["cue_rows"][1:]:
        assert normalize_text(start) in norm, f"{label}: missing cue start {start}"
        assert normalize_text(end) in norm, f"{label}: missing cue end {end}"
        assert normalize_text(caption) in norm, f"{label}: missing cue text {caption}"
        starts.append(_seconds(start))
        ends.append(_seconds(end))
    assert starts == sorted(starts)
    for i in range(1, len(starts)):
        assert starts[i] >= ends[i - 1]


def _require_pdf_show_notes(text):
    for phrase in ["Harbor Signals Episode 07", "Repair Night", "HS-07"]:
        assert normalize_text(phrase) in normalize_text(text), f"show notes pdf missing {phrase!r}"
    for phrase in ["Maya Torres", "repair cafe", "small appliances", "repair skills", "Blue Harbor Library"]:
        assert normalize_text(phrase) in normalize_text(text), f"show notes pdf missing {phrase!r}"
    for _, timecode, title, _ in EXPECT["chapter_rows"][1:]:
        assert normalize_text(timecode) in normalize_text(text), f"show notes pdf missing chapter time {timecode}"
        assert normalize_text(title) in normalize_text(text), f"show notes pdf missing chapter title {title}"
    norm = normalize_text(text)
    assert ("caption" in norm or "access transcript" in norm) and ("access" in norm or "accessible" in norm)


def _expected_header(rows):
    return [_norm(c) for c in rows[0]]


def _sheet(wb, expected_name):
    _assert_required_sheets(wb, [expected_name])
    return wb[expected_name]


def _table(ws, expected_name):
    for name in ws.tables.keys():
        if str(name).lower() == str(expected_name).lower():
            return ws.tables[name]
    raise AssertionError(f"Missing table {expected_name!r} on {ws.title!r}")


def _table_bounds(table):
    return range_boundaries(table.ref.replace("$", ""))


def _table_headers(ws, table):
    min_col, min_row, max_col, _ = _table_bounds(table)
    return [_norm(ws.cell(min_row, col).value) for col in range(min_col, max_col + 1)]


def _table_header_map(ws, table):
    min_col, _, _, _ = _table_bounds(table)
    return {_norm_key(header): idx for idx, header in enumerate(_table_headers(ws, table), start=min_col)}


def _require_headers(ws, table, expected_headers):
    actual = {_norm_key(header) for header in _table_headers(ws, table)}
    missing = [header for header in expected_headers if _norm_key(header) not in actual]
    assert not missing, f"{ws.title}: table {table.name} missing headers {missing}; found {_table_headers(ws, table)}"


def _table_has_size(table, min_data_rows, min_cols):
    min_col, min_row, max_col, max_row = _table_bounds(table)
    assert max_col - min_col + 1 >= min_cols, f"{table.name}: expected at least {min_cols} columns"
    assert max_row - min_row >= min_data_rows, f"{table.name}: expected at least {min_data_rows} data rows"


def _table_rows(ws, table):
    _, min_row, _, max_row = _table_bounds(table)
    header_map = _table_header_map(ws, table)
    rows = []
    for row_idx in range(min_row + 1, max_row + 1):
        row = {"__row_idx": row_idx}
        for key, col_idx in header_map.items():
            row[key] = ws.cell(row_idx, col_idx)
        if any(cell.value is not None for key, cell in row.items() if key != "__row_idx"):
            rows.append(row)
    return rows


def _cell_by_header(row, header):
    key = _norm_key(header)
    assert key in row, f"Missing header {header!r}"
    return row[key]


def _value_matches(actual, expected):
    actual_text = _norm(actual)
    expected_text = _norm(expected)
    if actual_text.casefold() == expected_text.casefold():
        return True
    if len(expected_text) >= 12 and expected_text.casefold() in actual_text.casefold():
        return True
    return False


FIELD_ALIASES = {
    "episode": {"episode", "episode title", "title"},
    "caption_packet": {"caption_packet", "caption packet", "output 1", "output: caption transcript"},
    "show_notes": {"show_notes", "show notes", "output 2", "output: public show notes"},
    "timing_qc": {"timing_qc", "timing qc", "output 3", "output: timing qc"},
    "cue_count": {"cue_count", "cue count"},
    "chapter_count": {"chapter_count", "chapter count"},
    "music_license": {"music_license", "music license"},
}


def _key_value_matches(actual, expected):
    if _value_matches(actual, expected):
        return True
    expected_key = _norm_key(expected)
    actual_key = _norm_key(actual)
    aliases = {_norm_key(item) for item in FIELD_ALIASES.get(_norm(expected).casefold(), set())}
    return expected_key in aliases and actual_key in aliases


def _find_row(rows, key_header, key_value):
    for row in rows:
        if _key_value_matches(_cell_by_header(row, key_header).value, key_value):
            return row
    raise AssertionError(f"Could not find row with {key_header!r}={key_value!r}")


def _assert_row_contains(rows, headers, expected_row, key_header=None, skip_headers=()):
    key_header = key_header or headers[0]
    row = _find_row(rows, key_header, expected_row[headers.index(key_header)])
    for header, expected_value in zip(headers, expected_row):
        if header in skip_headers:
            continue
        actual = _cell_by_header(row, header).value
        assert _value_matches(actual, expected_value), f"{header}: expected {expected_value!r}, found {actual!r}"
    return row


def _range_covered(ws, expected):
    min_col, min_row, max_col, max_row = range_boundaries(expected.replace("$", ""))
    target = {f"{get_column_letter(col)}{row}" for row in range(min_row, max_row + 1) for col in range(min_col, max_col + 1)}
    cells = set()
    for dv in ws.data_validations.dataValidation:
        for rng in dv.cells.ranges:
            min_c, min_r, max_c, max_r = range_boundaries(str(rng).replace("$", ""))
            cells.update(f"{get_column_letter(col)}{row}" for row in range(min_r, max_r + 1) for col in range(min_c, max_c + 1))
    return target.issubset(cells)


def _clean_area(value):
    if isinstance(value, (list, tuple)):
        value = ",".join(str(item) for item in value)
    text = _norm(value).replace("'", "").replace("$", "").replace(" ", "")
    if "!" in text:
        text = text.split("!", 1)[1]
    return text.upper()


def _area_covers(actual, target):
    actual = _clean_area(actual)
    target = _clean_area(target)
    if actual == target or target in actual or actual in target:
        return True
    try:
        a_min_c, a_min_r, a_max_c, a_max_r = range_boundaries(actual)
        t_min_c, t_min_r, t_max_c, t_max_r = range_boundaries(target)
    except ValueError:
        return False
    return a_min_c <= t_min_c and a_min_r <= t_min_r and a_max_c >= t_max_c and a_max_r >= t_max_r


def _sheet_text(wb, sheets):
    _assert_required_sheets(wb, sheets)
    parts = []
    for sheet in sheets:
        for row in wb[sheet].iter_rows():
            for cell in row:
                if cell.value is not None:
                    parts.append(str(cell.value))
    return "\n".join(parts)


def _assert_required_sheets(wb, sheets):
    missing = [sheet for sheet in sheets if sheet not in wb.sheetnames]
    assert not missing, f"Missing required workbook sheets: {missing}; found {wb.sheetnames}"


def _formula_refs(value, *refs):
    formula = _norm_formula(value)
    assert formula.startswith("="), f"expected formula, got {value!r}"
    for ref in refs:
        assert _norm_formula(ref) in formula, f"formula {value!r} does not reference {ref}"


def _formula_targets(value, sheet_name, table_name):
    formula = _norm_formula(value)
    assert formula.startswith("="), f"expected formula, got {value!r}"
    refs_target = sheet_name.replace(" ", "").upper() in formula.replace("'", "") or table_name.upper() in formula
    assert refs_target, f"formula {value!r} must reference {sheet_name} or {table_name}"


def _defined_name_names(wb):
    try:
        values = wb.defined_names.values()
    except AttributeError:
        values = getattr(wb.defined_names, "definedName", [])
    return {dn.name for dn in values if getattr(dn, "name", None)}


def _ready_formula_ok(value):
    formula = _norm_formula(value)
    assert formula.startswith("="), "package-ready control must be a formula"
    for token in ("READY", "HOLD"):
        assert token in formula, f"package-ready formula missing {token}"
    assert "IF" in formula or "IFS" in formula, "package-ready formula must be conditional"
    assert "AND" in formula or formula.count("=") >= 3, "package-ready formula must combine all readiness checks"
    assert any(token in formula for token in ("B4=10", "B6=10", "CUECOUNT=10", "CUETIMING", "TBLCUETIMING", "CUES")), "package-ready formula must check ten cues"
    assert any(token in formula for token in ("B5=5", "B7=5", "CHAPTERCOUNT=5", "TBLCHAPTERS", "CHAPTER")), "package-ready formula must check five chapters"
    assert any(token in formula for token in ("B6=6", "PRIVATE", "EXCLUDED")), "package-ready formula must check six private-note exclusions"


def _status_validation_ok(ws, table_name, status_header):
    table = _table(ws, table_name)
    _, min_row, _, max_row = _table_bounds(table)
    col_idx = _table_header_map(ws, table)[_norm_key(status_header)]
    target = f"{get_column_letter(col_idx)}{min_row + 1}:{get_column_letter(col_idx)}{max_row}"
    assert _range_covered(ws, target), f"{ws.title}: missing validation over {target}"


def _summary_rows(wb):
    ws = _sheet(wb, "Package Summary")
    table = _table(ws, "tblPackageSummary")
    return _table_rows(ws, table)


def _summary_value_cell(wb, field):
    return _cell_by_header(_find_row(_summary_rows(wb), "Field", field), "Value")


def _status_ok(value, expected):
    actual = _norm(value).casefold()
    return actual in {_norm(expected).casefold(), "approved", "ready", "pass", "passed", "ok", "complete", "completed"}


def _table_key(sheet):
    return {
        "Package Summary": "summary_rows",
        "Cue Timing QC": "cue_rows",
        "Chapter Markers": "chapter_rows",
        "Publish Manifest": "manifest_rows",
        "QA Checks": "qa_rows",
        "Private Notes": "private_rows",
    }[sheet]


def test_outputs_exist():
    assert _path("docx").exists()
    assert _path("pdf").exists()
    assert _path("xlsx").exists()
    assert _path("docx").suffix.lower() == ".docx"
    assert _path("pdf").suffix.lower() == ".pdf"
    assert _path("xlsx").suffix.lower() == ".xlsx"


def test_caption_docx_content_timing_privacy():
    doc = Document(_path("docx"))
    assert _heading_sequence_matches(_heading_order(doc), EXPECT["heading_order"])
    text = _doc_text(doc)
    for phrase in ["HS-07", "Maya Torres", "00:04:09.000", "Open captions", "Speaker labels"]:
        assert normalize_text(phrase) in normalize_text(text), f"caption packet missing {phrase!r}"
    _require_caption_cues(text, "caption packet")
    _forbid_private_and_raw_labels(text, "caption packet")


def test_show_notes_pdf_content_privacy():
    text = _pdf_text(_path("pdf"))
    _require_pdf_show_notes(text)
    _forbid_private_and_raw_labels(text, "show notes pdf")


def test_workbook_structure_values_controls():
    wb = load_workbook(_path("xlsx"), data_only=False)
    assert wb.sheetnames == EXPECT["sheet_order"]
    for sheet, (table_name, _ref) in EXPECT["tables"].items():
        ws = _sheet(wb, sheet)
        table = _table(ws, table_name)
        expected_rows = EXPECT[_table_key(sheet)]
        _require_headers(ws, table, _expected_header(expected_rows))
        _table_has_size(table, len(expected_rows) - 1, len(expected_rows[0]))
    for sheet in EXPECT["hidden_sheets"]:
        assert _sheet(wb, sheet).sheet_state in ("hidden", "veryHidden")
    names = _defined_name_names(wb)
    for name in EXPECT["defined_names"]:
        assert name in names
    _formula_targets(_summary_value_cell(wb, "Cue count").value, "Cue Timing QC", "tblCueTiming")
    _formula_targets(_summary_value_cell(wb, "Chapter count").value, "Chapter Markers", "tblChapters")
    _ready_formula_ok(_summary_value_cell(wb, "Package ready").value)
    _status_validation_ok(_sheet(wb, "Cue Timing QC"), "tblCueTiming", "QC")
    _status_validation_ok(_sheet(wb, "Chapter Markers"), "tblChapters", "Status")
    _status_validation_ok(_sheet(wb, "QA Checks"), "tblQA", "Status")
    for sheet, area in EXPECT["print_areas"].items():
        assert _area_covers(_sheet(wb, sheet).print_area, area)
    summary_rows = _summary_rows(wb)
    for expected_row in EXPECT["summary_rows"][1:3]:
        row = _find_row(summary_rows, "Field", expected_row[0])
        assert _value_matches(_cell_by_header(row, "Value").value, expected_row[1])
    try:
        private_count = _summary_value_cell(wb, "Private notes excluded").value
        if isinstance(private_count, str) and private_count.startswith("="):
            formula = _norm_formula(private_count)
            assert "PRIVATE" in formula and any(token in formula for token in ("COUNT", "COUNTA", "ROWS", "SUBTOTAL")), "private-notes-excluded formula must count private exclusions"
        else:
            assert _norm(private_count) == "6", "private-notes-excluded count must be 6"
    except AssertionError:
        ready_formula = _norm_formula(_summary_value_cell(wb, "Package ready").value)
        assert "PRIVATE" in ready_formula and any(token in ready_formula for token in ("COUNT", "COUNTA", "ROWS", "SUBTOTAL"))
    cue_rows = _table_rows(_sheet(wb, "Cue Timing QC"), _table(_sheet(wb, "Cue Timing QC"), "tblCueTiming"))
    for expected_row in EXPECT["cue_rows"][1:]:
        _assert_row_contains(cue_rows, EXPECT["cue_rows"][0], expected_row, key_header="Cue")
    chapter_rows = _table_rows(_sheet(wb, "Chapter Markers"), _table(_sheet(wb, "Chapter Markers"), "tblChapters"))
    for expected_row in EXPECT["chapter_rows"][1:]:
        row = _assert_row_contains(chapter_rows, EXPECT["chapter_rows"][0], expected_row, key_header="Chapter", skip_headers=("Status",))
        assert _status_ok(_cell_by_header(row, "Status").value, "Ready")
    manifest_text = normalize_text(_sheet_text(wb, ["Publish Manifest"]))
    for phrase in [
        "HS-07",
        "Harbor Signals Episode 07: Repair Night",
        "00:04:09.000",
        "episode_07_caption_transcript_packet.docx",
        "episode_07_public_show_notes.pdf",
        "episode_07_timing_qc.xlsx",
        "Blue Harbor Library",
    ]:
        assert normalize_text(phrase) in manifest_text
    qa_text = normalize_text(_sheet_text(wb, ["QA Checks"]))
    for token in ("cue", "chapter", "overlap", "speaker", "private", "docx", "pdf", "xlsx"):
        assert token in qa_text
    qa_rows = _table_rows(_sheet(wb, "QA Checks"), _table(_sheet(wb, "QA Checks"), "tblQA"))
    assert sum(1 for row in qa_rows if _status_ok(_cell_by_header(row, "Status").value, "Pass")) >= 6
    private_rows = _table_rows(_sheet(wb, "Private Notes"), _table(_sheet(wb, "Private Notes"), "tblPrivateNotes"))
    assert len(private_rows) == 6
    private_text = _sheet_text(wb, ["Private Notes"])
    for phrase in [row[0] for row in EXPECT["private_rows"][1:]]:
        assert normalize_text(phrase) in normalize_text(private_text)
    public = _sheet_text(wb, ["Package Summary", "Cue Timing QC", "Chapter Markers", "Publish Manifest", "QA Checks"])
    _forbid_private_and_raw_labels(public, "public workbook sheets")


def test_cross_output_consistency():
    doc_text = _doc_text(Document(_path("docx")))
    pdf_text = _pdf_text(_path("pdf"))
    wb = load_workbook(_path("xlsx"), data_only=False)
    _assert_required_sheets(wb, ["Package Summary", "Cue Timing QC", "Chapter Markers", "Publish Manifest"])
    for anchor in ["Maya Torres", "Tidepool Station", "Blue Harbor Library", "00:02:41"]:
        assert normalize_text(anchor) in normalize_text(doc_text)
        assert normalize_text(anchor) in normalize_text(pdf_text)
        assert normalize_text(anchor) in normalize_text(_sheet_text(wb, ["Cue Timing QC", "Chapter Markers", "Publish Manifest"]))
    _ready_formula_ok(_summary_value_cell(wb, "Package ready").value)
