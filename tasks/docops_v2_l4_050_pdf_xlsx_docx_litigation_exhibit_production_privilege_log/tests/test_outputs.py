import json
import os
import re
import sys
import zipfile
from pathlib import Path
from datetime import date, datetime

from docx import Document
from openpyxl import load_workbook
from openpyxl.utils.cell import get_column_letter, range_boundaries
from pypdf import PdfReader

sys.path.insert(0, str(Path(__file__).parent))
from verifier_utils import *  # noqa: F401,F403

META_PATH = Path(os.environ.get('TASK_METADATA_PATH', '/tests/task_metadata.json'))
META = json.loads(META_PATH.read_text(encoding='utf-8'))
EXPECT = META['verifier_expectations']


def _vopt_norm_text(value):
    if value is None:
        return ''
    if isinstance(value, datetime):
        if value.hour or value.minute or value.second:
            return value.strftime('%Y-%m-%d %H:%M')
        return value.strftime('%Y-%m-%d')
    if isinstance(value, date):
        return value.strftime('%Y-%m-%d')
    if isinstance(value, float) and value.is_integer():
        value = int(value)
    text = str(value)
    text = text.replace('\u2013', '-').replace('\u2014', '-').replace('\u2212', '-')
    text = text.replace('\xa0', ' ')
    text = re.sub(r'\s+', ' ', text).strip()
    return text


def _vopt_norm_formula(value):
    return re.sub(r'\s+', '', _vopt_norm_text(value)).upper()


def _vopt_norm_key(value):
    return re.sub(r'[^A-Z0-9]+', '', _vopt_norm_text(value).upper())


def _vopt_assert_formula_tokens(actual, tokens, label='formula'):
    norm = _vopt_norm_formula(actual)
    missing = [_vopt_norm_formula(token) for token in tokens if _vopt_norm_formula(token) not in norm]
    assert not missing, f"{label}: formula {actual!r} is missing purpose tokens {missing!r}"


def _vopt_norm_rows(rows):
    return [[_vopt_norm_text(cell) for cell in row] for row in rows]


def _vopt_rows_equal(actual, expected):
    return _vopt_norm_rows(actual) == _vopt_norm_rows(expected)


def _vopt_rows_text(rows):
    return '\n'.join('|'.join(_vopt_norm_text(cell) for cell in row) for row in rows)


def _vopt_header_has(header, required):
    actual = {_vopt_norm_text(cell).lower() for cell in header}
    missing = [cell for cell in required if _vopt_norm_text(cell).lower() not in actual]
    assert not missing, f"Missing expected header columns {missing!r}; actual={header!r}"


def _vopt_section_text(section_part):
    return '\n'.join(p.text for p in getattr(section_part, 'paragraphs', []) if p.text is not None)


def _vopt_header_text(doc):
    return '\n'.join(_vopt_section_text(section.header) for section in doc.sections)


def _vopt_footer_text(doc):
    return '\n'.join(_vopt_section_text(section.footer) for section in doc.sections)


def _vopt_table_ref(ws, expected_name=None):
    tables = ws.tables
    if expected_name:
        for name in tables.keys():
            if str(name).lower() == str(expected_name).lower():
                return tables[name].ref
    vals = list(tables.values())
    assert vals, f"No Excel native table found on sheet {ws.title!r}"
    return vals[0].ref


def _vopt_table(ws, expected_name):
    for name in ws.tables.keys():
        if str(name).lower() == str(expected_name).lower():
            return ws.tables[name]
    raise AssertionError(f"No Excel native table named {expected_name!r} found on sheet {ws.title!r}")


def _vopt_table_has_size(ws, expected_name, min_rows, min_cols):
    table = _vopt_table(ws, expected_name)
    min_col, min_row, max_col, max_row = range_boundaries(table.ref)
    actual_rows = max_row - min_row + 1
    actual_cols = max_col - min_col + 1
    assert actual_rows >= min_rows and actual_cols >= min_cols, (
        f"{ws.title}.{expected_name}: expected at least {min_rows} rows x {min_cols} cols, "
        f"found {actual_rows} rows x {actual_cols} cols at {table.ref}"
    )
    return table


def _vopt_table_rows(ws, table_name):
    table = _vopt_table(ws, table_name)
    min_col, min_row, max_col, max_row = range_boundaries(table.ref)
    return [
        [ws.cell(row, col).value for col in range(min_col, max_col + 1)]
        for row in range(min_row, max_row + 1)
    ]


def _vopt_find_header_index(header, accepted_names):
    accepted = {_vopt_norm_key(name) for name in accepted_names}
    for idx, value in enumerate(header):
        if _vopt_norm_key(value) in accepted:
            return idx
    raise AssertionError(f"Missing expected column {accepted_names!r}; header={header!r}")


def _vopt_clean_area(value):
    text = _vopt_norm_text(value).replace("'", '').replace('$', '').replace(' ', '')
    if '!' in text:
        text = text.split('!', 1)[1]
    return text.upper()


def _vopt_assert_print_area(ws, expected):
    actual = _vopt_clean_area(ws.print_area)
    target = _vopt_clean_area(expected)
    if target in actual or actual in target:
        return
    try:
        a_min_col, a_min_row, a_max_col, a_max_row = range_boundaries(actual)
        t_min_col, t_min_row, t_max_col, t_max_row = range_boundaries(target)
    except ValueError:
        raise AssertionError(f"{ws.title}: expected print area {expected!r}, found {ws.print_area!r}")
    encloses = (
        a_min_col <= t_min_col
        and a_min_row <= t_min_row
        and a_max_col >= t_max_col
        and a_max_row >= t_max_row
    )
    assert encloses, f"{ws.title}: expected print area {expected!r}, found {ws.print_area!r}"


def _vopt_freeze_pane(value):
    return getattr(value, 'coordinate', value)


def _vopt_validated_cells(ws):
    cells = set()
    for dv in ws.data_validations.dataValidation:
        for rng in dv.cells.ranges:
            text = str(rng).replace('$', '')
            try:
                min_col, min_row, max_col, max_row = range_boundaries(text)
            except ValueError:
                cells.add(text)
                continue
            if max_row > 10000:
                max_row = max(min_row, 200)
            for row in range(min_row, max_row + 1):
                for col in range(min_col, max_col + 1):
                    cells.add(f"{get_column_letter(col)}{row}")
    return cells


class _VoptValidationRanges(list):
    def __init__(self, ws, ranges):
        super().__init__(ranges)
        self.ws = ws

    def __contains__(self, expected):
        expected = str(expected).replace('$', '')
        if super().__contains__(expected):
            return True
        try:
            min_col, min_row, max_col, max_row = range_boundaries(expected)
        except ValueError:
            return super().__contains__(expected)
        target = {f"{get_column_letter(col)}{row}" for row in range(min_row, max_row + 1) for col in range(min_col, max_col + 1)}
        return target.issubset(_vopt_validated_cells(self.ws))


def _vopt_dv_ranges(ws):
    ranges = []
    for dv in ws.data_validations.dataValidation:
        ranges.extend(str(rng).replace('$', '') for rng in dv.cells.ranges)
    return _VoptValidationRanges(ws, ranges)



def _path(kind):
    env = {'pdf': 'PDF_OUTPUT_PATH', 'xlsx': 'XLSX_OUTPUT_PATH', 'docx': 'DOCX_OUTPUT_PATH'}[kind]
    key = {'pdf': 'pdf_output', 'xlsx': 'xlsx_output', 'docx': 'docx_output'}[kind]
    return Path(os.environ.get(env, EXPECT[key]))


def _pdf_text(path):
    return '\n'.join(page.extract_text() or '' for page in PdfReader(str(path)).pages)


def _outline_titles(path):
    return [title for _, title in flatten_outline(PdfReader(str(path)).outline)]


def _outline_matches(actual, expected):
    return [_vopt_norm_key(title) for title in actual] == [_vopt_norm_key(title) for title in expected]


def _doc_text(doc):
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append('|'.join(cell.text.strip() for cell in row.cells))
    return '\n'.join(parts)


def _doc_heading_order(doc):
    out = []
    for p in doc.paragraphs:
        style = p.style.name if p.style else ''
        if p.text.strip() and (style.startswith('Heading') or style == 'Title'):
            out.append(p.text.strip())
    return out


def _norm_heading_label(value):
    text = str(value or '')
    text = text.replace('\u2013', '-').replace('\u2014', '-').replace('\xa0', ' ')
    text = re.sub(r'\s+', ' ', text).strip()
    text = re.sub(r'^(?:\d+|[IVXLCDM]+)[\.)]\s+', '', text, flags=re.I)
    text = re.sub(r'^[A-Z][\.)]\s+', '', text)
    text = re.sub(r'[:\-]+$', '', text).strip()
    return text.casefold()


def _heading_sequence_matches(actual, expected):
    actual_norm = [_norm_heading_label(item) for item in actual]
    expected_norm = [_norm_heading_label(item) for item in expected]
    pos = 0
    for item in actual_norm:
        if pos < len(expected_norm) and item == expected_norm[pos]:
            pos += 1
    return pos == len(expected_norm)


def _doc_table_rows(table):
    return [[cell.text.strip() for cell in row.cells] for row in table.rows]


def _has_toc_field(path):
    with zipfile.ZipFile(path) as zf:
        xml = zf.read('word/document.xml').decode('utf-8', errors='ignore')
    return bool(re.search(r'TOC\s+(?:\\)?o|TOC\s*(?:&quot;|")', xml))


def _highlighted_para(doc, text):
    for p in doc.paragraphs:
        if _vopt_norm_text(text) in _vopt_norm_text(p.text):
            return docx_para_has_highlight(p)
    raise AssertionError(f'Missing paragraph: {text}')


def _highlighted_confidential_qc_para(doc):
    for p in doc.paragraphs:
        norm = _vopt_norm_text(p.text)
        if 'CONFIDENTIAL' in norm and 'PROD-003' in norm and 'PROD-004' in norm:
            assert docx_para_has_highlight(p), f"QC paragraph should be highlighted: {p.text!r}"
            return
    raise AssertionError('Missing highlighted QC paragraph for CONFIDENTIAL stamp confirmation')


def _find_doc_table_with_header(doc, required):
    for table in doc.tables:
        rows = _doc_table_rows(table)
        if rows:
            header = {_vopt_norm_key(cell) for cell in rows[0]}
            if all(any(_vopt_norm_key(req) in item or item in _vopt_norm_key(req) for item in header) for req in required):
                return rows
    raise AssertionError(f'Missing DOCX table with columns {required!r}')


def _require_docx_memo_semantics(text):
    require_all(text, ['produced_exhibit_binder.pdf', 'ACME-000001', 'ACME-000007'], 'production memo')
    for produced_id in EXPECT['produced_ids']:
        assert produced_id in text, f"production memo missing produced ID {produced_id}"
    for withheld_id in EXPECT['withheld_ids']:
        assert withheld_id in text, f"production memo missing withheld ID {withheld_id}"
    for excluded_id in EXPECT['excluded_ids']:
        assert excluded_id in text, f"production memo missing excluded ID {excluded_id}"
    require_all(text, ['privilege log', 'withheld', 'excluded'], 'production memo')


def _require_pdf_binder_semantics(text):
    for produced_id in EXPECT['produced_ids']:
        assert produced_id in text, f"Missing produced doc {produced_id}"
    for title in ['Maintenance Email Thread', 'Vendor Invoice', 'Incident Photo Contact Sheet', 'Safety Audit Extract']:
        assert title in text, f"Missing produced document title {title!r}"
    for label in EXPECT['bates_labels']:
        assert label in text, f"Missing Bates label {label}"


def _row_values(ws, row, max_col):
    return [ws.cell(row, col).value for col in range(1, max_col + 1)]


def _dv_ranges(ws):
    return _vopt_dv_ranges(ws)


def _sheet_formula_text(ws):
    return '\n'.join(
        _vopt_norm_formula(cell.value)
        for row in ws.iter_rows()
        for cell in row
        if isinstance(cell.value, str) and cell.value.startswith('=')
    )


def _assert_qc_summary_formulas(qc):
    formulas = _sheet_formula_text(qc)
    for token in ['PRODUCTIONINDEX', 'PRIVILEGELOG', 'EXCLUDEDITEMS']:
        assert token in formulas, f'QC Summary formulas must reference {token}'
    assert any(token in formulas for token in ['SUM', 'SUMIF']), 'QC Summary must formula-calculate produced pages'


def _assert_hidden_attorney_notes(notes_text):
    lower = _vopt_norm_text(notes_text).lower()
    assert ('privilege' in lower or 'withheld' in lower or 'priv-005' in lower), 'Attorney Notes missing privilege/withheld work record'
    assert 'confidential' in lower, 'Attorney Notes missing confidentiality work record'
    assert ('excluded' in lower or 'exclusion' in lower or 'nonresp-008' in lower), 'Attorney Notes missing excluded-item work record'


def test_outputs_exist():
    assert _path('pdf').exists(), f"Missing PDF output: {_path('pdf')}"
    assert _path('xlsx').exists(), f"Missing XLSX output: {_path('xlsx')}"
    assert _path('docx').exists(), f"Missing DOCX output: {_path('docx')}"
    assert _path('pdf').suffix.lower() == '.pdf'
    assert _path('xlsx').suffix.lower() == '.xlsx'
    assert _path('docx').suffix.lower() == '.docx'


def test_pdf_binder_scope_bates_bookmarks_and_confidentiality():
    pdf = _path('pdf')
    reader = PdfReader(str(pdf))
    assert len(reader.pages) == 7, f"Expected 7 produced PDF pages, found {len(reader.pages)}"
    assert _outline_matches(_outline_titles(pdf), EXPECT['pdf_bookmarks']), f"Unexpected PDF bookmarks: {_outline_titles(pdf)!r}"
    text = _pdf_text(pdf)
    _require_pdf_binder_semantics(text)
    forbid_any(text, EXPECT['pdf_forbidden_text'], 'produced PDF binder')
    assert text.count('CONFIDENTIAL') >= 6, 'Expected repeated confidentiality stamps on confidential document pages'


def test_workbook_native_controls_and_privilege_log_content():
    wb = load_workbook(_path('xlsx'), data_only=False)
    assert wb.sheetnames == EXPECT['workbook_sheet_order'], f"Unexpected sheets: {wb.sheetnames!r}"
    for sheet_name, (table_name, table_ref) in EXPECT['tables'].items():
        min_col, min_row, max_col, max_row = range_boundaries(table_ref)
        _vopt_table_has_size(
            wb[sheet_name],
            table_name,
            max_row - min_row + 1,
            max_col - min_col + 1,
        )
    prod = wb['Production Index']
    priv = wb['Privilege Log']
    excluded = wb['Excluded Items']
    qc = wb['QC Summary']
    assert prod.freeze_panes == 'A2'
    assert priv.freeze_panes == 'A2'
    _vopt_assert_print_area(prod, EXPECT['print_areas']['Production Index'])
    _vopt_assert_print_area(priv, EXPECT['print_areas']['Privilege Log'])
    _vopt_assert_print_area(qc, EXPECT['print_areas']['QC Summary'])
    for sheet in EXPECT['hidden_sheets']:
        assert wb[sheet].sheet_state in ('hidden', 'veryHidden'), f"{sheet} should be hidden"
    names = {dn.name for dn in wb.defined_names.values()}
    for name in EXPECT['defined_names']:
        assert name in names, f"Missing defined name: {name}"
    formula_expectations = {
        'Production Index!I2': ['LEFT', 'A2', 'PROD', 'PRODUCED'],
        'Production Index!I5': ['LEFT', 'A5', 'PROD', 'PRODUCED'],
    }
    for ref, tokens in formula_expectations.items():
        sheet, cell = ref.split('!')
        _vopt_assert_formula_tokens(wb[sheet][cell].value, tokens, ref)
    _assert_qc_summary_formulas(qc)
    prod_rows = _vopt_table_rows(prod, 'production_index')
    prod_header, prod_data = prod_rows[0], prod_rows[1:]
    prod_id_col = _vopt_find_header_index(prod_header, ['ID', 'Doc ID', 'Document ID', 'Candidate ID'])
    pages_col = _vopt_find_header_index(prod_header, ['Pages', 'Page Count'])
    conf_col = _vopt_find_header_index(prod_header, ['Confidentiality'])
    produced_ids = [_vopt_norm_text(row[prod_id_col]) for row in prod_data]
    assert produced_ids == EXPECT['produced_ids']
    assert sum(int(row[pages_col]) for row in prod_data) == 7
    conf_by_id = {row[prod_id_col]: _vopt_norm_key(row[conf_col]) for row in prod_data}
    for produced_id in EXPECT['confidential_pdf_ids']:
        assert 'CONFIDENTIAL' in conf_by_id[produced_id], f"{produced_id} should be marked CONFIDENTIAL"
    priv_rows = _vopt_table_rows(priv, 'privilege_log')
    priv_header, priv_data = priv_rows[0], priv_rows[1:]
    priv_id_col = _vopt_find_header_index(priv_header, ['ID', 'Doc ID', 'Priv ID', 'Candidate ID'])
    status_col = _vopt_find_header_index(priv_header, ['Status'])
    assert [_vopt_norm_text(row[priv_id_col]) for row in priv_data] == EXPECT['withheld_ids']
    for row in priv_data:
        assert 'WITHHELD' in _vopt_norm_key(row[status_col]), f"Privilege row should be withheld: {row!r}"
    priv_text = _vopt_rows_text(priv_data)
    require_all(priv_text, ['Attorney-client', 'Work product', 'Maya Ortiz'], 'privilege log')
    forbid_any(priv_text, ['Legal strategy:', 'settlement posture'], 'privilege log')
    excluded_rows = _vopt_table_rows(excluded, 'excluded_items')
    excluded_header, excluded_data = excluded_rows[0], excluded_rows[1:]
    excl_id_col = _vopt_find_header_index(excluded_header, ['ID', 'Doc ID', 'Item ID', 'Candidate ID'])
    assert [_vopt_norm_text(row[excl_id_col]) for row in excluded_data] == EXPECT['excluded_ids']
    excluded_text = _vopt_rows_text(excluded_data)
    require_all(excluded_text, ['Duplicate', 'Non-responsive'], 'excluded items')
    notes_text = '\n'.join(str(cell.value) for row in wb['Attorney Notes'].iter_rows() for cell in row if cell.value)
    _assert_hidden_attorney_notes(notes_text)
    assert len(prod.conditional_formatting) >= 1
    prod_validation = _dv_ranges(prod)
    assert 'E2:E5' in prod_validation, f"Missing confidentiality validation on Production Index: {prod_validation!r}"
    assert ('I2:I5' in prod_validation) or ('K2:K5' in prod_validation), (
        f"Missing production/QC status validation on Production Index: {prod_validation!r}"
    )
    priv_validation = _dv_ranges(priv)
    assert 'I2:I4' in priv_validation, f"Missing privilege status validation on Privilege Log: {priv_validation!r}"


def test_docx_production_memo_structure_and_public_boundary():
    docx = _path('docx')
    doc = Document(docx)
    text = _doc_text(doc)
    headings = _doc_heading_order(doc)
    if EXPECT['doc_heading_order'][0] not in headings and EXPECT['doc_heading_order'][0] in text:
        headings = [EXPECT['doc_heading_order'][0]] + headings
    assert _heading_sequence_matches(headings, EXPECT['doc_heading_order']), f"Unexpected heading order: {_doc_heading_order(doc)!r}"
    _require_docx_memo_semantics(text)
    forbid_any(text, EXPECT['doc_forbidden_text'], 'production memo')
    assert EXPECT['doc_header_contains'] in _vopt_header_text(doc)
    assert EXPECT['doc_footer_contains'] in _vopt_footer_text(doc)
    assert _has_toc_field(docx), 'Expected real Word TOC field'
    _highlighted_confidential_qc_para(doc)
    rows = _find_doc_table_with_header(doc, ['Doc ID', 'Title', 'Bates'])
    _vopt_header_has(rows[0], ['Doc ID', 'Title'])
    assert any('BATES' in _vopt_norm_key(cell) for cell in rows[0]), f"Missing Bates column in DOCX exhibit table: {rows[0]!r}"
    assert [row[0] for row in rows[1:]] == EXPECT['produced_ids']


def test_cross_output_production_and_privilege_consistency():
    pdf_text = _pdf_text(_path('pdf'))
    doc_text = _doc_text(Document(_path('docx')))
    wb = load_workbook(_path('xlsx'), data_only=False)
    prod = wb['Production Index']
    priv = wb['Privilege Log']
    prod_rows = _vopt_table_rows(prod, 'production_index')
    priv_rows = _vopt_table_rows(priv, 'privilege_log')
    prod_id_col = _vopt_find_header_index(prod_rows[0], ['ID', 'Doc ID', 'Document ID', 'Candidate ID'])
    priv_id_col = _vopt_find_header_index(priv_rows[0], ['ID', 'Doc ID', 'Priv ID', 'Candidate ID'])
    produced_ids = [_vopt_norm_text(row[prod_id_col]) for row in prod_rows[1:]]
    withheld_ids = [_vopt_norm_text(row[priv_id_col]) for row in priv_rows[1:]]
    assert produced_ids == EXPECT['produced_ids']
    assert withheld_ids == EXPECT['withheld_ids']
    for produced_id in produced_ids:
        assert produced_id in pdf_text
        assert produced_id in doc_text
    for withheld_id in withheld_ids:
        assert withheld_id not in pdf_text
        assert withheld_id in doc_text
    for excluded_id in EXPECT['excluded_ids']:
        assert excluded_id not in pdf_text
        assert excluded_id in doc_text
    for forbidden in ['settlement posture', 'reserve exposure', 'counsel theory']:
        assert forbidden not in normalize_text(pdf_text)
