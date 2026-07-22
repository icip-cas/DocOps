import json
import os
import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt

from verifier_utils import normalize_text, run_preflight

META_PATH = Path(os.environ.get("TASK_METADATA_PATH", "/tests/task_metadata.json"))
if not META_PATH.exists():
    META_PATH = Path(__file__).parent / "task_metadata.json"
META = json.loads(META_PATH.read_text(encoding="utf-8"))
EXPECT = META["verifier_expectations"]
INPUT_PATH = Path(os.environ.get("INPUT_PATH", META["input_path"]))
OUTPUT_PATH = Path(os.environ.get("OUTPUT_PATH", META["output_path"]))


def doc_text(doc):
    parts = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text.strip())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())
    for section in doc.sections:
        for paragraph in section.footer.paragraphs:
            if paragraph.text.strip():
                parts.append(paragraph.text.strip())
    core = doc.core_properties
    parts.extend(str(value or "") for value in [core.author, core.title, core.subject, core.keywords, core.comments])
    return "\n".join(parts)


def visible_heading_sequence(doc):
    wanted = set(EXPECT["headings"])
    return [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip() in wanted]


def paragraph_by_text(doc, text):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise AssertionError(f"Missing paragraph: {text}")


def first_run(paragraph):
    for run in paragraph.runs:
        if run.text.strip():
            return run
    return paragraph.runs[0] if paragraph.runs else None


def effective_font_name(paragraph, run):
    return (run.font.name if run is not None else None) or paragraph.style.font.name


def effective_font_size(paragraph, run):
    return (run.font.size if run is not None else None) or paragraph.style.font.size


def effective_bold(paragraph, run):
    value = run.font.bold if run is not None else None
    return value if value is not None else paragraph.style.font.bold


def effective_color(paragraph, run):
    rgb = None
    if run is not None and run.font.color is not None:
        rgb = run.font.color.rgb
    if rgb is None and paragraph.style.font.color is not None:
        rgb = paragraph.style.font.color.rgb
    return None if rgb is None else str(rgb)


def all_table_rows(doc):
    rows = []
    for table in doc.tables:
        table_rows = []
        for row in table.rows:
            table_rows.append([cell.text.strip() for cell in row.cells])
        rows.append(table_rows)
    return rows


def cell_fill(cell):
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    shd = cell._tc.find(".//w:shd", namespaces=ns)
    if shd is None:
        return None
    return shd.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill")


def has_page_break(paragraph):
    xml = paragraph._p.xml
    return 'w:type="page"' in xml or "w:type='page'" in xml


def cleanup_context(text, phrase):
    pattern = re.compile(rf"\b{re.escape(phrase)}\b", re.IGNORECASE)
    for match in pattern.finditer(text):
        window = text[max(0, match.start() - 120): match.end() + 180]
        if not re.search(r"\b(no|not|removed|cleaned|cleared|excluded|without|corrected|replaced)\b", window, re.IGNORECASE):
            return False
    return True


def assert_no_forbidden_text(text):
    normalized = normalize_text(text)
    hard = [
        "Lake Side School Bord", "2026-04-12", "4 members present", "3-2",
        "$1,257,000", "Policy 6154", "8:24 PM", "2026-05-16",
        "approvl", "attendence", "minits", "resoultion", "budgte", "calender",
    ]
    hits = [phrase for phrase in hard if normalize_text(phrase) in normalized]
    assert not hits, f"Forbidden error remnants still present: {hits}"
    contextual = [phrase for phrase in ["DRAFT", "PRIVATE", "TODO", "obsolete"] if normalize_text(phrase) in normalized and not cleanup_context(text, phrase)]
    assert not contextual, f"Forbidden cleanup markers still present: {contextual}"


def assert_table_header_style(doc):
    style = EXPECT["style"]
    for table in doc.tables:
        for cell in table.rows[0].cells:
            assert cell_fill(cell) == style["table_header_fill"]
            runs = [run for paragraph in cell.paragraphs for run in paragraph.runs if run.text.strip()]
            assert runs, "Header cell has no text run"
            assert all(run.font.bold for run in runs)
            assert all(str(run.font.color.rgb) == style["table_header_font"] for run in runs)


def test_output_exists_and_is_docx():
    run_preflight(META, INPUT_PATH, OUTPUT_PATH)
    assert OUTPUT_PATH != INPUT_PATH


def test_headings_required_text_tables_and_forbidden_errors_removed():
    doc = Document(OUTPUT_PATH)
    assert visible_heading_sequence(doc) == EXPECT["headings"]
    text = doc_text(doc)
    for phrase in [
        "Lakeside School Board", "2026-04-21", "5 members present", "4-1",
        "$1,275,000", "Policy 6145", "8:42 PM", "2026-05-06",
    ]:
        assert phrase in text, f"Missing corrected public value: {phrase}"
    for phrase in ["correction register", "public", "release"]:
        assert normalize_text(phrase) in normalize_text(text), f"Missing public repair context: {phrase}"
    assert_no_forbidden_text(text)
    assert len(doc.tables) >= len(EXPECT["tables"]), "Missing public native tables"


def test_style_layout_footer_and_page_breaks():
    doc = Document(OUTPUT_PATH)
    style = EXPECT["style"]
    title = paragraph_by_text(doc, EXPECT["headings"][0])
    title_run = first_run(title)
    assert effective_font_name(title, title_run) == style["title_font"]
    assert effective_bold(title, title_run)
    assert effective_font_size(title, title_run) == Pt(style["title_size_pt"])
    assert effective_color(title, title_run) == style["title_color"]
    for heading in EXPECT["headings"][1:]:
        paragraph = paragraph_by_text(doc, heading)
        run = first_run(paragraph)
        assert effective_color(paragraph, run) == style["heading_color"]
        assert effective_bold(paragraph, run)
    for section in doc.sections:
        assert abs(section.top_margin.inches - style["margin_inches"]) < 0.02
        assert abs(section.bottom_margin.inches - style["margin_inches"]) < 0.02
        assert abs(section.left_margin.inches - style["margin_inches"]) < 0.02
        assert abs(section.right_margin.inches - style["margin_inches"]) < 0.02
        footer_text = "\n".join(p.text for p in section.footer.paragraphs)
        assert style["footer_text"] in footer_text
    assert_table_header_style(doc)


def test_public_metadata_applied():
    doc = Document(OUTPUT_PATH)
    core = doc.core_properties
    assert core.author == EXPECT["metadata"]["author"]
    assert core.title == EXPECT["metadata"]["title"]
    meta_text = "\n".join(str(value or "") for value in [core.author, core.title, core.subject, core.keywords, core.comments])
    assert_no_forbidden_text(meta_text)


def test_source_artifact_was_not_modified():
    source = Document(INPUT_PATH)
    text = doc_text(source)
    missing = [phrase for phrase in EXPECT["source_must_contain"] if normalize_text(phrase) not in normalize_text(text)]
    assert not missing, f"Source artifact no longer contains expected errors: {missing}"
