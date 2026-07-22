import json
import os
import re
import sys
import zipfile
from pathlib import Path
from datetime import date, datetime

from docx import Document
from openpyxl import load_workbook
from openpyxl.utils.cell import get_column_letter, range_boundaries

sys.path.insert(0, str(Path(__file__).parent))
from verifier_utils import *  # noqa: F401,F403

META_PATH = Path(os.environ.get('TASK_METADATA_PATH', '/tests/task_metadata.json'))
META = json.loads(META_PATH.read_text(encoding='utf-8'))
EXPECT = META['verifier_expectations']


def _vopt_norm_text(value):
    if value is None:
        return ''
    if isinstance(value, datetime):
        if value.hour or value.minute or value.second:
            return value.strftime('%Y-%m-%d %H:%M')
        return value.strftime('%Y-%m-%d')
    if isinstance(value, date):
        return value.strftime('%Y-%m-%d')
    if isinstance(value, float) and value.is_integer():
        value = int(value)
    text = str(value)
    text = text.replace('\u2013', '-').replace('\u2014', '-').replace('\u2212', '-')
    text = text.replace('\xa0', ' ')
    text = re.sub(r'\s+', ' ', text).strip()
    return text


def _vopt_norm_formula(value):
    return re.sub(r'\s+', '', _vopt_norm_text(value)).upper()


def _vopt_norm_key(value):
    return re.sub(r'[^A-Z0-9]+', '', _vopt_norm_text(value).upper())


def _vopt_assert_formula_tokens(actual, tokens, label='formula'):
    norm = _vopt_norm_formula(actual)
    missing = [_vopt_norm_formula(token) for token in tokens if _vopt_norm_formula(token) not in norm]
    assert not missing, f"{label}: formula {actual!r} is missing purpose tokens {missing!r}"


def _vopt_sheet_has_formula_tokens(ws, tokens, label):
    for row in ws.iter_rows():
        for cell in row:
            value = cell.value
            if isinstance(value, str) and value.startswith('='):
                norm = _vopt_norm_formula(value)
                if all(_vopt_norm_formula(token) in norm for token in tokens):
                    return
    raise AssertionError(f"{label}: no formula found with purpose tokens {tokens!r}")


def _vopt_norm_rows(rows):
    return [[_vopt_norm_text(cell) for cell in row] for row in rows]


def _vopt_rows_equal(actual, expected):
    return _vopt_norm_rows(actual) == _vopt_norm_rows(expected)


def _vopt_rows_text(rows):
    return '\n'.join('|'.join(_vopt_norm_text(cell) for cell in row) for row in rows)


def _vopt_header_has(header, required):
    actual = {_vopt_norm_text(cell).lower() for cell in header}
    missing = [cell for cell in required if _vopt_norm_text(cell).lower() not in actual]
    assert not missing, f"Missing expected header columns {missing!r}; actual={header!r}"


def _vopt_section_text(section_part):
    return '\n'.join(p.text for p in getattr(section_part, 'paragraphs', []) if p.text is not None)


def _vopt_header_text(doc):
    return '\n'.join(_vopt_section_text(section.header) for section in doc.sections)


def _vopt_footer_text(doc):
    return '\n'.join(_vopt_section_text(section.footer) for section in doc.sections)


def _vopt_contains_text(haystack, needle):
    return _vopt_norm_text(needle).casefold() in _vopt_norm_text(haystack).casefold()


def _vopt_table_ref(ws, expected_name=None):
    tables = ws.tables
    if expected_name:
        for name in tables.keys():
            if str(name).lower() == str(expected_name).lower():
                return tables[name].ref
    vals = list(tables.values())
    assert vals, f"No Excel native table found on sheet {ws.title!r}"
    return vals[0].ref


def _vopt_table(ws, expected_name):
    for name in ws.tables.keys():
        if str(name).lower() == str(expected_name).lower():
            return ws.tables[name]
    raise AssertionError(f"No Excel native table named {expected_name!r} found on sheet {ws.title!r}")


def _vopt_table_has_size(ws, expected_name, min_rows, min_cols):
    table = _vopt_table(ws, expected_name)
    min_col, min_row, max_col, max_row = range_boundaries(table.ref)
    actual_rows = max_row - min_row + 1
    actual_cols = max_col - min_col + 1
    assert actual_rows >= min_rows and actual_cols >= min_cols, (
        f"{ws.title}.{expected_name}: expected at least {min_rows} rows x {min_cols} cols, "
        f"found {actual_rows} rows x {actual_cols} cols at {table.ref}"
    )
    return table


def _vopt_table_rows(ws, table_name):
    table = _vopt_table(ws, table_name)
    min_col, min_row, max_col, max_row = range_boundaries(table.ref)
    return [
        [ws.cell(row, col).value for col in range(min_col, max_col + 1)]
        for row in range(min_row, max_row + 1)
    ]


def _vopt_find_header_index(header, accepted_names):
    accepted = {_vopt_norm_key(name) for name in accepted_names}
    for idx, value in enumerate(header):
        norm = _vopt_norm_key(value)
        if norm in accepted:
            return idx
        for item in accepted:
            if item == 'VENDOR' and norm != 'VENDOR':
                continue
            if len(item) > 2 and norm.startswith(item):
                return idx
    raise AssertionError(f"Missing expected column {accepted_names!r}; header={header!r}")


def _vopt_try_header_index(header, accepted_names):
    try:
        return _vopt_find_header_index(header, accepted_names)
    except AssertionError:
        return None


def _vopt_clean_area(value):
    text = _vopt_norm_text(value).replace("'", '').replace('$', '').replace(' ', '')
    if '!' in text:
        text = text.split('!', 1)[1]
    return text.upper()


def _vopt_assert_print_area(ws, expected):
    actual = _vopt_clean_area(ws.print_area)
    target = _vopt_clean_area(expected)
    assert target in actual or actual in target, f"{ws.title}: expected print area {expected!r}, found {ws.print_area!r}"


def _vopt_bounds(area):
    return range_boundaries(_vopt_clean_area(area))


def _vopt_assert_print_area_covers_table(ws, table_name, expected):
    try:
        _vopt_assert_print_area(ws, expected)
        return
    except AssertionError:
        pass
    actual = _vopt_clean_area(ws.print_area)
    assert actual, f"{ws.title}: missing print area"
    a_min_col, a_min_row, a_max_col, a_max_row = _vopt_bounds(actual)
    t_min_col, t_min_row, t_max_col, t_max_row = _vopt_bounds(_vopt_table(ws, table_name).ref)
    assert (
        a_min_col <= t_min_col
        and a_min_row <= t_min_row
        and a_max_col >= t_max_col
        and a_max_row >= t_max_row
    ), f"{ws.title}: print area {ws.print_area!r} does not cover table {table_name} at {_vopt_table(ws, table_name).ref}"


def _vopt_freeze_pane(value):
    return getattr(value, 'coordinate', value)


def _vopt_validated_cells(ws):
    cells = set()
    for dv in ws.data_validations.dataValidation:
        for rng in dv.cells.ranges:
            text = str(rng).replace('$', '')
            try:
                min_col, min_row, max_col, max_row = range_boundaries(text)
            except ValueError:
                cells.add(text)
                continue
            if max_row > 10000:
                max_row = max(min_row, 200)
            for row in range(min_row, max_row + 1):
                for col in range(min_col, max_col + 1):
                    cells.add(f"{get_column_letter(col)}{row}")
    return cells


class _VoptValidationRanges(list):
    def __init__(self, ws, ranges):
        super().__init__(ranges)
        self.ws = ws

    def __contains__(self, expected):
        expected = str(expected).replace('$', '')
        if super().__contains__(expected):
            return True
        try:
            min_col, min_row, max_col, max_row = range_boundaries(expected)
        except ValueError:
            return super().__contains__(expected)
        target = {f"{get_column_letter(col)}{row}" for row in range(min_row, max_row + 1) for col in range(min_col, max_col + 1)}
        return target.issubset(_vopt_validated_cells(self.ws))


def _vopt_dv_ranges(ws):
    ranges = []
    for dv in ws.data_validations.dataValidation:
        ranges.extend(str(rng).replace('$', '') for rng in dv.cells.ranges)
    return _VoptValidationRanges(ws, ranges)



def _path(kind):
    env = {'docx': 'DOCX_OUTPUT_PATH', 'xlsx': 'XLSX_OUTPUT_PATH'}[kind]
    key = {'docx': 'docx_output', 'xlsx': 'xlsx_output'}[kind]
    return Path(os.environ.get(env, EXPECT[key]))


def _doc_text(doc):
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append('|'.join(cell.text.strip() for cell in row.cells))
    return '\n'.join(parts)


def _doc_heading_order(doc):
    out = []
    for p in doc.paragraphs:
        style = p.style.name if p.style else ''
        if p.text.strip() and (style.startswith('Heading') or style == 'Title'):
            out.append(p.text.strip())
    return out


def _norm_heading_label(value):
    text = str(value or '')
    text = text.replace('\u2013', '-').replace('\u2014', '-').replace('\xa0', ' ')
    text = re.sub(r'\s+', ' ', text).strip()
    text = re.sub(r'^(?:\d+|[IVXLCDM]+)[\.)]\s+', '', text, flags=re.I)
    text = re.sub(r'^[A-Z][\.)]\s+', '', text)
    text = re.sub(r'[:\-]+$', '', text).strip()
    return text.casefold()


def _heading_sequence_matches(actual, expected):
    actual_norm = [_norm_heading_label(item) for item in actual]
    expected_norm = [_norm_heading_label(item) for item in expected]
    pos = 0
    for item in actual_norm:
        if pos < len(expected_norm) and item == expected_norm[pos]:
            pos += 1
    return pos == len(expected_norm)


def _doc_table_rows(table):
    return [[cell.text.strip() for cell in row.cells] for row in table.rows]


def _has_toc_field(path):
    with zipfile.ZipFile(path) as zf:
        xml = zf.read('word/document.xml').decode('utf-8', errors='ignore')
    return bool(re.search(r'TOC\s+(?:\\)?o|TOC\s*(?:&quot;|")', xml))


def _highlighted_para(doc, text):
    for p in doc.paragraphs:
        if _vopt_norm_text(text) in _vopt_norm_text(p.text):
            return docx_para_has_highlight(p)
    raise AssertionError(f'Missing paragraph: {text}')


def _highlighted_conflict_control_para(doc):
    for p in doc.paragraphs:
        norm = _vopt_norm_text(p.text)
        if 'conflicted' in norm.lower() and 'CivicStack' in norm and 'excluded' in norm.lower() and 'final averaging' in norm.lower():
            assert docx_para_has_highlight(p), f"Conflict-control paragraph should be highlighted: {p.text!r}"
            return
    raise AssertionError('Missing highlighted conflict-control paragraph explaining the CivicStack score exclusion')


def _find_doc_table_with_header(doc, required):
    for table in doc.tables:
        rows = _doc_table_rows(table)
        if rows:
            header = {_vopt_norm_key(cell) for cell in rows[0]}
            if all(any(_vopt_norm_key(req) in item or item in _vopt_norm_key(req) for item in header) for req in required):
                return rows
    raise AssertionError(f'Missing DOCX table with columns {required!r}')


def _rank_rows_by_header(rows):
    header, data = rows[0], rows[1:]
    rank_col = _vopt_find_header_index(header, ['Rank'])
    vendor_id_col = _vopt_try_header_index(header, ['Vendor ID', 'ID'])
    vendor_col = _vopt_find_header_index(header, ['Vendor', 'Vendor Name'])
    score_col = _vopt_try_header_index(header, ['Score', 'Weighted Total', 'Total Weighted Score', 'Composite Score'])
    status_col = _vopt_find_header_index(header, ['Disposition', 'Status', 'Award Status', 'Recommendation', 'Recommendation/Status'])
    return [
        [
            _vopt_norm_text(row[rank_col]),
            _vopt_norm_text(row[vendor_id_col]) if vendor_id_col is not None else '',
            _vopt_norm_text(row[vendor_col]),
            _vopt_norm_text(row[score_col]) if score_col is not None else '',
            _vopt_norm_text(row[status_col]),
        ]
        for row in data[:4]
    ]


def _cell_ref_points_to_vendor_score(formula, wb, expected_vendor_id, expected_vendor):
    if not wb:
        return False
    match = re.fullmatch(r"=\s*'?([^'!]+)'?!\$?([A-Z]+)\$?(\d+)\s*", _vopt_norm_text(formula), flags=re.I)
    if not match:
        return False
    sheet_name, col, row_text = match.groups()
    if sheet_name not in wb.sheetnames:
        return False
    ws = wb[sheet_name]
    row = int(row_text)
    header = [ws.cell(1, c).value for c in range(1, ws.max_column + 1)]
    vendor_id_col = _vopt_try_header_index(header, ['Vendor ID', 'ID'])
    vendor_col = _vopt_try_header_index(header, ['Vendor', 'Vendor Name'])
    score_col = _vopt_try_header_index(header, ['Score', 'Weighted Total', 'Total Weighted Score', 'Composite Score'])
    target_col = range_boundaries(f'{col}{row}:{col}{row}')[0]
    if vendor_id_col is not None:
        actual_id = _vopt_norm_text(ws.cell(row, vendor_id_col + 1).value)
        if expected_vendor_id and actual_id != expected_vendor_id:
            return False
    elif vendor_col is not None and _vopt_norm_text(ws.cell(row, vendor_col + 1).value) != expected_vendor:
        return False
    return score_col is None or target_col == score_col + 1


def _score_matches(got_score, want_score, wb=None, expected_vendor_id='', expected_vendor='', require_score=True):
    text = _vopt_norm_text(got_score)
    if not require_score:
        return True
    if text.startswith('='):
        return _cell_ref_points_to_vendor_score(text, wb, expected_vendor_id, expected_vendor)
    cleaned = re.sub(r'[^0-9.\-]+', '', text)
    if cleaned:
        try:
            return abs(float(cleaned) - float(want_score)) <= 0.05
        except ValueError:
            pass
    return False


def _status_matches(actual, expected):
    norm = _vopt_norm_key(actual)
    expected_norm = _vopt_norm_key(expected)
    if 'RECOMMENDED' in expected_norm:
        return 'RECOMMENDED' in norm and 'AWARD' in norm
    if 'INELIGIBLE' in expected_norm:
        return 'INELIGIBLE' in norm or 'MANDATORY' in norm
    if 'ALTERNATE' in expected_norm:
        return 'ALTERNATE' in norm
    return expected_norm in norm or norm in expected_norm


def _assert_public_rank_rows(rows, label, wb=None, require_vendor_id=True, require_score=True):
    actual = _rank_rows_by_header(rows)
    assert len(actual) >= 4, f"{label}: expected at least 4 rank rows, found {len(actual)}"
    expected = EXPECT['public_rank_rows']
    for got, want in zip(actual[:4], expected):
        if require_vendor_id:
            assert got[1] == want[1], f"{label}: expected vendor ID {want[1]}, found {got[1]}"
        assert got[2] == want[2], f"{label}: expected vendor {want[2]}, found {got[2]}"
        if want[0] == 'Not ranked':
            rank_key = _vopt_norm_key(got[0])
            status_key = _vopt_norm_key(got[4])
            assert (
                rank_key in {'NA', 'NOTRANKED'}
                or 'NOTRANKED' in status_key
                or 'INELIGIBLE' in status_key
                or 'MANDATORY' in status_key
            ), f"{label}: Northstar should be not ranked/ineligible, found {got!r}"
        else:
            assert got[0] == want[0], f"{label}: expected rank {want[0]}, found {got[0]}"
            assert _score_matches(got[3], want[3], wb, want[1], want[2], require_score), f"{label}: expected score {want[3]}, found {got[3]}"
        assert _status_matches(got[4], want[4]), f"{label}: expected disposition like {want[4]!r}, found {got[4]!r}"


def _require_award_memo_semantics(text):
    require_all(text, ['CivicStack', 'Northstar Data', 'mandatory data-residency', 'best-value'], 'award memo')
    for factor in ['technical', 'implementation', 'accessibility', 'price', 'past performance']:
        assert factor in text.lower(), f"award memo: missing evaluation factor {factor!r}"
    require_all(text, ['Apex Archive', 'BrightLedger'], 'award memo')


def _row_values(ws, row, max_col):
    return [ws.cell(row, col).value for col in range(1, max_col + 1)]


def _dv_ranges(ws):
    return _vopt_dv_ranges(ws)


def _find_row_containing(ws, needle, min_row=1, max_row=None):
    target = _vopt_norm_text(needle).lower()
    max_row = max_row or ws.max_row
    for row in range(min_row, max_row + 1):
        row_text = ' '.join(_vopt_norm_text(ws.cell(row, col).value).lower() for col in range(1, ws.max_column + 1))
        if target in row_text:
            return row
    raise AssertionError(f"Could not find row containing {needle!r} on sheet {ws.title!r}")


def test_outputs_exist():
    assert _path('docx').exists(), f"Missing DOCX output: {_path('docx')}"
    assert _path('xlsx').exists(), f"Missing XLSX output: {_path('xlsx')}"
    assert _path('docx').suffix.lower() == '.docx'
    assert _path('xlsx').suffix.lower() == '.xlsx'


def test_award_memo_structure_boundary_and_ranking():
    docx = _path('docx')
    doc = Document(docx)
    assert _heading_sequence_matches(_doc_heading_order(doc), EXPECT['doc_heading_order']), f"Unexpected heading order: {_doc_heading_order(doc)!r}"
    text = _doc_text(doc)
    _require_award_memo_semantics(text)
    forbid_any(text, EXPECT['doc_forbidden_text'], 'award memo')
    assert _vopt_contains_text(_vopt_header_text(doc), EXPECT['doc_header_contains'])
    assert _vopt_contains_text(_vopt_footer_text(doc), EXPECT['doc_footer_contains'])
    assert _has_toc_field(docx)
    _highlighted_conflict_control_para(doc)
    rows = _find_doc_table_with_header(doc, ['Rank', 'Vendor'])
    _assert_public_rank_rows(rows, 'award memo rank table', require_vendor_id=False, require_score=False)


def test_scorebook_native_controls_formulas_and_public_debrief():
    wb = load_workbook(_path('xlsx'), data_only=False)
    assert wb.sheetnames == EXPECT['sheet_order']
    for sheet_name, (table_name, ref) in EXPECT['tables'].items():
        min_col, min_row, max_col, max_row = range_boundaries(ref)
        table = _vopt_table_has_size(wb[sheet_name], table_name, max_row - min_row + 1, max_col - min_col + 1 if sheet_name != 'Score Matrix' else 10)
        if sheet_name == 'Score Matrix':
            header = _vopt_table_rows(wb[sheet_name], table_name)[0]
            for accepted in [
                ['Vendor ID', 'ID'],
                ['Vendor', 'Vendor Name'],
                ['Technical'],
                ['Implementation'],
                ['Accessibility'],
                ['Price', 'Price Score'],
                ['Past Performance'],
                ['Weighted Total', 'Total Weighted Score', 'Composite Score'],
                ['Rank'],
                ['Recommendation', 'Status', 'Disposition', 'Recommendation/Status'],
            ]:
                _vopt_find_header_index(header, accepted)
    for sheet_name in EXPECT['hidden_sheets']:
        assert wb[sheet_name].sheet_state in ('hidden', 'veryHidden')
    names = {dn.name for dn in wb.defined_names.values()}
    for name in EXPECT['defined_names']:
        assert name in names
    for sheet_name, print_area in EXPECT['print_areas'].items():
        _vopt_assert_print_area_covers_table(wb[sheet_name], EXPECT['tables'][sheet_name][0], print_area)
    formula_expectations = {
        'Eligibility': ['IF', 'Pass', 'Eligible'],
        'Score Matrix price score': ['IF', 'Eligibility'],
        'Score Matrix weighted total': ['Eligibility', 'Weights'],
        'Score Matrix rank': ['RANK'],
        'Price Analysis eligibility': ['Eligibility'],
        'Price Analysis score link': ['Score Matrix'],
    }
    _vopt_sheet_has_formula_tokens(wb['Eligibility'], formula_expectations['Eligibility'], 'Eligibility')
    _vopt_sheet_has_formula_tokens(wb['Score Matrix'], formula_expectations['Score Matrix price score'], 'Score Matrix price score')
    _vopt_sheet_has_formula_tokens(wb['Score Matrix'], formula_expectations['Score Matrix weighted total'], 'Score Matrix weighted total')
    _vopt_sheet_has_formula_tokens(wb['Score Matrix'], formula_expectations['Score Matrix rank'], 'Score Matrix rank')
    _vopt_sheet_has_formula_tokens(wb['Price Analysis'], formula_expectations['Price Analysis eligibility'], 'Price Analysis eligibility')
    _vopt_sheet_has_formula_tokens(wb['Price Analysis'], formula_expectations['Price Analysis score link'], 'Price Analysis score link')
    for sheet_name, ranges in EXPECT['data_validation_ranges'].items():
        actual = _dv_ranges(wb[sheet_name])
        for expected in ranges:
            if sheet_name == 'Score Matrix' and expected == 'K2:K5':
                rows = _vopt_table_rows(wb['Score Matrix'], 'weighted_score_matrix')
                status_col = _vopt_find_header_index(rows[0], ['Recommendation', 'Status', 'Disposition', 'Recommendation/Status']) + 1
                status_range = f"{get_column_letter(status_col)}2:{get_column_letter(status_col)}5"
                assert status_range in actual, f"Score Matrix: expected data validation on status range {status_range}, found {list(actual)!r}"
            else:
                assert expected in actual
    assert sum(len(cf.rules) for sheet in ['Eligibility', 'Score Matrix', 'Price Analysis', 'Public Debrief'] for cf in wb[sheet].conditional_formatting) >= 1
    rows = _vopt_table_rows(wb['Public Debrief'], 'public_debrief_summary')
    _assert_public_rank_rows(rows, 'Public Debrief', wb=wb)


def test_scorebook_values_and_hidden_conflict_details():
    wb = load_workbook(_path('xlsx'), data_only=False)
    score_rows = _vopt_table_rows(wb['Score Matrix'], 'weighted_score_matrix')
    score_header, score_data = score_rows[0], score_rows[1:]
    vendor_col = _vopt_find_header_index(score_header, ['Vendor ID', 'ID'])
    status_col = _vopt_find_header_index(score_header, ['Recommendation', 'Status', 'Disposition', 'Recommendation/Status'])
    score_by_vendor = {_vopt_norm_text(row[vendor_col]): row for row in score_data}
    assert _status_matches(score_by_vendor['VND-B'][status_col], 'Recommended award')
    assert 'INELIGIBLE' in _vopt_norm_key(score_by_vendor['VND-C'][status_col])
    conflict_text = '\n'.join(str(cell.value) for row in wb['Conflict Log'].iter_rows() for cell in row if cell.value)
    require_all(conflict_text, ['Taylor', 'Prior employment with CivicStack', "Exclude Taylor's CivicStack score"], 'hidden conflict log')
    rules_text = '\n'.join(str(cell.value) for row in wb['Rules'].iter_rows() for cell in row if cell.value)
    require_all(rules_text, ['Recommended award', 'Alternate', 'Ineligible'], 'hidden rules sheet')


def test_cross_output_award_consistency():
    memo_text = _doc_text(Document(_path('docx')))
    wb = load_workbook(_path('xlsx'), data_only=False)
    debrief_rows = _vopt_table_rows(wb['Public Debrief'], 'public_debrief_summary')
    _assert_public_rank_rows(debrief_rows, 'Public Debrief', wb=wb)
    assert EXPECT['winner'] in memo_text
    assert EXPECT['ineligible_vendor'] in memo_text
    score_rows = _vopt_table_rows(wb['Score Matrix'], 'weighted_score_matrix')
    score_header, score_data = score_rows[0], score_rows[1:]
    vendor_col = _vopt_find_header_index(score_header, ['Vendor', 'Vendor Name'])
    status_col = _vopt_find_header_index(score_header, ['Recommendation', 'Status', 'Disposition', 'Recommendation/Status'])
    score_by_vendor = {_vopt_norm_text(row[vendor_col]): row for row in score_data}
    assert _status_matches(score_by_vendor[EXPECT['winner']][status_col], 'Recommended award')
    eligibility_rows = _vopt_table_rows(wb['Eligibility'], 'vendor_eligibility')
    eligibility_text = _vopt_rows_text(eligibility_rows)
    assert EXPECT['ineligible_vendor'] in eligibility_text
    northstar_line = next(
        (_vopt_norm_text('|'.join(row)) for row in _vopt_norm_rows(eligibility_rows) if EXPECT['ineligible_vendor'] in '|'.join(row)),
        '',
    )
    northstar_key = _vopt_norm_key(northstar_line)
    assert 'RESIDENCY' in northstar_key and ('FAIL' in northstar_key or 'INELIGIBLE' in northstar_key), (
        f"Eligibility should show Northstar's data-residency gate failure, found {northstar_line!r}"
    )
