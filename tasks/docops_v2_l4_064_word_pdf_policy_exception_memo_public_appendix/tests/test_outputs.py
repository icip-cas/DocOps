import json
import os
import re
import sys
import zipfile
from pathlib import Path
from datetime import date, datetime

import pdfplumber
from docx import Document
from pypdf import PdfReader

sys.path.insert(0, str(Path(__file__).parent))
from verifier_utils import *  # noqa: F401,F403

META_PATH = Path(os.environ.get('TASK_METADATA_PATH', '/tests/task_metadata.json'))
META = json.loads(META_PATH.read_text(encoding='utf-8'))
EXPECT = META['verifier_expectations']


def _vopt_norm_text(value):
    if value is None:
        return ''
    if isinstance(value, datetime):
        if value.hour or value.minute or value.second:
            return value.strftime('%Y-%m-%d %H:%M')
        return value.strftime('%Y-%m-%d')
    if isinstance(value, date):
        return value.strftime('%Y-%m-%d')
    if isinstance(value, float) and value.is_integer():
        value = int(value)
    text = str(value)
    text = text.replace('\u2013', '-').replace('\u2014', '-').replace('\u2212', '-')
    text = text.replace('\xa0', ' ')
    text = re.sub(r'\s+', ' ', text).strip()
    return text


def _vopt_norm_formula(value):
    return re.sub(r'\s+', '', _vopt_norm_text(value)).upper()


def _vopt_assert_formula(actual, expected, label='formula'):
    assert _vopt_norm_formula(actual) == _vopt_norm_formula(expected), f"{label}: expected {expected!r}, found {actual!r}"


def _vopt_norm_rows(rows):
    return [[_vopt_norm_text(cell) for cell in row] for row in rows]


def _vopt_rows_equal(actual, expected):
    return _vopt_norm_rows(actual) == _vopt_norm_rows(expected)


def _vopt_rows_text(rows):
    return '\n'.join('|'.join(_vopt_norm_text(cell) for cell in row) for row in rows)


def _vopt_header_has(header, required):
    actual = {_vopt_norm_text(cell).lower() for cell in header}
    missing = [cell for cell in required if _vopt_norm_text(cell).lower() not in actual]
    assert not missing, f"Missing expected header columns {missing!r}; actual={header!r}"


def _vopt_section_text(section_part):
    return '\n'.join(p.text for p in getattr(section_part, 'paragraphs', []) if p.text is not None)


def _vopt_header_text(doc):
    return '\n'.join(_vopt_section_text(section.header) for section in doc.sections)


def _vopt_footer_text(doc):
    return '\n'.join(_vopt_section_text(section.footer) for section in doc.sections)


def _vopt_table_ref(ws, expected_name=None):
    tables = ws.tables
    if expected_name:
        for name in tables.keys():
            if str(name).lower() == str(expected_name).lower():
                return tables[name].ref
    vals = list(tables.values())
    assert vals, f"No Excel native table found on sheet {ws.title!r}"
    return vals[0].ref


def _vopt_clean_area(value):
    text = _vopt_norm_text(value).replace("'", '').replace('$', '').replace(' ', '')
    if '!' in text:
        text = text.split('!', 1)[1]
    return text.upper()


def _vopt_assert_print_area(ws, expected):
    actual = _vopt_clean_area(ws.print_area)
    target = _vopt_clean_area(expected)
    assert target in actual or actual in target, f"{ws.title}: expected print area {expected!r}, found {ws.print_area!r}"


def _vopt_freeze_pane(value):
    return getattr(value, 'coordinate', value)


def _vopt_validated_cells(ws):
    cells = set()
    for dv in ws.data_validations.dataValidation:
        for rng in dv.cells.ranges:
            text = str(rng).replace('$', '')
            try:
                min_col, min_row, max_col, max_row = range_boundaries(text)
            except ValueError:
                cells.add(text)
                continue
            if max_row > 10000:
                max_row = max(min_row, 200)
            for row in range(min_row, max_row + 1):
                for col in range(min_col, max_col + 1):
                    cells.add(f"{get_column_letter(col)}{row}")
    return cells


class _VoptValidationRanges(list):
    def __init__(self, ws, ranges):
        super().__init__(ranges)
        self.ws = ws

    def __contains__(self, expected):
        expected = str(expected).replace('$', '')
        if super().__contains__(expected):
            return True
        try:
            min_col, min_row, max_col, max_row = range_boundaries(expected)
        except ValueError:
            return super().__contains__(expected)
        target = {f"{get_column_letter(col)}{row}" for row in range(min_row, max_row + 1) for col in range(min_col, max_col + 1)}
        return target.issubset(_vopt_validated_cells(self.ws))


def _vopt_dv_ranges(ws):
    ranges = []
    for dv in ws.data_validations.dataValidation:
        ranges.extend(str(rng).replace('$', '') for rng in dv.cells.ranges)
    return _VoptValidationRanges(ws, ranges)


def _vopt_ordered_subset(expected, actual):
    actual_iter = iter([_vopt_norm_text(item) for item in actual])
    for item in [_vopt_norm_text(value) for value in expected]:
        for candidate in actual_iter:
            if item == candidate:
                break
        else:
            return False
    return True



def _resolve_output(kind):
    if kind == 'memo':
        return Path(os.environ.get('DOCX_OUTPUT_PATH', EXPECT['memo_output']))
    if kind == 'appendix':
        return Path(os.environ.get('PDF_OUTPUT_PATH', EXPECT['appendix_output']))
    raise KeyError(kind)


def _doc_heading_order(doc):
    out = []
    for p in doc.paragraphs:
        style = p.style.name if p.style else ''
        if p.text.strip() and (style.startswith('Heading') or style == 'Title'):
            out.append(p.text.strip())
    return out


def _doc_text(doc):
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append('|'.join(cell.text.strip() for cell in row.cells))
    return '\n'.join(parts)


def _table_rows(table):
    return [[cell.text.strip() for cell in row.cells] for row in table.rows]


def _has_toc_field(path):
    with zipfile.ZipFile(path) as zf:
        xml = zf.read('word/document.xml').decode('utf-8', errors='ignore')
    return bool(re.search(r'TOC\s+(?:\\)?o|TOC\s*(?:&quot;|")', xml))


def _highlighted_para(doc, text):
    for p in doc.paragraphs:
        if _vopt_norm_text(text) in _vopt_norm_text(p.text):
            return any(run.font.highlight_color is not None for run in p.runs) or '<w:highlight' in p._p.xml
    raise AssertionError(f'Paragraph not found: {text}')


def _heading_key(text):
    text = _vopt_norm_text(text)
    text = re.sub(r'^\d+\.\s*', '', text)
    text = text.replace('–', '-').replace(':', ' - ')
    return normalize_text(text)


def _assert_major_headings(doc):
    headings = [_heading_key(text) for text in _doc_heading_order(doc)]
    assert headings and headings[0] == normalize_text('Policy Exception Release Memo'), f'Unexpected first heading: {_doc_heading_order(doc)!r}'
    required = ['Scope', 'Release Matrix', 'Exception Register']
    for item in required:
        assert normalize_text(item) in headings, f'Missing major heading {item!r}; headings={_doc_heading_order(doc)!r}'
    assert any('appendix a' in heading for heading in headings), (
        f'Missing Appendix A public heading; headings={_doc_heading_order(doc)!r}'
    )


def _assert_memo_semantics(text):
    require_all(text, ['EX-102', 'EX-118', 'EX-131', 'EX-144'], 'memo')
    for eid in EXPECT['public_exception_ids']:
        assert eid in text, f'memo: missing public exception ID {eid}'
    norm = normalize_text(text)
    assert 'public' in norm and any(token in norm for token in ('appendix', 'pdf', 'exception summary', 'release')), (
        'memo: must describe the public release/appendix boundary'
    )
    assert 'EX-144' in text and any(token in normalize_text(text) for token in ('closed', 'internal tracking', 'internal only')), (
        'memo: EX-144 must remain only as closed/internal tracking context'
    )


def _assert_release_matrix(doc):
    text = _doc_text(doc)
    expected = {
        'EX-102': ['CTRL-1', 'blocking'],
        'EX-118': ['CTRL-2', 'watch'],
        'EX-131': ['CTRL-4', 'blocking'],
    }
    norm = normalize_text(text)
    for eid, tokens in expected.items():
        assert normalize_text(eid) in norm, f'Release matrix/register missing {eid}'
        for token in tokens:
            assert normalize_text(token) in norm, f'{eid}: missing {token!r} release context'


def _assert_exception_register(doc):
    text = _doc_text(doc)
    expectations = {
        'EX-102': ['CTRL-1', 'MFA exception', 'Rosa Iyer', '2026-09-15'],
        'EX-118': ['CTRL-2', 'backup restore', 'Noah Patel', '2026-08-30'],
        'EX-131': ['CTRL-4', 'privileged access', 'Ava Chen', '2026-08-22'],
        'EX-144': ['CTRL-5', 'closed', 'Mina Zhou', '2026-07-19'],
    }
    for eid, tokens in expectations.items():
        require_all(text, [eid] + tokens, f'memo register {eid}')


def _highlighted_exception(doc, eid):
    for p in doc.paragraphs:
        if eid in p.text:
            if any(run.font.highlight_color is not None for run in p.runs) or '<w:highlight' in p._p.xml:
                return True
    return False


def _assert_pdf_outline(reader):
    titles = [title for _level, title in flatten_outline(reader.outline)]
    norm = normalize_text(' '.join(titles))
    assert 'appendix' in norm or 'summary' in norm or 'cover' in norm, f'Missing appendix/cover outline: {titles!r}'
    for eid in EXPECT['public_exception_ids']:
        assert normalize_text(eid) in norm, f'Missing outline entry for {eid}: {titles!r}'


def _assert_public_pdf_boundary(text):
    norm = normalize_text(text).replace('-', ' ')
    assert 'ex 144' not in norm, 'Closed exception EX-144 must not be published in public appendix'
    assert 'internal legal strategy' not in norm, 'Public appendix must not mention internal legal strategy'
    assert 'compensating control' not in norm, 'Public appendix must not mention compensating-control details'


def _pdf_texts(path):
    with pdfplumber.open(str(path)) as pdf:
        return [page.extract_text() or '' for page in pdf.pages]


def _page_titles(texts):
    out = []
    for text in texts:
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        out.append(lines[0] if lines else '')
    return out


def test_outputs_exist():
    memo = _resolve_output('memo')
    appendix = _resolve_output('appendix')
    assert memo.exists(), f'Missing memo output: {memo}'
    assert appendix.exists(), f'Missing appendix output: {appendix}'
    assert memo.suffix.lower() == '.docx'
    assert appendix.suffix.lower() == '.pdf'


def test_memo_structure_and_content():
    memo = _resolve_output('memo')
    doc = Document(memo)
    _assert_major_headings(doc)
    text = _doc_text(doc)
    _assert_memo_semantics(text)
    forbid_any(text, ['Manual TOC placeholder', 'Draft Notes - Remove', 'Internal legal strategy:'], 'memo')
    assert _has_toc_field(memo), 'Expected TOC field'
    assert len(doc.sections) >= EXPECT['section_count_min']
    assert EXPECT['header_contains'] in _vopt_header_text(doc)
    assert 'Internal Draft' not in _vopt_header_text(doc) + '\n' + _vopt_footer_text(doc), 'Draft header/footer must be removed'
    _assert_release_matrix(doc)
    _assert_exception_register(doc)
    for eid in ('EX-102', 'EX-131'):
        assert _highlighted_exception(doc, eid), f'Missing highlight on blocking exception {eid}'


def test_public_pdf_appendix():
    appendix = _resolve_output('appendix')
    reader = PdfReader(str(appendix))
    assert len(reader.pages) >= 1
    texts = _pdf_texts(appendix)
    full = '\n'.join(texts)
    require_all(full, ['Public Policy Exception Appendix', 'EX-102', 'CTRL-1', 'EX-118', 'CTRL-2', 'EX-131', 'CTRL-4'], 'public appendix')
    public_id_count = sum(1 for eid in EXPECT['public_exception_ids'] if eid in full)
    assert public_id_count == 3, f'public appendix should contain exactly the three public IDs, found {public_id_count}'
    _assert_public_pdf_boundary(full)
    _assert_pdf_outline(reader)


def test_cross_output_public_exception_ids():
    doc = Document(_resolve_output('memo'))
    memo_text = _doc_text(doc)
    pdf_text = '\n'.join(_pdf_texts(_resolve_output('appendix')))
    for eid in EXPECT['public_exception_ids']:
        assert eid in memo_text, f'{eid} missing from memo'
        assert eid in pdf_text, f'{eid} missing from public appendix'
    assert 'EX-144' in memo_text, 'Closed exception should remain in internal memo register'
    assert 'EX-144' not in pdf_text, 'Closed exception must not be published in public appendix'
