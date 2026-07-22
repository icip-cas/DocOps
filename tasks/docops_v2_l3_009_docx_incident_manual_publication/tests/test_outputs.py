import json
import sys
from pathlib import Path

from docx import Document

sys.path.insert(0, str(Path(__file__).parent))
from verifier_utils import *  # noqa: F401,F403

META_PATH = Path(__import__('os').environ.get('TASK_METADATA_PATH', '/tests/task_metadata.json'))
if not META_PATH.exists():
    META_PATH = Path(__file__).parent / 'task_metadata.json'
META = json.loads(META_PATH.read_text())
INPUT_PATH = Path(META['input_path'])
OUTPUT_PATH = Path(META['output_path'])


def headings(doc):
    return [
        p.text.strip()
        for p in doc.paragraphs
        if p.text.strip() and p.style and p.style.name.startswith('Heading')
    ]


def table_rows(table):
    return [[norm_cell(cell.text) for cell in row.cells] for row in table.rows]


def table_by_header(doc, expected_header):
    expected = [normalize_text(item) for item in expected_header]
    for table in doc.tables:
        rows = table_rows(table)
        if not rows:
            continue
        header = [normalize_text(item) for item in rows[0]]
        if header[:len(expected)] == expected:
            return rows
    raise AssertionError(f'Table with header not found: {expected_header!r}')


def all_doc_text(doc):
    chunks = docx_texts(doc)
    chunks.append(docx_header_text(doc))
    chunks.append(docx_footer_text(doc))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                chunks.append(cell.text)
    return '\n'.join(chunks)


def test_incident_manual_publication_longflow():
    run_preflight(META, INPUT_PATH, OUTPUT_PATH)
    doc = Document(OUTPUT_PATH)
    text = all_doc_text(doc)

    actual_headings = headings(doc)
    expected_headings = [
        'Publication Summary',
        'Response Timeline',
        'Customer Communication',
        'Control Exceptions',
        'Action Register',
        'Appendix A - Evidence Index',
        'Appendix B - Approval Record',
    ]
    if actual_headings and actual_headings[0] not in expected_headings:
        actual_headings = actual_headings[1:]
    assert actual_headings == expected_headings
    assert docx_has_toc_field(doc), 'Expected a real Word TOC field'
    assert 'Incident Response Manual | Public Release | 2026-06-05' in docx_header_text(doc)
    assert 'Page' in docx_footer_text(doc)

    for heading in ['Control Exceptions', 'Action Register', 'Appendix A - Evidence Index', 'Appendix B - Approval Record']:
        assert has_page_break_before(doc, heading), f'Missing required page break before {heading}'

    required = [
        'This public manual summarizes the June 2026 incident response without privileged root-cause analysis.',
        'Customer-facing language has been normalized and internal notes have been removed.',
        'Exception: MFA evidence owner remains unassigned.',
        'Resolved exception: log retention evidence was exported and approved.',
    ]
    for item in required:
        assert normalize_text(item) in normalize_text(text), f'Missing required publication text: {item}'

    forbidden = [
        'Draft:',
        'Working Copy',
        'Manual TOC placeholder',
        '[INTERNAL]',
        'vendor root cause suspected',
        'firewall rule owner was late',
        '555-0199',
        'Do not publish',
        'Draft footer',
        'Internal Only',
    ]
    for item in forbidden:
        assert normalize_text(item) not in normalize_text(text), f'Forbidden internal/draft text remained: {item}'

    assert len(doc.tables) >= 6
    timeline = table_by_header(doc, ['Time', 'Milestone', 'Public Summary'])
    assert timeline[0] == ['Time', 'Milestone', 'Public Summary']
    assert timeline[1] == ['2026-06-01 08:12', 'Alert opened', 'Incident response process started.']
    assert timeline[3] == ['2026-06-01 15:10', 'Service restored', 'Customer service was restored.']

    communication = table_by_header(doc, ['Audience', 'Message Owner', 'Required Update', 'Status'])
    assert communication[0] == ['Audience', 'Message Owner', 'Required Update', 'Status']
    assert communication[1] == ['Customers', 'Maya Chen', 'Send public restoration summary.', 'Ready']
    assert communication[2] == ['Support desk', 'Jon Bell', 'Use approved FAQ language only.', 'Ready']

    exceptions = table_by_header(doc, ['Exception ID', 'Control', 'Owner', 'Release Status'])
    assert exceptions[0] == ['Exception ID', 'Control', 'Owner', 'Release Status']
    assert exceptions[1] == ['EX-01', 'MFA evidence', 'Unassigned', 'Open - owner required']
    assert exceptions[2] == ['EX-02', 'Log retention', 'Rina Patel', 'Closed - evidence exported']

    actions = table_by_header(doc, ['Action ID', 'Owner', 'Due Date', 'Publish Status'])
    assert actions[0] == ['Action ID', 'Owner', 'Due Date', 'Publish Status']
    assert actions[1] == ['ACT-01', 'Unassigned', '2026-06-07', 'Blocked until owner assigned']
    assert actions[2] == ['ACT-02', 'Maya Chen', '2026-06-06', 'Ready']

    evidence = table_by_header(doc, ['Evidence ID', 'Public Label', 'Disposition'])
    assert evidence[0] == ['Evidence ID', 'Public Label', 'Disposition']
    assert evidence[1] == ['EV-01', 'Alert chronology', 'Publish']
    assert evidence[2] == ['EV-02', 'Internal Slack excerpt', 'Exclude']
    assert evidence[3] == ['EV-03', 'Customer notification record', 'Publish']

    approvals = table_by_header(doc, ['Approver', 'Role', 'Decision'])
    assert approvals[0] == ['Approver', 'Role', 'Decision']
    assert approvals[1] == ['Maya Chen', 'Incident Lead', 'Approved']
    assert approvals[2] == ['Jon Bell', 'Legal', 'Approved after privileged text removal']

    para = docx_para_by_text(doc, 'Exception: MFA evidence owner remains unassigned.')
    assert docx_para_has_highlight(para), 'Open exception paragraph should be highlighted'
