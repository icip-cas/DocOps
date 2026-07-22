
import json, sys
from pathlib import Path
from docx import Document
sys.path.insert(0, str(Path(__file__).parent))
from verifier_utils import *
META_PATH=Path(__import__('os').environ.get('TASK_METADATA_PATH','/tests/task_metadata.json'))
if not META_PATH.exists():
    META_PATH=Path(__file__).parent/'task_metadata.json'
META=json.loads(META_PATH.read_text()); INPUT_PATH=Path(META['input_path']); OUTPUT_PATH=Path(META['output_path'])
def heads(doc): return [p.text.strip() for p in doc.paragraphs if p.text.strip() and p.style and p.style.name.startswith('Heading')]
def test_policy_exception_cascade():
    run_preflight(META, INPUT_PATH, OUTPUT_PATH)
    doc=Document(OUTPUT_PATH)
    assert heads(doc)[-5:]==['Scope','Controls','Exceptions','Appendix A - Definitions','Appendix B - Evidence']
    assert docx_has_toc_field(doc)
    assert has_page_break_before(doc,'Appendix A - Definitions') and has_page_break_before(doc,'Appendix B - Evidence')
    texts=docx_texts(doc)
    assert 'Manual TOC placeholder' not in texts and not any(t.startswith('Draft:') for t in texts)
    header_text = docx_header_text(doc)
    footer_text = docx_footer_text(doc)
    assert 'Policy Manual | Exceptions Open: 2 | 2026-06-05' in header_text
    assert 'Page' in footer_text
    for forbidden in ['Confidential', 'Draft']:
        assert forbidden not in header_text
        assert forbidden not in footer_text
    rel=doc.tables[0]; exc=doc.tables[1]
    assert_table_header_contains(rel, 4, ['Control','Requirement','Exception Owner','Release Status'], 'release table')
    assert len(rel.rows)==3 and len(rel.columns)==4
    assert rel.cell(1,0).text.strip()=='CTRL-1'
    assert rel.cell(1,1).text.strip()=='MFA required'
    assert rel.cell(1,2).text.strip()=='Unassigned'
    assert rel.cell(1,3).text.strip()=='Blocked - MFA rollout missing owner'
    assert rel.cell(2,0).text.strip()=='CTRL-2'
    assert rel.cell(2,1).text.strip()=='Vendor review required'
    assert rel.cell(2,2).text.strip()=='Compliance team'
    assert rel.cell(2,3).text.strip()=='Blocked - Vendor review overdue'
    assert_table_header_contains(exc, 3, ['Exception','Owner','Due Date'], 'exceptions table')
    assert len(exc.rows)==3 and len(exc.columns)==3
    assert exc.cell(1,0).text.strip()=='MFA rollout missing owner'
    assert exc.cell(1,1).text.strip()=='Unassigned'
    assert exc.cell(1,2).text.strip()=='2026-06-12'
    assert exc.cell(2,0).text.strip()=='Vendor review overdue'
    assert exc.cell(2,1).text.strip()=='Compliance team'
    assert exc.cell(2,2).text.strip()=='2026-06-14'
    assert docx_para_has_highlight(docx_para_by_text(doc,'Exception: MFA rollout missing owner.'))
