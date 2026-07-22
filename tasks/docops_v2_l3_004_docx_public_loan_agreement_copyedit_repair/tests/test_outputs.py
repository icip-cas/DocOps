import json
import os
import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt

from verifier_utils import normalize_text, run_preflight

META_PATH = Path(os.environ.get("TASK_METADATA_PATH", "/tests/task_metadata.json"))
if not META_PATH.exists():
    META_PATH = Path(__file__).parent / "task_metadata.json"
META = json.loads(META_PATH.read_text(encoding="utf-8"))
EXPECT = META["verifier_expectations"]
INPUT_PATH = Path(os.environ.get("INPUT_PATH", META["input_path"]))
OUTPUT_PATH = Path(os.environ.get("OUTPUT_PATH", META["output_path"]))


def doc_text(doc):
    parts = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text.strip())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())
    for section in doc.sections:
        for paragraph in section.footer.paragraphs:
            if paragraph.text.strip():
                parts.append(paragraph.text.strip())
    core = doc.core_properties
    parts.extend(str(value or "") for value in [core.author, core.title, core.subject, core.keywords, core.comments])
    return "\n".join(parts)


def visible_heading_sequence(doc):
    wanted = set(EXPECT["headings"])
    return [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip() in wanted]


def paragraph_by_text(doc, text):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise AssertionError(f"Missing paragraph: {text}")


def require_all_terms(text, phrases, label):
    text_norm = normalize_text(text)
    missing = [phrase for phrase in phrases if normalize_text(phrase) not in text_norm]
    assert not missing, f"{label}: missing required public terms: {missing}"


def require_any(text, phrases, label):
    text_norm = normalize_text(text)
    assert any(normalize_text(phrase) in text_norm for phrase in phrases), (
        f"{label}: expected at least one of {phrases!r}"
    )


def allowed_cleanup_context(text, start, end):
    window = text[max(0, start - 150): min(len(text), end + 220)]
    return re.search(
        r"\b(no|not|removed|repaired|cleaned|cleanup|without|free of|obsolete|remnants? remain)\b",
        window,
        re.IGNORECASE,
    ) is not None


def assert_no_forbidden_remnants(text):
    text_norm = normalize_text(text)
    hard_forbidden = [
        "Riverbend Musuem",
        "Riverbend Mueseum",
        "Northstar Gallary",
        "40 objects",
        "$24,000",
        "$2,4000",
        "2026-09-31",
        "2026-10-12",
        "2026-10-05",
        "2027-01-02",
        "35-65% RH",
        "recieve",
        "enviroment",
        "accomodation",
        "insurence",
        "liablity",
        "temperatre",
        "securty",
    ]
    hits = [phrase for phrase in hard_forbidden if normalize_text(phrase) in text_norm]
    assert not hits, f"Forbidden error remnants still present: {hits}"

    contextual_hits = []
    for phrase in ["DRAFT", "PRIVATE", "TODO", "uncertain", "maybe"]:
        flags = 0 if phrase in {"DRAFT", "PRIVATE", "TODO"} else re.IGNORECASE
        for match in re.finditer(rf"\b{re.escape(phrase)}\b", text, flags):
            if not allowed_cleanup_context(text, match.start(), match.end()):
                contextual_hits.append(phrase)
                break
    assert not contextual_hits, f"Draft/private/uncertainty markers still present: {contextual_hits}"


def first_run(paragraph):
    for run in paragraph.runs:
        if run.text.strip():
            return run
    return paragraph.runs[0] if paragraph.runs else None


def all_table_rows(doc):
    rows = []
    for table in doc.tables:
        table_rows = []
        for row in table.rows:
            table_rows.append([cell.text.strip() for cell in row.cells])
        rows.append(table_rows)
    return rows


def cell_fill(cell):
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    shd = cell._tc.find(".//w:shd", namespaces=ns)
    if shd is None:
        return None
    return shd.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill")


def has_page_break(paragraph):
    xml = paragraph._p.xml
    return 'w:type="page"' in xml or "w:type='page'" in xml


def assert_break_before(doc, heading):
    previous = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == heading:
            assert previous is not None, f"{heading}: no paragraph before heading"
            assert has_page_break(previous) or has_page_break(paragraph), f"{heading}: missing page break before heading"
            return
        previous = paragraph
    raise AssertionError(f"Heading not found for page break: {heading}")


def assert_table_header_style(doc):
    style = EXPECT["style"]
    for table in doc.tables:
        for cell in table.rows[0].cells:
            assert cell_fill(cell) == style["table_header_fill"]
            runs = [run for paragraph in cell.paragraphs for run in paragraph.runs if run.text.strip()]
            assert runs, "Header cell has no text run"
            assert all(run.font.bold for run in runs)
            assert all(str(run.font.color.rgb) == style["table_header_font"] for run in runs)


def test_output_exists_and_is_docx():
    run_preflight(META, INPUT_PATH, OUTPUT_PATH)
    assert OUTPUT_PATH != INPUT_PATH


def test_headings_required_text_tables_and_forbidden_errors_removed():
    doc = Document(OUTPUT_PATH)
    assert visible_heading_sequence(doc) == EXPECT["headings"]
    text = doc_text(doc)
    require_all_terms(
        text,
        [
            "Riverbend Museum",
            "Northstar Gallery",
            "42",
            "$240,000",
            "2026-09-30",
            "2026-10-02",
            "2026-10-15",
            "2027-01-20",
            "68-72 F and 45-55% RH",
        ],
        "corrected public facts",
    )
    require_any(text, ["correction register", "corrections applied", "applied corrections"], "correction register")
    require_any(text, ["workflow", "condition report", "condition reporting"], "condition workflow")
    require_any(text, ["checklist", "public communication", "public wording"], "public checklist")
    require_any(text, ["appendix", "final review", "consistency"], "final appendix")
    require_any(text, ["spelling", "terminology", "writing"], "spelling repair")
    assert_no_forbidden_remnants(text)
    tables = all_table_rows(doc)
    assert len(tables) >= 4, "Expected multiple rebuilt public tables"


def test_style_layout_footer_and_page_breaks():
    doc = Document(OUTPUT_PATH)
    style = EXPECT["style"]
    title = paragraph_by_text(doc, EXPECT["headings"][0])
    title_run = first_run(title)
    assert title_run.font.name == style["title_font"]
    assert title_run.font.bold
    assert title_run.font.size == Pt(style["title_size_pt"])
    assert str(title_run.font.color.rgb) == style["title_color"]
    for heading in EXPECT["headings"][1:]:
        run = first_run(paragraph_by_text(doc, heading))
        assert str(run.font.color.rgb) == style["heading_color"]
        assert run.font.bold
    for section in doc.sections:
        assert abs(section.top_margin.inches - style["margin_inches"]) < 0.02
        assert abs(section.bottom_margin.inches - style["margin_inches"]) < 0.02
        assert abs(section.left_margin.inches - style["margin_inches"]) < 0.02
        assert abs(section.right_margin.inches - style["margin_inches"]) < 0.02
        footer_text = "\n".join(p.text for p in section.footer.paragraphs)
        assert style["footer_text"] in footer_text
    assert_table_header_style(doc)
    for heading in EXPECT["page_break_before"]:
        assert_break_before(doc, heading)


def test_public_metadata_applied():
    doc = Document(OUTPUT_PATH)
    core = doc.core_properties
    assert core.author == EXPECT["metadata"]["author"]
    assert core.title == EXPECT["metadata"]["title"]
    meta_text = "\n".join(str(value or "") for value in [core.author, core.title, core.subject, core.keywords, core.comments])
    hits = [phrase for phrase in EXPECT["forbidden_phrases"] if normalize_text(phrase) in normalize_text(meta_text)]
    assert not hits, f"Forbidden metadata remnants still present: {hits}"


def test_source_artifact_was_not_modified():
    source = Document(INPUT_PATH)
    text = doc_text(source)
    missing = [phrase for phrase in EXPECT["source_must_contain"] if normalize_text(phrase) not in normalize_text(text)]
    assert not missing, f"Source artifact no longer contains expected errors: {missing}"
