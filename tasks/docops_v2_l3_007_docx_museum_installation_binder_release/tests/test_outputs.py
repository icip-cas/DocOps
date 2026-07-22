import json
import os
import re
from pathlib import Path

from docx import Document

from verifier_utils import (
    docx_footer_text,
    docx_has_page_break,
    docx_has_toc_field,
    docx_header_text,
    docx_texts,
    forbid_any,
    normalize_text,
    run_preflight,
)

META_PATH = Path(os.environ.get("TASK_METADATA_PATH", "/tests/task_metadata.json"))
if not META_PATH.exists():
    META_PATH = Path(__file__).parent / "task_metadata.json"
META = json.loads(META_PATH.read_text(encoding="utf-8"))
EXPECT = META["verifier_expectations"]
INPUT_PATH = Path(os.environ.get("INPUT_PATH", META["input_path"]))
OUTPUT_PATH = Path(os.environ.get("OUTPUT_PATH", META["output_path"]))


def heading_order(doc):
    out = []
    first = next((p.text.strip() for p in doc.paragraphs if p.text.strip()), "")
    if first:
        out.append(first)
    for p in doc.paragraphs:
        style = p.style.name if p.style else ""
        text = p.text.strip()
        if normalize_text(text) in {"table of contents", "toc"}:
            continue
        if text and (style == "Title" or style.startswith("Heading")):
            if not out or normalize_text(text) != normalize_text(out[-1]):
                out.append(text)
    return out


def comparable_text(value):
    text = normalize_text(str(value))
    text = re.sub(r"\s*,\s*and\s+", " and ", text)
    text = re.sub(r"[.。;；:：]+$", "", text)
    return " ".join(text.split())


def contains_all(text, phrases):
    norm = comparable_text(text)
    return all(comparable_text(phrase) in norm for phrase in phrases)


def assert_public_content(text):
    checks = {
        "binder purpose": ["installation binder", "East Wing textile exhibition"],
        "public object scope": ["11 loaned objects", "three galleries", "6 mounts", "4 case rotations", "2 courier-supervised installations"],
        "environmental targets": ["68-72 degrees Fahrenheit", "45-55 percent relative humidity"],
        "label production": ["label copy", "approved", "production"],
    }
    missing = [label for label, phrases in checks.items() if not contains_all(text, phrases)]
    assert not missing, f"museum installation binder: missing public content groups: {missing}"


def value_matches(header, actual, expected):
    actual_norm = comparable_text(actual)
    expected_norm = comparable_text(expected)
    if actual_norm == expected_norm:
        return True
    if expected_norm in actual_norm or actual_norm in expected_norm:
        return True
    if expected_norm == "loose edge threads noted before mount approval":
        return actual_norm == "loose edge threads noted before mount application"
    if expected_norm == "light sensitivity requires stricter gallery b limits":
        return actual_norm == "light sensitivity requires stricter gallery b settings"
    return False


def has_page_break_before_flexible(doc, heading_text):
    prev = None
    wanted = comparable_text(heading_text)
    for p in doc.paragraphs:
        if comparable_text(p.text.strip()) == wanted:
            if prev is None:
                return docx_has_page_break(p)
            return docx_has_page_break(prev) or docx_has_page_break(p)
        prev = p
    raise AssertionError(f"Heading not found for page-break check: {heading_text}")


def all_visible_text(doc):
    parts = docx_texts(doc)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    parts.append(text)
    header = docx_header_text(doc)
    footer = docx_footer_text(doc)
    if header:
        parts.append(header)
    if footer:
        parts.append(footer)
    return "\n".join(parts)


def table_with_headers(doc, headers):
    wanted = [normalize_text(h) for h in headers]
    for table in doc.tables:
        if not table.rows:
            continue
        actual = [normalize_text(cell.text) for cell in table.rows[0].cells]
        if all(h in actual for h in wanted):
            return table, {h: actual.index(normalize_text(h)) for h in headers}
    raise AssertionError(f"Missing table with headers: {headers}")


def row_by_key(table, key_col, key_value):
    for row in table.rows[1:]:
        if normalize_text(row.cells[key_col].text) == normalize_text(str(key_value)):
            return row
    raise AssertionError(f"Missing row keyed by {key_value!r}")


def assert_table_exact(doc, expected_rows, key_header):
    table, cols = table_with_headers(doc, expected_rows[0])
    assert len(table.rows) == len(expected_rows), f"{key_header}: wrong row count"
    assert len(table.columns) >= len(expected_rows[0]), f"{key_header}: wrong column count"
    key_col = cols[key_header]
    for expected in expected_rows[1:]:
        row = row_by_key(table, key_col, expected[expected_rows[0].index(key_header)])
        for header, expected_value in zip(expected_rows[0], expected):
            actual = row.cells[cols[header]].text.strip()
            assert value_matches(header, actual, expected_value), (
                f"{key_header} row {expected[0]!r} {header}: expected {expected_value!r}, found {actual!r}"
            )


def test_output_exists_and_is_docx():
    run_preflight(META, INPUT_PATH, OUTPUT_PATH)
    assert OUTPUT_PATH != INPUT_PATH


def test_title_sections_and_public_text():
    doc = Document(OUTPUT_PATH)
    headings = heading_order(doc)
    if headings and normalize_text(headings[0]) == normalize_text(EXPECT["title"]):
        section_headings = headings[1:]
    else:
        raise AssertionError(f"Document title missing or wrong: {headings[:2]!r}")
    assert [normalize_text(h) for h in section_headings] == [normalize_text(h) for h in EXPECT["heading_order"]]
    text = all_visible_text(doc)
    assert_public_content(text)
    forbid_any(text, EXPECT["forbidden_phrases"], "museum installation binder")


def test_native_tables_and_installation_workflow():
    doc = Document(OUTPUT_PATH)
    assert_table_exact(doc, EXPECT["loan_rows"], "Object ID")
    assert_table_exact(doc, EXPECT["schedule_rows"], "Phase")
    assert_table_exact(doc, EXPECT["environment_rows"], "Zone")
    assert_table_exact(doc, EXPECT["label_rows"], "Label ID")
    assert_table_exact(doc, EXPECT["condition_rows"], "Exception ID")
    assert_table_exact(doc, EXPECT["accessibility_rows"], "Item")
    assert_table_exact(doc, EXPECT["contact_rows"], "Role")


def test_appendix_toc_header_footer():
    doc = Document(OUTPUT_PATH)
    assert docx_has_toc_field(doc), "Expected a real Word TOC field"
    for heading in EXPECT["page_break_before"]:
        assert has_page_break_before_flexible(doc, heading), f"Missing page break before {heading}"
    header = docx_header_text(doc)
    footer = docx_footer_text(doc)
    for phrase in EXPECT["header_contains"]:
        assert phrase in header, f"Missing header phrase: {phrase}"
    for phrase in EXPECT["footer_contains"]:
        assert phrase in footer, f"Missing footer phrase: {phrase}"
    assert "DRAFT" not in header.upper()
    assert "INTERNAL" not in header.upper()


def test_source_artifact_was_not_modified():
    source = Document(INPUT_PATH)
    source_text = all_visible_text(source)
    assert "DRAFT - East Wing Textile Exhibition Installation Binder Working Draft" in source_text
    assert "parking lot" in normalize_text(source_text)
    assert "unreleased valuation" in normalize_text(source_text)
