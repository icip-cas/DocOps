import json
import os
import re
from datetime import date, datetime
from pathlib import Path

from docx import Document
from openpyxl import load_workbook
from openpyxl.utils.cell import get_column_letter, range_boundaries

from verifier_utils import *  # noqa: F401,F403

META_PATH = Path(os.environ.get("TASK_METADATA_PATH", "/tests/task_metadata.json"))
META = json.loads(META_PATH.read_text(encoding="utf-8"))
EXPECT = META["verifier_expectations"]


def _norm(value):
    if value is None:
        return ""
    if isinstance(value, datetime):
        return value.strftime("%Y-%m-%d")
    if isinstance(value, date):
        return value.strftime("%Y-%m-%d")
    if isinstance(value, float) and value.is_integer():
        value = int(value)
    text = str(value).replace("\u2013", "-").replace("\u2014", "-").replace("\xa0", " ")
    return re.sub(r"\s+", " ", text).strip()


def _norm_formula(value):
    return re.sub(r"\s+", "", _norm(value)).upper()


def _rows(ws, start, end, max_col):
    return [[_norm(ws.cell(r, c).value) for c in range(1, max_col + 1)] for r in range(start, end + 1)]


def _expected_rows(rows):
    return [[_norm(c) for c in row] for row in rows]


def _doc_text(doc):
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append("|".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts)


def _heading_order(doc):
    out = []
    for p in doc.paragraphs:
        style = p.style.name if p.style else ""
        if p.text.strip() and (style == "Title" or style.startswith("Heading")):
            out.append(p.text.strip())
    return out


def _header_footer_text(doc):
    headers = []
    footers = []
    for section in doc.sections:
        headers.extend(p.text for p in section.header.paragraphs)
        footers.extend(p.text for p in section.footer.paragraphs)
    return "\n".join(headers), "\n".join(footers)


def _table_rows(table):
    return [[cell.text.strip() for cell in row.cells] for row in table.rows]


def _path(kind):
    env = {"docx": "DOCX_OUTPUT_PATH", "xlsx": "XLSX_OUTPUT_PATH"}[kind]
    key = {"docx": "docx_output", "xlsx": "xlsx_output"}[kind]
    return Path(os.environ.get(env, EXPECT[key]))


def _table_ref(ws, expected_name):
    for name in ws.tables.keys():
        if str(name).lower() == str(expected_name).lower():
            return ws.tables[name].ref
    raise AssertionError(f"Missing table {expected_name!r} on {ws.title!r}; found {list(ws.tables.keys())!r}")


def _clean_area(value):
    text = _norm(value).replace("'", "").replace("$", "").replace(" ", "")
    if "!" in text:
        text = text.split("!", 1)[1]
    return text.upper()


def _assert_print_area(ws, expected):
    actual = _clean_area(ws.print_area)
    target = _clean_area(expected)
    assert target in actual or actual in target, f"{ws.title}: expected print area {expected!r}, found {ws.print_area!r}"


def _validated_cells(ws):
    cells = set()
    for dv in ws.data_validations.dataValidation:
        for rng in dv.cells.ranges:
            text = str(rng).replace("$", "")
            min_col, min_row, max_col, max_row = range_boundaries(text)
            for row in range(min_row, max_row + 1):
                for col in range(min_col, max_col + 1):
                    cells.add(f"{get_column_letter(col)}{row}")
    return cells


def _range_covered(ws, expected):
    expected = expected.replace("$", "")
    min_col, min_row, max_col, max_row = range_boundaries(expected)
    target = {f"{get_column_letter(col)}{row}" for row in range(min_row, max_row + 1) for col in range(min_col, max_col + 1)}
    return target.issubset(_validated_cells(ws))


def _sheet_text(wb, sheets):
    parts = []
    for sheet in sheets:
        ws = wb[sheet]
        for row in ws.iter_rows():
            for cell in row:
                if cell.value is not None:
                    parts.append(str(cell.value))
    return "\n".join(parts)


def test_outputs_exist():
    assert _path("docx").exists()
    assert _path("xlsx").exists()
    assert _path("docx").suffix.lower() == ".docx"
    assert _path("xlsx").suffix.lower() == ".xlsx"


def test_appeal_letter_structure_content_and_privacy():
    doc = Document(_path("docx"))
    assert _heading_order(doc) == EXPECT["doc_heading_order"]
    text = _doc_text(doc)
    require_all(text, EXPECT["doc_required"], "appeal letter")
    forbid_any(text, EXPECT["doc_forbidden"], "appeal letter")
    header, footer = _header_footer_text(doc)
    assert EXPECT["doc_header"] in header
    assert EXPECT["doc_footer"] in footer
    assert len(doc.tables) >= 1
    assert _table_rows(doc.tables[0]) == EXPECT["criteria_rows"]


def test_workbook_structure_tables_controls():
    wb = load_workbook(_path("xlsx"), data_only=False)
    assert wb.sheetnames == EXPECT["sheet_order"]
    for sheet, (table_name, ref) in EXPECT["tables"].items():
        assert _table_ref(wb[sheet], table_name) == ref
    for sheet in EXPECT["hidden_sheets"]:
        assert wb[sheet].sheet_state in ("hidden", "veryHidden")
    names = {dn.name for dn in wb.defined_names.values()}
    for name in EXPECT["defined_names"]:
        assert name in names
    for ref, expected in EXPECT["formula_cells"].items():
        sheet, cell = ref.split("!")
        assert _norm_formula(wb[sheet][cell].value) == _norm_formula(expected), f"{ref}: expected {expected!r}, found {wb[sheet][cell].value!r}"
    for sheet, ranges in EXPECT["data_validation"].items():
        for rng in ranges:
            assert _range_covered(wb[sheet], rng), f"{sheet}: missing validation over {rng}"
    for sheet, area in EXPECT["print_areas"].items():
        _assert_print_area(wb[sheet], area)


def test_workbook_values_and_privacy():
    wb = load_workbook(_path("xlsx"), data_only=False)
    assert _rows(wb["Appeal Summary"], 1, 13, 4) == _expected_rows(EXPECT["summary_rows"])
    assert _rows(wb["Criteria Mapping"], 1, 6, 6) == _expected_rows(EXPECT["criteria_mapping_rows"])
    assert _rows(wb["Evidence Index"], 1, 7, 5) == _expected_rows(EXPECT["evidence_rows"])
    assert _rows(wb["Timeline"], 1, 5, 6) == _expected_rows(EXPECT["timeline_rows"])
    assert _rows(wb["Attachments Checklist"], 1, 7, 5) == _expected_rows(EXPECT["attachments_rows"])
    assert _rows(wb["Privacy Review"], 1, 5, 4) == _expected_rows(EXPECT["privacy_rows"])
    assert _rows(wb["Raw Claims"], 1, 7, 4) == _expected_rows(EXPECT["raw_claim_rows"])
    assert _rows(wb["Payer Rules"], 1, 4, 2) == _expected_rows(EXPECT["payer_rule_rows"])
    public = _sheet_text(wb, ["Appeal Summary", "Criteria Mapping", "Evidence Index", "Timeline", "Attachments Checklist", "Privacy Review"])
    forbid_any(public, EXPECT["forbidden_public"], "public workbook sheets")


def test_cross_output_consistency():
    doc_text = _doc_text(Document(_path("docx")))
    wb = load_workbook(_path("xlsx"), data_only=False)
    assert wb["Appeal Summary"]["B6"].value == "=B5+60"
    assert wb["Appeal Summary"]["B12"].value == '=COUNTIF(\'Criteria Mapping\'!E2:E6,"Met")'
    assert wb["Appeal Summary"]["B13"].value == '=IF(B12=5,"Ready","Hold")'
    for anchor in ["DEN-7842", "AC-77824", "Rimegepant ODT 75 mg", "8 tablets per 30 days", "G43.709", "2026-08-06"]:
        assert anchor in doc_text
        assert anchor in _sheet_text(wb, ["Appeal Summary", "Criteria Mapping", "Evidence Index", "Timeline", "Attachments Checklist", "Privacy Review"])
    assert [wb["Criteria Mapping"].cell(row, 5).value for row in range(2, 7)] == ["Met", "Met", "Met", "Met", "Met"]
    assert wb["Attachments Checklist"]["C7"].value == "No"
    for forbidden in EXPECT["forbidden_public"]:
        assert normalize_text(forbidden) not in normalize_text(doc_text)
