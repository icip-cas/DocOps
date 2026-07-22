import json
import os
import re
from datetime import date, datetime
from pathlib import Path

from docx import Document
from openpyxl import load_workbook
from openpyxl.utils.cell import get_column_letter, range_boundaries

from verifier_utils import *  # noqa: F401,F403

META_PATH = Path(os.environ.get("TASK_METADATA_PATH", "/tests/task_metadata.json"))
META = json.loads(META_PATH.read_text(encoding="utf-8"))
EXPECT = META["verifier_expectations"]

PUBLIC_FORBIDDEN = [item for item in EXPECT["forbidden_public"] if item.casefold() != "private ap risk note"] + [
    "fraudwatch",
    "fraud-watch",
    "blacklisted",
    "slow-pay",
    "slow pay",
    "embarrass",
    "do not loop procurement",
    "don't tell procurement",
    "hold every invoice",
    "hold invoices forever",
]


def _norm(value):
    if value is None:
        return ""
    if isinstance(value, datetime):
        return value.strftime("%Y-%m-%d")
    if isinstance(value, date):
        return value.strftime("%Y-%m-%d")
    if isinstance(value, float) and value.is_integer():
        value = int(value)
    text = str(value).replace("\u2013", "-").replace("\u2014", "-").replace("\xa0", " ")
    return re.sub(r"\s+", " ", text).strip()


def _norm_key(value):
    return re.sub(r"[^a-z0-9]+", "", _norm(value).lower())


def _norm_formula(value):
    return re.sub(r"\s+", "", _norm(value)).upper()


def _path(kind):
    env = {"docx": "DOCX_OUTPUT_PATH", "xlsx": "XLSX_OUTPUT_PATH"}[kind]
    key = {"docx": "docx_output", "xlsx": "xlsx_output"}[kind]
    return Path(os.environ.get(env, EXPECT[key]))


def _doc_text(doc):
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append("|".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts)


def _heading_order(doc):
    out = []
    for p in doc.paragraphs:
        style = p.style.name if p.style else ""
        text = p.text.strip()
        if text and (style == "Title" or style.startswith("Heading")):
            out.append(text)
        elif text:
            out.append(text)
    return out


def _norm_heading_label(value):
    text = _norm(value)
    text = re.sub(r"^(?:\d+|[IVXLCDM]+)[\.)]\s+", "", text, flags=re.I)
    text = re.sub(r"^[A-Z][\.)]\s+", "", text)
    text = re.sub(r"[:\-]+$", "", text).strip()
    return text.casefold()


def _heading_sequence_matches(actual, expected):
    actual_norm = [_norm_heading_label(item) for item in actual]
    expected_norm = [_norm_heading_label(item) for item in expected]
    pos = 0
    for item in actual_norm:
        if pos < len(expected_norm) and item == expected_norm[pos]:
            pos += 1
    return pos == len(expected_norm)


def _header_footer_text(doc):
    h, f = [], []
    for section in doc.sections:
        h.extend(p.text for p in section.header.paragraphs)
        f.extend(p.text for p in section.footer.paragraphs)
    return "\n".join(h), "\n".join(f)


def _table_rows(table):
    return [[cell.text.strip() for cell in row.cells] for row in table.rows]


def _has_decision_table(doc):
    doc_text = normalize_text(_doc_text(doc))
    for table in doc.tables:
        rows = _table_rows(table)
        if not rows:
            continue
        text = normalize_text("\n".join(" | ".join(row) for row in rows))
        header = normalize_text(" | ".join(rows[0]))
        has_cols = all(term in header for term in ["vendor", "action", "backup withholding", "reason"])
        has_orion = "orion creative llc" in text and any(term in text for term in ["$0.00", "none", "no withholding"])
        has_mesa = (
            "mesa field services" in text
            and "$2,016.00" in text
            and "inv-7781" in doc_text
            and "$6,384.00" in doc_text
        )
        has_brightline = "brightline legal pllc" in text and any(term in text for term in ["$0.00", "none", "no withholding"])
        if has_cols and has_orion and has_mesa and has_brightline:
            return True
    return False


def _sheet(wb, sheet_name):
    assert sheet_name in wb.sheetnames, f"Missing sheet {sheet_name!r}; found {wb.sheetnames!r}"
    return wb[sheet_name]


def _table(ws, expected_name):
    for name in ws.tables.keys():
        if str(name).lower() == str(expected_name).lower():
            return ws.tables[name]
    raise AssertionError(f"Missing table {expected_name!r} on {ws.title!r}; found {list(ws.tables.keys())!r}")


def _table_bounds(table):
    return range_boundaries(str(table.ref).replace("$", ""))


def _table_has_size(ws, expected_name, min_rows, min_cols):
    table = _table(ws, expected_name)
    min_col, min_row, max_col, max_row = _table_bounds(table)
    assert max_row - min_row + 1 >= min_rows, f"{expected_name}: table has too few rows"
    assert max_col - min_col + 1 >= min_cols, f"{expected_name}: table has too few columns"
    return table


def _table_headers(ws, table):
    min_col, min_row, max_col, _ = _table_bounds(table)
    return [_norm(ws.cell(min_row, col).value) for col in range(min_col, max_col + 1)]


def _table_header_map(ws, table):
    min_col, min_row, max_col, _ = _table_bounds(table)
    mapping = {}
    for col in range(min_col, max_col + 1):
        mapping[_norm_key(ws.cell(min_row, col).value)] = col
    return mapping


def _require_headers(ws, table, expected_headers):
    actual = [_norm_key(h) for h in _table_headers(ws, table)]
    missing = [h for h in expected_headers if _norm_key(h) not in actual]
    assert not missing, f"{ws.title}: missing table headers {missing!r}"


def _find_sheet_row(ws, row_token, table, start_data=True):
    min_col, min_row, max_col, max_row = _table_bounds(table)
    start = min_row + 1 if start_data else min_row
    token = _norm_key(row_token)
    for row in range(start, max_row + 1):
        values = [_norm(ws.cell(row, col).value) for col in range(min_col, max_col + 1)]
        if any(token in _norm_key(value) for value in values):
            return row, values
    raise AssertionError(f"{ws.title}: could not find row containing {row_token!r}")


def _find_sheet_row_any(ws, row_tokens, table):
    for token in row_tokens:
        try:
            return _find_sheet_row(ws, token, table)
        except AssertionError:
            continue
    raise AssertionError(f"{ws.title}: could not find row containing any of {row_tokens!r}")


def _cell_by_header(ws, table, row, header):
    mapping = _table_header_map(ws, table)
    key = _norm_key(header)
    assert key in mapping, f"{ws.title}: missing header {header!r}"
    return ws.cell(row, mapping[key])


def _formula_has_tokens(formula, tokens, label):
    normalized = _norm_formula(formula)
    assert normalized.startswith("="), f"{label}: expected formula, found {formula!r}"
    missing = [token for token in tokens if _norm_formula(token).lstrip("=") not in normalized]
    assert not missing, f"{label}: formula missing tokens {missing!r}: {formula!r}"


def _formula_has_any(formula, tokens, label):
    normalized = _norm_formula(formula)
    assert normalized.startswith("="), f"{label}: expected formula, found {formula!r}"
    assert any(_norm_formula(token).lstrip("=") in normalized for token in tokens), f"{label}: formula missing any of {tokens!r}: {formula!r}"


def _clean_area(value):
    text = _norm(value).replace("'", "").replace("$", "").replace(" ", "")
    if "!" in text:
        text = text.split("!", 1)[1]
    return text.upper()


def _range_covered(ws, expected):
    expected = expected.replace("$", "")
    min_col, min_row, max_col, max_row = range_boundaries(expected)
    target = {f"{get_column_letter(col)}{row}" for row in range(min_row, max_row + 1) for col in range(min_col, max_col + 1)}
    cells = set()
    for dv in ws.data_validations.dataValidation:
        for rng in dv.cells.ranges:
            min_c, min_r, max_c, max_r = range_boundaries(str(rng).replace("$", ""))
            cells.update(f"{get_column_letter(col)}{row}" for row in range(min_r, max_r + 1) for col in range(min_c, max_c + 1))
    return target.issubset(cells)


def _sheet_text(wb, sheets):
    parts = []
    for sheet in sheets:
        ws = _sheet(wb, sheet)
        for row in ws.iter_rows():
            for cell in row:
                if cell.value is not None:
                    parts.append(str(cell.value))
    return "\n".join(parts)


def _defined_name_names(wb):
    names = set()
    defined_names = wb.defined_names
    if hasattr(defined_names, "values"):
        names.update(dn.name for dn in defined_names.values())
    else:
        names.update(dn.name for dn in defined_names.definedName)
    return names


def _assert_row_contains(row_values, tokens, label):
    text = " | ".join(row_values)
    missing = [token for token in tokens if _norm_key(token) not in _norm_key(text)]
    assert not missing, f"{label}: missing {missing!r} in {row_values!r}"


def _has_row_containing(rows, tokens):
    for row in rows:
        text = " | ".join(row)
        if all(_norm_key(token) in _norm_key(text) for token in tokens):
            return True
    return False


def _table_data_rows(ws, table):
    min_col, min_row, max_col, max_row = _table_bounds(table)
    return [
        [_norm(ws.cell(row, col).value) for col in range(min_col, max_col + 1)]
        for row in range(min_row + 1, max_row + 1)
    ]


def _formula_or_value_mentions(value, tokens, label):
    text = _norm(value)
    missing = [token for token in tokens if _norm_key(token) not in _norm_key(text)]
    assert not missing, f"{label}: missing {missing!r}: {value!r}"


def test_outputs_exist():
    assert _path("docx").exists(), f"Missing DOCX output at {_path('docx')}"
    assert _path("xlsx").exists(), f"Missing XLSX output at {_path('xlsx')}"
    assert _path("docx").suffix.lower() == ".docx"
    assert _path("xlsx").suffix.lower() == ".xlsx"


def test_memo_structure_content_privacy():
    doc = Document(_path("docx"))
    assert _heading_sequence_matches(_heading_order(doc), EXPECT["doc_heading_order"])
    text = _doc_text(doc)
    for anchor in [
        "CP2100A", "2026-06-15", "first B-notice workflow", "24%",
        "Orion Creative LLC", "88-4519021", "Mesa Field Services",
        "2026-07-12", "$2,016.00", "$6,384.00", "Brightline Legal PLLC",
        "27-6401180", "2026-06-30",
    ]:
        assert anchor in text
    norm_text = normalize_text(text)
    assert "risk" in norm_text and "excluded" in norm_text and "public" in norm_text, (
        "action memo: missing privacy-boundary statement for internal AP risk notes"
    )
    forbid_any(text, PUBLIC_FORBIDDEN, "action memo")
    header, footer = _header_footer_text(doc)
    assert EXPECT["doc_header"] in header
    assert EXPECT["doc_footer"] in footer
    assert _has_decision_table(doc)


def test_workbook_structure_formulas_controls():
    wb = load_workbook(_path("xlsx"), data_only=False)
    assert wb.sheetnames == EXPECT["sheet_order"]
    min_table_rows = {
        "Action Summary": 12,
        "Vendor Actions": 4,
        "Payment Decisions": 5,
        "Communications Log": 5,
        "TIN Corrections": 4,
        "Privacy Review": 4,
        "Raw IRS Notice": 4,
        "Rules": 3,
    }
    for sheet, (table_name, ref) in EXPECT["tables"].items():
        ws = _sheet(wb, sheet)
        expected_min_cols = range_boundaries(ref)[2] - range_boundaries(ref)[0] + 1
        expected_min_rows = min_table_rows[sheet]
        _table_has_size(ws, table_name, expected_min_rows, expected_min_cols)

    expected_headers = {
        "Action Summary": ["Field", "Value", "Source", "Review"],
        "Vendor Actions": ["Vendor ID", "Vendor", "Notice Reason", "Current Master TIN", "Corrected TIN", "W-9 Status", "Action", "Public Note"],
        "Payment Decisions": ["Vendor ID", "Vendor", "Invoice", "Gross Reportable Payment", "Backup Withholding", "Net Payable", "Decision"],
        "Communications Log": ["Item", "Owner", "Due Date", "Status", "Evidence"],
        "TIN Corrections": ["Vendor ID", "Old Name/TIN", "Corrected Name/TIN", "1099 Box", "Status"],
        "Privacy Review": ["Check", "Rule", "Result", "Owner"],
        "Raw IRS Notice": ["Vendor ID", "Notice Vendor Name", "Notice TIN", "Reason", "Notice Date"],
        "Rules": ["Rule", "Value"],
    }
    for sheet, headers in expected_headers.items():
        ws = _sheet(wb, sheet)
        table = _table(ws, EXPECT["tables"][sheet][0])
        _require_headers(ws, table, headers)

    for sheet in EXPECT["hidden_sheets"]:
        assert _sheet(wb, sheet).sheet_state in ("hidden", "veryHidden")
    names = _defined_name_names(wb)
    for name in EXPECT["defined_names"]:
        assert name in names

    summary = _sheet(wb, "Action Summary")
    summary_table = _table(summary, "tblActionSummary")
    for label, token_options in {
        "Total reportable July payments": [["=", "Payment Decisions"], ["=", "Payment_Decisions_Gross"]],
        "Total backup withholding": [["=", "Payment Decisions"], ["=", "TotalBackupWithholding"], ["=", "Payment_Decisions_BW"]],
        "Net payable after withholding": [["=", "Payment Decisions"], ["=", "MesaNetPayable"], ["=", "Payment_Decisions_Net"]],
        "Packet ready": [["=", "IF"]],
    }.items():
        row, _ = _find_sheet_row(summary, label, summary_table)
        formula = _cell_by_header(summary, summary_table, row, "Value").value
        for tokens in token_options:
            try:
                _formula_has_tokens(formula, tokens, f"Action Summary {label}")
                break
            except AssertionError:
                continue
        else:
            raise AssertionError(f"Action Summary {label}: formula missing accepted tokens: {formula!r}")

    payments = _sheet(wb, "Payment Decisions")
    payment_table = _table(payments, "tblPaymentDecisions")
    for vendor_id in ["V-1042", "V-2198", "V-3301", "V-4410"]:
        row, _ = _find_sheet_row(payments, vendor_id, payment_table)
        backup = _cell_by_header(payments, payment_table, row, "Backup Withholding").value
        net = _cell_by_header(payments, payment_table, row, "Net Payable").value
        assert backup is not None, f"{vendor_id} backup withholding missing"
        assert net is not None, f"{vendor_id} net payable missing"
    row, _ = _find_sheet_row(payments, "V-2198", payment_table)
    _formula_has_tokens(_cell_by_header(payments, payment_table, row, "Backup Withholding").value, ["="], "Mesa backup withholding")
    _formula_has_tokens(_cell_by_header(payments, payment_table, row, "Net Payable").value, ["="], "Mesa net payable")
    _formula_has_any(_cell_by_header(payments, payment_table, row, "Backup Withholding").value, ["0.24", "BackupWithholdingRate", "Rules"], "Mesa backup withholding")
    try:
        row, _ = _find_sheet_row(payments, "TOTAL", payment_table)
    except AssertionError:
        row = None
    if row is not None:
        for header in ["Gross Reportable Payment", "Backup Withholding", "Net Payable"]:
            _formula_has_tokens(_cell_by_header(payments, payment_table, row, header).value, ["SUM"], f"TOTAL {header}")

    for sheet, ranges in EXPECT["data_validation"].items():
        ws = _sheet(wb, sheet)
        for rng in ranges:
            assert _range_covered(ws, rng), f"{sheet}: missing validation over {rng}"
    for sheet, area in EXPECT["print_areas"].items():
        actual = _clean_area(_sheet(wb, sheet).print_area)
        table = _table(_sheet(wb, sheet), EXPECT["tables"][sheet][0])
        target = _clean_area(table.ref)
        assert target in actual or actual in target


def test_workbook_values_privacy():
    wb = load_workbook(_path("xlsx"), data_only=False)

    summary = _sheet(wb, "Action Summary")
    summary_table = _table(summary, "tblActionSummary")
    for label, tokens in {
        "IRS notice": ["CP2100A"],
        "Notice date": ["2026-06-15"],
        "B-notice packet due": ["2026-06-30"],
        "Vendors on notice": ["3"],
        "Validated corrected W-9s": ["2"],
        "Vendors requiring July withholding": ["1"],
    }.items():
        if label == "Validated corrected W-9s":
            _, row_values = _find_sheet_row_any(summary, ["Validated corrected W-9s", "Validated corrected W-9 count"], summary_table)
        else:
            _, row_values = _find_sheet_row(summary, label, summary_table)
        _assert_row_contains(row_values, tokens, f"Action Summary {label}")
    row, row_values = _find_sheet_row(summary, "Backup withholding rate", summary_table)
    rate_value = _cell_by_header(summary, summary_table, row, "Value").value
    assert (
        _norm_key("0.24") in _norm_key(" | ".join(row_values))
        or _norm_key("24%") in _norm_key(" | ".join(row_values))
        or (isinstance(rate_value, str) and rate_value.startswith("=") and any(token in _norm_formula(rate_value) for token in ["RULES", "BACKUPWITHHOLDINGRATE"]))
    ), f"Action Summary Backup withholding rate: missing rate or rules reference in {row_values!r}"

    vendors = _sheet(wb, "Vendor Actions")
    vendor_table = _table(vendors, "tblVendorActions")
    for vendor_id, tokens in {
        "V-1042": ["Orion Creative LLC", "88-4519020", "88-4519021", "W-9", "no withholding"],
        "V-2198": ["Mesa Field Services", "Missing TIN", "No valid W-9", "Apply 24% backup withholding"],
        "V-3301": ["Brightline Legal PLLC", "27-6401188", "27-6401180", "W-9", "no withholding"],
    }.items():
        _, row_values = _find_sheet_row(vendors, vendor_id, vendor_table)
        _assert_row_contains(row_values, tokens, f"Vendor Actions {vendor_id}")

    payments = _sheet(wb, "Payment Decisions")
    payment_table = _table(payments, "tblPaymentDecisions")
    for vendor_id, tokens in {
        "V-1042": ["Orion Creative LLC", "INV-6142", "12500", "withholding"],
        "V-2198": ["Mesa Field Services", "INV-7781", "8400", "24", "withholding"],
        "V-3301": ["Brightline Legal PLLC", "INV-9034", "6100", "withholding"],
        "V-4410": ["Northstar Supply Corp", "INV-5530", "2700"],
    }.items():
        _, row_values = _find_sheet_row(payments, vendor_id, payment_table)
        _assert_row_contains(row_values, tokens, f"Payment Decisions {vendor_id}")
    _, northstar = _find_sheet_row(payments, "V-4410", payment_table)
    assert any(_norm_key(token) in _norm_key(" | ".join(northstar)) for token in ["Corporation", "exempt", "No withholding"]), (
        f"Payment Decisions V-4410: missing corporation/exempt/no-withholding decision in {northstar!r}"
    )

    comm = _sheet(wb, "Communications Log")
    comm_table = _table(comm, "tblCommunicationsLog")
    comm_rows = _table_data_rows(comm, comm_table)
    assert _has_row_containing(comm_rows, ["B-notice", "2026-06-30"]), "Communications Log: missing first B-notice packet action"
    assert _has_row_containing(comm_rows, ["Orion", "88-4519021"]) or _has_row_containing(comm_rows, ["Orion", "W-9"]), "Communications Log: missing Orion correction action"
    assert _has_row_containing(comm_rows, ["Brightline", "27-6401180"]) or _has_row_containing(comm_rows, ["Brightline", "W-9"]), "Communications Log: missing Brightline correction action"
    assert _has_row_containing(comm_rows, ["Mesa", "withholding"]), "Communications Log: missing Mesa withholding action"
    assert _has_row_containing(comm_rows, ["Remit", "withholding"]) or _has_row_containing(comm_rows, ["withholding", "2026-07-31"]), "Communications Log: missing withholding remittance follow-up"

    tin = _sheet(wb, "TIN Corrections")
    tin_table = _table(tin, "tblTinCorrections")
    for vendor_id, tokens in {
        "V-1042": ["88-4519020", "88-4519021"],
        "V-2198": ["Missing", "Pending"],
        "V-3301": ["27-6401188", "27-6401180"],
    }.items():
        _, row_values = _find_sheet_row(tin, vendor_id, tin_table)
        _assert_row_contains(row_values, tokens, f"TIN Corrections {vendor_id}")

    privacy = _sheet(wb, "Privacy Review")
    privacy_table = _table(privacy, "tblPrivacyReview")
    privacy_rows = _table_data_rows(privacy, privacy_table)
    assert _has_row_containing(privacy_rows, ["risk", "Pass"]), "Privacy Review: missing internal-risk pass check"
    assert _has_row_containing(privacy_rows, ["public", "Pass"]) or _has_row_containing(privacy_rows, ["private", "Pass"]), "Privacy Review: missing public/private boundary pass check"
    assert _has_row_containing(privacy_rows, ["vendor", "Pass"]), "Privacy Review: missing vendor-communications pass check"

    raw = _sheet(wb, "Raw IRS Notice")
    raw_table = _table(raw, "tblRawIRSNotice")
    for vendor_id in ["V-1042", "V-2198", "V-3301"]:
        _find_sheet_row(raw, vendor_id, raw_table)
    rules = _sheet(wb, "Rules")
    rules_table = _table(rules, "tblRules")
    rules_rows = _table_data_rows(rules, rules_table)
    assert _has_row_containing(rules_rows, ["withholding", "0.24"]) or _has_row_containing(rules_rows, ["withholding", "24"]), "Rules: missing backup withholding rate"
    assert _has_row_containing(rules_rows, ["B-notice"]) or _has_row_containing(rules_rows, ["Notice", "2026-06-30"]), "Rules: missing first B-notice rule context"
    assert _has_row_containing(rules_rows, ["W-9"]) or _has_row_containing(rules_rows, ["TIN"]), "Rules: missing corrected W-9/TIN rule context"

    public = _sheet_text(wb, ["Action Summary", "Vendor Actions", "Payment Decisions", "Communications Log", "TIN Corrections", "Privacy Review"])
    forbid_any(public, PUBLIC_FORBIDDEN, "public workbook sheets")


def test_cross_output_consistency():
    doc_text = _doc_text(Document(_path("docx")))
    wb = load_workbook(_path("xlsx"), data_only=False)
    for anchor in ["CP2100A", "2026-06-15", "Orion Creative LLC", "Mesa Field Services", "Brightline Legal PLLC", "$2,016.00", "$6,384.00", "2026-06-30"]:
        assert anchor in doc_text
    public_workbook_text = _sheet_text(wb, ["Action Summary", "Vendor Actions", "Payment Decisions", "Communications Log", "TIN Corrections"])
    for anchor in ["CP2100A", "2026-06-15", "Orion Creative LLC", "Mesa Field Services", "Brightline Legal PLLC", "2026-06-30"]:
        assert anchor in public_workbook_text
    payments = _sheet(wb, "Payment Decisions")
    payment_table = _table(payments, "tblPaymentDecisions")
    row, _ = _find_sheet_row(payments, "V-2198", payment_table)
    _formula_has_any(_cell_by_header(payments, payment_table, row, "Backup Withholding").value, ["0.24", "BackupWithholdingRate", "Rules"], "Mesa withholding formula")
    _formula_has_tokens(_cell_by_header(payments, payment_table, row, "Net Payable").value, ["="], "Mesa net payable formula")
    vendors = _sheet(wb, "Vendor Actions")
    vendor_table = _table(vendors, "tblVendorActions")
    row, _ = _find_sheet_row(vendors, "V-2198", vendor_table)
    assert "withholding" in _norm_key(_cell_by_header(vendors, vendor_table, row, "Action").value)
    forbid_any(doc_text, PUBLIC_FORBIDDEN, "action memo")
