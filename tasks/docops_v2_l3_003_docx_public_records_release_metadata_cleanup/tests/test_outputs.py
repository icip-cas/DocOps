import json
import os
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt

from verifier_utils import normalize_text, run_preflight

META_PATH = Path(os.environ.get("TASK_METADATA_PATH", "/tests/task_metadata.json"))
if not META_PATH.exists():
    META_PATH = Path(__file__).parent / "task_metadata.json"
META = json.loads(META_PATH.read_text(encoding="utf-8"))
EXPECT = META["verifier_expectations"]
INPUT_PATH = Path(os.environ.get("INPUT_PATH", META["input_path"]))
OUTPUT_PATH = Path(os.environ.get("OUTPUT_PATH", META["output_path"]))


def doc_text(doc):
    parts = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text.strip())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            if paragraph.text.strip():
                parts.append(paragraph.text.strip())
        for paragraph in section.footer.paragraphs:
            if paragraph.text.strip():
                parts.append(paragraph.text.strip())
    return "\n".join(parts)


def package_parts(path):
    with zipfile.ZipFile(path) as zf:
        return set(zf.namelist())


def package_xml_text(path):
    chunks = []
    with zipfile.ZipFile(path) as zf:
        for name in zf.namelist():
            if name.endswith(".xml") or name.endswith(".rels"):
                chunks.append(zf.read(name).decode("utf-8", errors="ignore"))
    return "\n".join(chunks)


def visible_heading_sequence(doc):
    wanted = set(EXPECT["headings"])
    return [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip() in wanted]


def first_run(paragraph):
    for run in paragraph.runs:
        if run.text.strip():
            return run
    return paragraph.runs[0] if paragraph.runs else None


def paragraph_by_text(doc, text):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise AssertionError(f"Missing paragraph: {text}")


def all_table_rows(doc):
    rows = []
    for table in doc.tables:
        table_rows = []
        for row in table.rows:
            table_rows.append([cell.text.strip() for cell in row.cells])
        rows.append(table_rows)
    return rows


def cell_fill(cell):
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    shd = cell._tc.find(".//w:shd", namespaces=ns)
    if shd is None:
        return None
    return shd.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill")


def has_page_break(paragraph):
    xml = paragraph._p.xml
    return 'w:type="page"' in xml or "w:type='page'" in xml


def assert_break_before(doc, heading):
    previous = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == heading:
            assert previous is not None, f"{heading}: no paragraph before heading"
            assert has_page_break(previous) or has_page_break(paragraph), f"{heading}: missing page break before heading"
            return
        previous = paragraph
    raise AssertionError(f"Heading not found for page break: {heading}")


def assert_table_header_style(doc):
    style = EXPECT["style"]
    for table in doc.tables:
        for cell in table.rows[0].cells:
            assert cell_fill(cell) == style["table_header_fill"]
            runs = [run for paragraph in cell.paragraphs for run in paragraph.runs if run.text.strip()]
            assert runs, "Header cell has no run"
            assert all(run.font.bold for run in runs)
            assert all(str(run.font.color.rgb) == style["table_header_font"] for run in runs)


def assert_public_content(text):
    normalized = normalize_text(text)
    required_terms = [
        "2026-01-18",
        "37",
        "$650",
        "Public Records Desk",
        "2026-08-15",
    ]
    missing = [term for term in required_terms if normalize_text(term) not in normalized]
    assert not missing, f"Missing corrected public facts: {missing}"
    concept_groups = {
        "public release packet": ["public release packet", "public records release"],
        "redaction/privacy cleanup": ["claimant identifiers", "identifiers", "pii", "redaction"],
        "appendix A": ["Appendix A"],
        "appendix B": ["Appendix B"],
        "appendix C": ["Appendix C"],
    }
    missing_concepts = [
        label
        for label, options in concept_groups.items()
        if not any(normalize_text(option) in normalized for option in options)
    ]
    assert not missing_concepts, f"Missing required public concepts: {missing_concepts}"


def assert_public_tables(doc):
    tables = all_table_rows(doc)
    assert len(tables) >= 3, f"Expected at least three rebuilt public tables, found {len(tables)}"
    table_text = normalize_text(
        "\n".join("\n".join("\t".join(row) for row in table) for table in tables)
    )
    fact_groups = {
        "incident date": ["incident date"],
        "claim count": ["claim count", "claims reviewed", "total public claims"],
        "reimbursable cap": ["reimbursable cap", "$650"],
        "public contact": ["public contact", "Public Records Desk"],
        "response deadline": ["response deadline", "2026-08-15"],
    }
    for label, options in fact_groups.items():
        assert any(normalize_text(option) in table_text for option in options), (
            f"Corrected Determination Facts table missing {label}"
        )
    category_groups = {
        "claim category": ["claim category", "cleanup", "property", "temporary"],
        "release handling": ["release treatment", "public determination", "release value", "aggregate"],
        "cap or reimbursement": ["$650", "reimbursement", "reimbursable"],
    }
    for label, options in category_groups.items():
        assert any(normalize_text(option) in table_text for option in options), f"Claim Category Table missing {label}"
    checklist_groups = {
        "tracked changes": ["tracked changes", "tracked revisions", "deleted text", "markup"],
        "comments": ["comments", "comment references"],
        "document properties": ["document properties", "metadata"],
        "completion status": ["complete", "completed", "removed", "cleaned"],
    }
    for label, options in checklist_groups.items():
        assert any(normalize_text(option) in table_text for option in options), (
            f"Redaction and Metadata Checklist missing {label}"
        )


def effective_bold(paragraph, run):
    if run.font.bold is not None:
        return bool(run.font.bold)
    if paragraph.style is not None and paragraph.style.font.bold is not None:
        return bool(paragraph.style.font.bold)
    return False


def test_output_exists_and_is_docx():
    run_preflight(META, INPUT_PATH, OUTPUT_PATH)
    assert OUTPUT_PATH != INPUT_PATH


def test_headings_required_text_tables_and_privacy_cleanup():
    doc = Document(OUTPUT_PATH)
    assert visible_heading_sequence(doc) == EXPECT["headings"]
    text = doc_text(doc)
    assert_public_content(text)
    hits = [phrase for phrase in EXPECT["forbidden_phrases"] if normalize_text(phrase) in normalize_text(text)]
    assert not hits, f"Forbidden visible phrases still present: {hits}"
    assert_public_tables(doc)


def test_ooxml_comments_tracked_changes_and_hidden_text_removed():
    parts = package_parts(OUTPUT_PATH)
    for forbidden_part in EXPECT["forbidden_package_parts"]:
        assert forbidden_part not in parts
    raw = package_xml_text(OUTPUT_PATH)
    hidden_hits = [phrase for phrase in EXPECT["forbidden_phrases"] if normalize_text(phrase) in normalize_text(raw)]
    assert not hidden_hits, f"Forbidden phrases remain in package XML: {hidden_hits}"
    marker_hits = [marker for marker in EXPECT["forbidden_ooxml_markers"] if marker in raw]
    assert not marker_hits, f"Tracked-change/comment OOXML markers remain: {marker_hits}"


def test_public_metadata_applied():
    doc = Document(OUTPUT_PATH)
    core = doc.core_properties
    assert core.author == EXPECT["metadata"]["author"]
    assert core.subject == EXPECT["metadata"]["subject"]
    assert core.title == EXPECT["metadata"]["title"]
    metadata_text = "\n".join(str(value or "") for value in [core.author, core.subject, core.title, core.keywords, core.comments, core.last_modified_by])
    hits = [phrase for phrase in EXPECT["forbidden_phrases"] if normalize_text(phrase) in normalize_text(metadata_text)]
    assert not hits, f"Forbidden metadata phrases remain: {hits}"


def test_style_layout_footer_and_page_breaks():
    doc = Document(OUTPUT_PATH)
    style = EXPECT["style"]
    title = paragraph_by_text(doc, EXPECT["headings"][0])
    title_run = first_run(title)
    assert title_run.font.name == style["title_font"]
    assert title_run.font.bold
    assert title_run.font.size == Pt(style["title_size_pt"])
    assert str(title_run.font.color.rgb) == style["title_color"]
    for heading in EXPECT["headings"][1:]:
        paragraph = paragraph_by_text(doc, heading)
        run = first_run(paragraph)
        assert str(run.font.color.rgb) == style["heading_color"]
        assert effective_bold(paragraph, run)
    for section in doc.sections:
        assert abs(section.top_margin.inches - style["margin_inches"]) < 0.02
        assert abs(section.bottom_margin.inches - style["margin_inches"]) < 0.02
        assert abs(section.left_margin.inches - style["margin_inches"]) < 0.02
        assert abs(section.right_margin.inches - style["margin_inches"]) < 0.02
        footer_text = "\n".join(p.text for p in section.footer.paragraphs)
        assert style["footer_text"] in footer_text
    assert_table_header_style(doc)
    for heading in EXPECT["page_break_before"]:
        assert_break_before(doc, heading)


def test_source_artifact_was_not_modified():
    raw = package_xml_text(INPUT_PATH)
    parts = package_parts(INPUT_PATH)
    for phrase in EXPECT["source_must_contain"]:
        if phrase.startswith("word/"):
            assert phrase in parts
        else:
            assert phrase in raw, f"Source no longer contains expected marker: {phrase}"
