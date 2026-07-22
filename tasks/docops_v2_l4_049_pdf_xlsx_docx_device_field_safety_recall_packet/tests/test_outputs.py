import json
import os
import re
import sys
import zipfile
from pathlib import Path
from datetime import date, datetime

from docx import Document
from openpyxl import load_workbook
from openpyxl.utils.cell import get_column_letter, range_boundaries
from pypdf import PdfReader

sys.path.insert(0, str(Path(__file__).parent))
from verifier_utils import *  # noqa: F401,F403

META_PATH = Path(os.environ.get('TASK_METADATA_PATH', '/tests/task_metadata.json'))
META = json.loads(META_PATH.read_text(encoding='utf-8'))
EXPECT = META['verifier_expectations']


def _vopt_norm_text(value):
    if value is None:
        return ''
    if isinstance(value, datetime):
        if value.hour or value.minute or value.second:
            return value.strftime('%Y-%m-%d %H:%M')
        return value.strftime('%Y-%m-%d')
    if isinstance(value, date):
        return value.strftime('%Y-%m-%d')
    if isinstance(value, float) and value.is_integer():
        value = int(value)
    text = str(value)
    text = text.replace('\u2013', '-').replace('\u2014', '-').replace('\u2212', '-')
    text = text.replace('\xa0', ' ')
    text = re.sub(r'\s+', ' ', text).strip()
    return text


def _vopt_norm_formula(value):
    return re.sub(r'\s+', '', _vopt_norm_text(value)).upper()


def _vopt_assert_formula(actual, expected, label='formula'):
    assert _vopt_norm_formula(actual) == _vopt_norm_formula(expected), f"{label}: expected {expected!r}, found {actual!r}"


def _vopt_norm_rows(rows):
    return [[_vopt_norm_text(cell) for cell in row] for row in rows]


def _vopt_rows_equal(actual, expected):
    return _vopt_norm_rows(actual) == _vopt_norm_rows(expected)


def _vopt_rows_text(rows):
    return '\n'.join('|'.join(_vopt_norm_text(cell) for cell in row) for row in rows)


def _vopt_norm_key(value):
    return re.sub(r'[^a-z0-9]+', '', _vopt_norm_text(value).lower())


def _vopt_header_has(header, required):
    actual = {_vopt_norm_text(cell).lower() for cell in header}
    missing = [cell for cell in required if _vopt_norm_text(cell).lower() not in actual]
    assert not missing, f"Missing expected header columns {missing!r}; actual={header!r}"


def _vopt_section_text(section_part):
    return '\n'.join(p.text for p in getattr(section_part, 'paragraphs', []) if p.text is not None)


def _vopt_header_text(doc):
    return '\n'.join(_vopt_section_text(section.header) for section in doc.sections)


def _vopt_footer_text(doc):
    return '\n'.join(_vopt_section_text(section.footer) for section in doc.sections)


def _vopt_table_ref(ws, expected_name=None):
    tables = ws.tables
    if expected_name:
        for name in tables.keys():
            if str(name).lower() == str(expected_name).lower():
                return tables[name].ref
    vals = list(tables.values())
    assert vals, f"No Excel native table found on sheet {ws.title!r}"
    return vals[0].ref


def _vopt_clean_area(value):
    text = _vopt_norm_text(value).replace("'", '').replace('$', '').replace(' ', '')
    if '!' in text:
        text = text.split('!', 1)[1]
    return text.upper()


def _vopt_assert_print_area(ws, expected):
    actual = _vopt_clean_area(ws.print_area)
    target = _vopt_clean_area(expected)
    assert target in actual or actual in target, f"{ws.title}: expected print area {expected!r}, found {ws.print_area!r}"


def _vopt_freeze_pane(value):
    return getattr(value, 'coordinate', value)


def _vopt_validated_cells(ws):
    cells = set()
    for dv in ws.data_validations.dataValidation:
        for rng in dv.cells.ranges:
            text = str(rng).replace('$', '')
            try:
                min_col, min_row, max_col, max_row = range_boundaries(text)
            except ValueError:
                cells.add(text)
                continue
            if max_row > 10000:
                max_row = max(min_row, 200)
            for row in range(min_row, max_row + 1):
                for col in range(min_col, max_col + 1):
                    cells.add(f"{get_column_letter(col)}{row}")
    return cells


class _VoptValidationRanges(list):
    def __init__(self, ws, ranges):
        super().__init__(ranges)
        self.ws = ws

    def __contains__(self, expected):
        expected = str(expected).replace('$', '')
        if super().__contains__(expected):
            return True
        try:
            min_col, min_row, max_col, max_row = range_boundaries(expected)
        except ValueError:
            return super().__contains__(expected)
        target = {f"{get_column_letter(col)}{row}" for row in range(min_row, max_row + 1) for col in range(min_col, max_col + 1)}
        return target.issubset(_vopt_validated_cells(self.ws))


def _vopt_dv_ranges(ws):
    ranges = []
    for dv in ws.data_validations.dataValidation:
        ranges.extend(str(rng).replace('$', '') for rng in dv.cells.ranges)
    return _VoptValidationRanges(ws, ranges)



def _path(kind):
    env = {
        'pdf': 'PDF_OUTPUT_PATH',
        'xlsx': 'XLSX_OUTPUT_PATH',
        'docx': 'DOCX_OUTPUT_PATH',
    }[kind]
    key = {
        'pdf': 'pdf_output',
        'xlsx': 'xlsx_output',
        'docx': 'docx_output',
    }[kind]
    return Path(os.environ.get(env, EXPECT[key]))


def _pdf_text(path):
    reader = PdfReader(str(path))
    return '\n'.join(page.extract_text() or '' for page in reader.pages)


def _outline_titles(path):
    return [title for _, title in flatten_outline(PdfReader(str(path)).outline)]


def _outline_matches(actual, expected):
    actual_norm = [_vopt_norm_key(item) for item in actual]
    expected_norm = [_vopt_norm_key(item) for item in expected]
    aliases = {
        '1fieldsafetynoticesummary': {'summary', 'noticesummary', 'fieldsafetynoticesummary'},
        '2affectedproducts': {'affectedproducts'},
        '3requiredcustomeractions': {'customeractions', 'requiredcustomeractions'},
        '4responseformandcontacts': {'responseformandcontact', 'responseformandcontacts'},
    }
    if len(actual_norm) != len(expected_norm):
        return False
    for actual_item, expected_item in zip(actual_norm, expected_norm):
        allowed = aliases.get(expected_item, {expected_item})
        if actual_item not in allowed and expected_item not in actual_item and actual_item not in expected_item:
            return False
    return True


def _doc_text(doc):
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append('|'.join(cell.text.strip() for cell in row.cells))
    return '\n'.join(parts)


def _doc_heading_order(doc):
    out = []
    for p in doc.paragraphs:
        style = p.style.name if p.style else ''
        if p.text.strip() and (style.startswith('Heading') or style == 'Title'):
            out.append(p.text.strip())
    return out


def _norm_heading_label(value):
    text = str(value or '')
    text = text.replace('\u2013', '-').replace('\u2014', '-').replace('\xa0', ' ')
    text = re.sub(r'\s+', ' ', text).strip()
    text = re.sub(r'^(?:\d+|[IVXLCDM]+)[\.)]\s+', '', text, flags=re.I)
    text = re.sub(r'^[A-Z][\.)]\s+', '', text)
    text = re.sub(r'[:\-]+$', '', text).strip()
    return text.casefold()


def _heading_sequence_matches(actual, expected):
    actual_norm = [_norm_heading_label(item) for item in actual]
    expected_norm = [_norm_heading_label(item) for item in expected]
    pos = 0
    for item in actual_norm:
        if pos < len(expected_norm) and item == expected_norm[pos]:
            pos += 1
    return pos == len(expected_norm)


def _doc_table_rows(table):
    return [[cell.text.strip() for cell in row.cells] for row in table.rows]


def _has_toc_field(path):
    with zipfile.ZipFile(path) as zf:
        xml = zf.read('word/document.xml').decode('utf-8', errors='ignore')
    return bool(re.search(r'TOC\s+(?:\\)?o|TOC\s*(?:&quot;|")', xml))


def _highlighted_para(doc, text):
    for p in doc.paragraphs:
        if _vopt_norm_text(text) in _vopt_norm_text(p.text):
            return docx_para_has_highlight(p)
    raise AssertionError(f'Missing paragraph: {text}')


def _row_values(ws, row, max_col):
    return [ws.cell(row, col).value for col in range(1, max_col + 1)]


def _dv_ranges(ws):
    return _vopt_dv_ranges(ws)


def _table_ref_has_size(ref, min_rows, min_cols):
    min_col, min_row, max_col, max_row = range_boundaries(ref)
    return (max_row - min_row + 1) >= min_rows and (max_col - min_col + 1) >= min_cols


def _assert_formula_contains(ws, cell, tokens, context):
    formula = _vopt_norm_formula(ws[cell].value)
    assert formula.startswith('='), f'{context}: expected live formula, found {ws[cell].value!r}'
    missing = [token for token in tokens if token.upper() not in formula]
    assert not missing, f'{context}: formula missing {missing!r}; found {ws[cell].value!r}'


def _all_formula_text(ws):
    return '\n'.join(
        _vopt_norm_formula(cell.value)
        for row in ws.iter_rows()
        for cell in row
        if isinstance(cell.value, str) and cell.value.startswith('=')
    )


def _worksheet_text(ws):
    return '\n'.join(
        _vopt_norm_text(cell.value)
        for row in ws.iter_rows()
        for cell in row
        if cell.value is not None
    )


def _assert_formula_semantics(lot, resp, summary):
    lot_formulas = _all_formula_text(lot)
    resp_formulas = _all_formula_text(resp)
    summary_formulas = _all_formula_text(summary)
    assert 'F' in lot_formulas and ('PUBLIC' in lot_formulas or 'NOTICE' in lot_formulas or 'INNOTICE' in lot_formulas), (
        'Lot Reconciliation must classify public notice rows from Publish values'
    )
    assert 'INTERNAL' in lot_formulas or 'RESTRICTED' in lot_formulas, (
        'Lot Reconciliation must flag restricted/internal rows'
    )
    assert 'COUNTIF' in lot_formulas and 'CUSTOMERRESPONSES' in lot_formulas, (
        'Lot Reconciliation must count customer responses by lot/action'
    )
    assert any(token in lot_formulas for token in ['FOLLOWUP', 'PENDINGRESPONSE', 'TRAININGSCHEDULED']) or any(
        token in resp_formulas for token in ['FOLLOWUP', 'PENDINGRESPONSE', 'TRAININGSCHEDULED']
    ), 'Workbook must identify follow-up needs with formulas'
    assert 'LOTRECONCILIATION' in summary_formulas and 'CUSTOMERRESPONSES' in summary_formulas, (
        'Response Summary must use formulas tied to reconciliation and response sheets'
    )


def _assert_customer_rows_retained(resp):
    text = _worksheet_text(resp)
    for row in EXPECT['customer_rows']:
        missing = [_vopt_norm_text(token) for token in row if _vopt_norm_text(token) not in text]
        assert not missing, f'Customer Responses missing source-row details {missing!r}'


def _has_validated_cells(ws, cells):
    return all(cell in _vopt_validated_cells(ws) for cell in cells)


def _assert_pdf_public_lot_semantics(text):
    for field_action_id, product, lot, region, _action, deadline in EXPECT['public_lots']:
        require_all(text, [field_action_id, product, lot, region, deadline], 'customer PDF public lot table')
    lower = _vopt_norm_text(text).lower()
    action_groups = [
        ['quarantine', 'firmware', '2.8'],
        ['corrected', 'ifu', 'training'],
        ['stop use', 'high', 'humidity'],
    ]
    for terms in action_groups:
        missing = [term for term in terms if term not in lower]
        assert not missing, f'customer PDF missing action semantics {terms!r}'


def _assert_cover_memo_semantics(text):
    for public_id in EXPECT['public_ids']:
        assert public_id in text, f'cover memo missing public row {public_id}'
    lower = _vopt_norm_text(text).lower()
    for terms in [
        ['customer field safety notice', 'workbook'],
        ['six', 'public'],
        ['validation', 'print'],
        ['pdf', 'workbook', 'fa-01'],
        ['internal', 'notes'],
    ]:
        assert all(term in lower for term in terms), f'cover memo missing semantic terms {terms!r}'


def _has_highlighted_followup(doc):
    for p in doc.paragraphs:
        text = _vopt_norm_text(p.text).lower()
        if not text:
            continue
        if docx_para_has_highlight(p) and ('follow' in text or 'outreach' in text) and (
            'pending' in text or 'training scheduled' in text or 'ch-101' in text
        ):
            return True
    return False


def test_outputs_exist():
    assert _path('pdf').exists(), f"Missing PDF output: {_path('pdf')}"
    assert _path('xlsx').exists(), f"Missing XLSX output: {_path('xlsx')}"
    assert _path('docx').exists(), f"Missing DOCX output: {_path('docx')}"
    assert _path('pdf').suffix.lower() == '.pdf'
    assert _path('xlsx').suffix.lower() == '.xlsx'
    assert _path('docx').suffix.lower() == '.docx'


def test_public_pdf_notice_scope_and_navigation():
    pdf = _path('pdf')
    reader = PdfReader(str(pdf))
    assert len(reader.pages) == 4, f"Expected 4 PDF pages, found {len(reader.pages)}"
    assert _outline_matches(_outline_titles(pdf), EXPECT['pdf_outline']), f"Unexpected PDF outline: {_outline_titles(pdf)!r}"
    text = _pdf_text(pdf)
    require_all(text, EXPECT['pdf_required_phrases'], 'customer PDF')
    _assert_pdf_public_lot_semantics(text)
    forbid_any(text, EXPECT['public_forbidden_text'], 'customer PDF public boundary')


def test_reconciliation_workbook_structure_controls_and_internal_retention():
    wb = load_workbook(_path('xlsx'), data_only=False)
    assert wb.sheetnames == EXPECT['workbook_sheet_order'], f"Unexpected sheets: {wb.sheetnames!r}"
    lot = wb['Lot Reconciliation']
    resp = wb['Customer Responses']
    summary = wb['Response Summary']
    lot_ref = _vopt_table_ref(lot, 'field_action_lot_reconciliation')
    resp_ref = _vopt_table_ref(resp, 'customer_response_matrix')
    assert _table_ref_has_size(lot_ref, 9, 8), f'Lot reconciliation table too small: {lot_ref}'
    assert _table_ref_has_size(resp_ref, 8, 6), f'Customer response table too small: {resp_ref}'
    freeze_panes = EXPECT.get('freeze_panes', {'Lot Reconciliation': 'A2', 'Customer Responses': 'A2'})
    assert _vopt_freeze_pane(lot.freeze_panes) == freeze_panes['Lot Reconciliation']
    assert _vopt_freeze_pane(resp.freeze_panes) == freeze_panes['Customer Responses']
    assert lot.print_area, 'Lot Reconciliation must have a print area'
    assert resp.print_area, 'Customer Responses must have a print area'
    assert summary.print_area, 'Response Summary must have a print area'
    for sheet in EXPECT['hidden_sheets']:
        assert wb[sheet].sheet_state in ('hidden', 'veryHidden'), f"{sheet} should be hidden"
    names = {dn.name for dn in wb.defined_names.values()}
    for name in EXPECT['defined_names']:
        assert name in names, f"Missing defined name: {name}"
    _assert_formula_semantics(lot, resp, summary)
    lot_public = [lot.cell(row, 1).value for row in range(2, 10) if lot.cell(row, 6).value == 'Yes']
    lot_internal = [lot.cell(row, 1).value for row in range(2, 10) if lot.cell(row, 6).value == 'No']
    assert lot_public == EXPECT['public_ids']
    assert lot_internal == EXPECT['internal_ids']
    internal_text = '\n'.join(str(cell.value) for ws in [wb['Internal Notes']] for row in ws.iter_rows() for cell in row if cell.value)
    require_all(internal_text, ['ASIC comparator drift', 'Attorney work product', 'FA-INT-07', 'FA-LEGAL-08'], 'internal workbook notes')
    assert len(lot.conditional_formatting) >= 1
    assert len(resp.conditional_formatting) >= 1
    assert _has_validated_cells(lot, ['F2', 'F9']), 'Lot Reconciliation must validate Publish values'
    assert _has_validated_cells(resp, ['D2', 'D8']) or _has_validated_cells(resp, ['B2', 'B8']) or _has_validated_cells(resp, ['C2', 'C8']), (
        'Customer Responses must validate response-status values'
    )
    _assert_customer_rows_retained(resp)


def test_cover_memo_structure_and_reconciliation_language():
    docx = _path('docx')
    doc = Document(docx)
    headings = _doc_heading_order(doc)
    if EXPECT['doc_heading_order'][0] not in headings and EXPECT['doc_heading_order'][0] in _doc_text(doc):
        headings = [EXPECT['doc_heading_order'][0]] + headings
    assert _heading_sequence_matches(headings, EXPECT['doc_heading_order']), f"Unexpected heading order: {_doc_heading_order(doc)!r}"
    text = _doc_text(doc)
    _assert_cover_memo_semantics(text)
    forbid_any(text, EXPECT['doc_forbidden_text'], 'cover memo')
    assert EXPECT['doc_header_contains'] in _vopt_header_text(doc)
    assert EXPECT['doc_footer_contains'] in _vopt_footer_text(doc)
    assert _has_toc_field(docx), 'Expected real Word TOC field'
    assert _has_highlighted_followup(doc), 'Expected highlighted high-risk follow-up paragraph'
    table_rows = _doc_table_rows(doc.tables[0])
    _vopt_header_has(table_rows[0], ['Field Action ID', 'Product', 'Lot'])
    assert [row[0] for row in table_rows[1:]] == EXPECT['public_ids']


def test_cross_output_public_internal_boundary_consistency():
    pdf_text = _pdf_text(_path('pdf'))
    doc_text = _doc_text(Document(_path('docx')))
    wb = load_workbook(_path('xlsx'), data_only=False)
    assert 'Lot Reconciliation' in wb.sheetnames, f"Missing sheet: Lot Reconciliation; actual={wb.sheetnames!r}"
    lot = wb['Lot Reconciliation']
    public_ids = [lot.cell(row, 1).value for row in range(2, 10) if lot.cell(row, 6).value == 'Yes']
    internal_ids = [lot.cell(row, 1).value for row in range(2, 10) if lot.cell(row, 6).value == 'No']
    assert public_ids == EXPECT['public_ids']
    assert internal_ids == EXPECT['internal_ids']
    for public_id in public_ids:
        assert public_id in pdf_text, f"{public_id} missing from public PDF"
        assert public_id in doc_text, f"{public_id} missing from cover memo"
    for internal_id in internal_ids:
        assert internal_id not in pdf_text, f"{internal_id} leaked into public PDF"
    assert 'ASIC comparator drift' not in pdf_text
    assert 'ASIC comparator drift' not in doc_text
