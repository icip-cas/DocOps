import json
import os
import re
import zipfile
from datetime import date, datetime
from html import unescape
from pathlib import Path

from docx import Document
from openpyxl import load_workbook
from openpyxl.utils.cell import get_column_letter, range_boundaries

from verifier_utils import *  # noqa: F401,F403

META_PATH = Path(os.environ.get("TASK_METADATA_PATH", "/tests/task_metadata.json"))
META = json.loads(META_PATH.read_text(encoding="utf-8"))
EXPECT = META["verifier_expectations"]


def _norm(value):
    if value is None:
        return ""
    if isinstance(value, datetime):
        return value.strftime("%Y-%m-%d")
    if isinstance(value, date):
        return value.strftime("%Y-%m-%d")
    if isinstance(value, float) and value.is_integer():
        value = int(value)
    text = str(value).replace("\u2013", "-").replace("\u2014", "-").replace("\xa0", " ")
    return re.sub(r"\s+", " ", text).strip()


def _norm_key(value):
    return re.sub(r"[^a-z0-9]+", "", _norm(value).lower())


def _norm_formula(value):
    return re.sub(r"\s+", "", _norm(value)).upper()


def _path(kind):
    env = {"clean": "CLEAN_DOCX_OUTPUT_PATH", "redline": "REDLINE_DOCX_OUTPUT_PATH", "xlsx": "XLSX_OUTPUT_PATH"}[kind]
    key = {"clean": "clean_docx_output", "redline": "redline_docx_output", "xlsx": "xlsx_output"}[kind]
    return Path(os.environ.get(env, EXPECT[key]))


def _doc_text(doc):
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append("|".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts)


def _heading_order(doc):
    out = []
    for p in doc.paragraphs:
        style = p.style.name if p.style else ""
        text = p.text.strip()
        if text and (style == "Title" or style.startswith("Heading")):
            out.append(text)
    return out


def _norm_heading_label(value):
    text = _norm(value)
    text = re.sub(r"^(?:\d+|[IVXLCDM]+)[\.)]\s+", "", text, flags=re.I)
    text = re.sub(r"^Section\s+\d+\s*[-:\u2013\u2014]\s*", "", text, flags=re.I)
    text = re.sub(r"^[A-Z][\.)]\s+", "", text)
    text = re.sub(r"[:\-]+$", "", text).strip()
    return text.casefold()


def _heading_sequence_matches(actual, expected):
    actual_norm = [_norm_heading_label(item) for item in actual]
    expected_norm = [_norm_heading_label(item) for item in expected]
    pos = 0
    for item in actual_norm:
        if pos < len(expected_norm) and item == expected_norm[pos]:
            pos += 1
    return pos == len(expected_norm)


def _token_stream(text):
    text = normalize_text(text)
    text = re.sub(r"\((?:one|two|three|four|five|six|twelve|\d+)(?:st|nd|rd|th)?\)", " ", text)
    text = text.replace("6%", "six percent").replace("3%", "three percent").replace("12%", "twelve percent")
    text = text.replace("5th", "fifth").replace("2", "two").replace("3", "three").replace("4", "four").replace("6", "six").replace("12", "twelve")
    return re.findall(r"[a-z0-9]+", text)


def _tokens_in_order(text, phrase):
    tokens = _token_stream(text)
    wanted = _token_stream(phrase)
    pos = 0
    for token in tokens:
        if pos < len(wanted) and token == wanted[pos]:
            pos += 1
    return pos == len(wanted)


def _doc_xml(path):
    with zipfile.ZipFile(path) as zf:
        return zf.read("word/document.xml").decode("utf-8", errors="ignore")


def _strip_xml(xml):
    return unescape(re.sub(r"<[^>]+>", "", xml))


def _tracked_change_text(xml, tag):
    chunks = re.findall(rf"<w:{tag}\b[^>]*>.*?</w:{tag}>", xml, flags=re.DOTALL)
    return normalize_text(" ".join(_strip_xml(chunk) for chunk in chunks))


def _redline_phrase_ok(insertion_text, deletion_text, flat, phrase, tag):
    haystack = insertion_text if tag == "ins" else deletion_text
    if _tokens_in_order(haystack, phrase):
        return True
    phrase_norm = normalize_text(phrase)
    if tag == "ins":
        if "6%" in phrase or "6 percent" in phrase_norm:
            return any(_tokens_in_order(haystack, alt) for alt in ["6%", "six percent", "6 percent"])
        if "two business hours" in phrase_norm:
            return _tokens_in_order(haystack, "two business hours") or (
                _tokens_in_order(haystack, "two") and _tokens_in_order(flat, "business hours")
            )
        if "monthly security evidence package" in phrase_norm:
            return (
                _tokens_in_order(haystack, "monthly security evidence package")
                and _tokens_in_order(haystack, "fifth business day")
            ) or (
                _tokens_in_order(haystack, "monthly")
                and _tokens_in_order(haystack, "fifth business day")
                and _tokens_in_order(flat, "security evidence package")
            )
        if "custom data extract requests are excluded" in phrase_norm:
            return _tokens_in_order(haystack, "custom data extract requests are excluded")
        if "vendor-notices@northwind.example" in phrase_norm:
            return "vendor-notices@northwind.example" in normalize_text(haystack)
    if tag == "del":
        if "3%" in phrase or "3 percent" in phrase_norm:
            return any(_tokens_in_order(haystack, alt) for alt in ["3%", "three percent", "3 percent"])
        if "four business hours" in phrase_norm:
            return _tokens_in_order(haystack, "four business hours") or (
                _tokens_in_order(haystack, "four") and _tokens_in_order(flat, "business hours")
            )
        if "quarterly security evidence package" in phrase_norm:
            return _tokens_in_order(haystack, "quarterly security evidence package") or (
                _tokens_in_order(haystack, "quarterly") and _tokens_in_order(flat, "security evidence package")
            )
        if "ad hoc data exports are included" in phrase_norm:
            return _tokens_in_order(haystack, "ad hoc data exports are included")
        if "contracts@northwind.example" in phrase_norm:
            return "contracts@northwind.example" in normalize_text(haystack)
    return _tokens_in_order(flat, phrase) and _tokens_in_order(haystack, phrase.split()[0])


def _clean_required_ok(text, phrase):
    if _tokens_in_order(text, phrase):
        return True
    norm_phrase = normalize_text(phrase)
    if "priority 1 response time" in norm_phrase:
        return _tokens_in_order(text, "Priority 1") and _tokens_in_order(text, "two business hours")
    if "priority 2 response time" in norm_phrase:
        return _tokens_in_order(text, "Priority 2") and _tokens_in_order(text, "one business day")
    if "monthly security evidence package" in norm_phrase:
        return _tokens_in_order(text, "monthly security evidence package") and _tokens_in_order(text, "fifth business day")
    if "service credit is 6%" in norm_phrase:
        return (
            _tokens_in_order(text, "Priority 1")
            and any(_tokens_in_order(text, alt) for alt in ["6%", "six percent", "6 percent"])
            and _tokens_in_order(text, "monthly SaaS fees")
        )
    if "aggregate monthly service credits" in norm_phrase:
        return _tokens_in_order(text, "aggregate monthly") and any(
            _tokens_in_order(text, alt) for alt in ["12%", "twelve percent", "12 percent"]
        )
    if "customer notice mailbox" in norm_phrase:
        return "vendor-notices@northwind.example" in normalize_text(text)
    return False


def _sheet(wb, sheet_name):
    assert sheet_name in wb.sheetnames, f"Missing sheet {sheet_name!r}; found {wb.sheetnames!r}"
    return wb[sheet_name]


def _table(ws, expected_name):
    for name in ws.tables.keys():
        if str(name).lower() == str(expected_name).lower():
            return ws.tables[name]
    raise AssertionError(f"Missing table {expected_name!r} on {ws.title!r}; found {list(ws.tables.keys())!r}")


def _table_bounds(table):
    return range_boundaries(str(table.ref).replace("$", ""))


def _table_has_size(ws, expected_name, min_rows, min_cols):
    table = _table(ws, expected_name)
    min_col, min_row, max_col, max_row = _table_bounds(table)
    assert max_row - min_row + 1 >= min_rows, f"{expected_name}: table has too few rows"
    assert max_col - min_col + 1 >= min_cols, f"{expected_name}: table has too few columns"
    return table


def _table_headers(ws, table):
    min_col, min_row, max_col, _ = _table_bounds(table)
    return [_norm(ws.cell(min_row, col).value) for col in range(min_col, max_col + 1)]


def _table_header_map(ws, table):
    min_col, min_row, max_col, _ = _table_bounds(table)
    return {_norm_key(ws.cell(min_row, col).value): col for col in range(min_col, max_col + 1)}


def _require_headers(ws, table, expected_headers):
    actual = [_norm_key(h) for h in _table_headers(ws, table)]
    missing = [h for h in expected_headers if _norm_key(h) not in actual]
    assert not missing, f"{ws.title}: missing table headers {missing!r}"


def _find_sheet_row(ws, row_token, table, start_data=True):
    min_col, min_row, max_col, max_row = _table_bounds(table)
    start = min_row + 1 if start_data else min_row
    token = _norm_key(row_token)
    for row in range(start, max_row + 1):
        values = [_norm(ws.cell(row, col).value) for col in range(min_col, max_col + 1)]
        if any(token in _norm_key(value) for value in values):
            return row, values
    raise AssertionError(f"{ws.title}: could not find row containing {row_token!r}")


def _find_sheet_row_any(ws, row_tokens, table):
    for token in row_tokens:
        try:
            return _find_sheet_row(ws, token, table)
        except AssertionError:
            continue
    raise AssertionError(f"{ws.title}: could not find row containing any of {row_tokens!r}")


def _has_row_containing(rows, tokens):
    for row in rows:
        text = " | ".join(row)
        if all(_norm_key(token) in _norm_key(text) for token in tokens):
            return True
    return False


def _table_data_rows(ws, table):
    min_col, min_row, max_col, max_row = _table_bounds(table)
    return [
        [_norm(ws.cell(row, col).value) for col in range(min_col, max_col + 1)]
        for row in range(min_row + 1, max_row + 1)
    ]


def _cell_by_header(ws, table, row, header):
    mapping = _table_header_map(ws, table)
    key = _norm_key(header)
    assert key in mapping, f"{ws.title}: missing header {header!r}"
    return ws.cell(row, mapping[key])


def _assert_row_contains(row_values, tokens, label):
    text = " | ".join(row_values)
    missing = [token for token in tokens if _norm_key(token) not in _norm_key(text)]
    assert not missing, f"{label}: missing {missing!r} in {row_values!r}"


def _formula_has_tokens(formula, tokens, label):
    normalized = _norm_formula(formula)
    assert normalized.startswith("="), f"{label}: expected formula, found {formula!r}"
    missing = [token for token in tokens if _norm_formula(token).lstrip("=") not in normalized]
    assert not missing, f"{label}: formula missing tokens {missing!r}: {formula!r}"


def _range_covered(ws, expected):
    min_col, min_row, max_col, max_row = range_boundaries(expected.replace("$", ""))
    target = {f"{get_column_letter(col)}{row}" for row in range(min_row, max_row + 1) for col in range(min_col, max_col + 1)}
    cells = set()
    for dv in ws.data_validations.dataValidation:
        for rng in dv.cells.ranges:
            min_c, min_r, max_c, max_r = range_boundaries(str(rng).replace("$", ""))
            cells.update(f"{get_column_letter(col)}{row}" for row in range(min_r, max_r + 1) for col in range(min_c, max_c + 1))
    return target.issubset(cells)


def _clean_area(value):
    text = _norm(value).replace("'", "").replace("$", "").replace(" ", "")
    if "!" in text:
        text = text.split("!", 1)[1]
    return text.upper()


def _sheet_text(wb, sheets):
    parts = []
    for sheet in sheets:
        ws = _sheet(wb, sheet)
        for row in ws.iter_rows():
            for cell in row:
                if cell.value is not None:
                    parts.append(str(cell.value))
    return "\n".join(parts)


def _defined_name_names(wb):
    names = set()
    defined_names = wb.defined_names
    if hasattr(defined_names, "values"):
        names.update(dn.name for dn in defined_names.values())
    else:
        names.update(dn.name for dn in defined_names.definedName)
    return names


def test_outputs_exist():
    assert _path("clean").exists(), f"Missing clean DOCX at {_path('clean')}"
    assert _path("redline").exists(), f"Missing redline DOCX at {_path('redline')}"
    assert _path("xlsx").exists(), f"Missing XLSX at {_path('xlsx')}"
    assert _path("clean").suffix.lower() == ".docx"
    assert _path("redline").suffix.lower() == ".docx"
    assert _path("xlsx").suffix.lower() == ".xlsx"


def test_clean_doc_content_and_privacy():
    doc = Document(_path("clean"))
    assert _heading_sequence_matches(_heading_order(doc), EXPECT["clean_heading_order"])
    text = _doc_text(doc)
    missing = [phrase for phrase in EXPECT["clean_required"] if not _clean_required_ok(text, phrase)]
    assert not missing, f"clean amendment: missing required phrases: {missing}"
    forbid_any(text, EXPECT["forbidden_public"], "clean amendment")


def test_redline_ooxml_tracked_changes():
    xml = _doc_xml(_path("redline"))
    assert "<w:ins" in xml, "redline doc must contain OOXML insertions"
    assert "<w:del" in xml, "redline doc must contain OOXML deletions"
    flat = _strip_xml(xml)
    insertion_text = _tracked_change_text(xml, "ins")
    deletion_text = _tracked_change_text(xml, "del")
    for phrase in EXPECT["redline_insertions"]:
        assert _redline_phrase_ok(insertion_text, deletion_text, flat, phrase, "ins"), f"missing tracked insertion phrase {phrase!r}"
    for phrase in EXPECT["redline_deletions"]:
        assert _redline_phrase_ok(insertion_text, deletion_text, flat, phrase, "del"), f"missing tracked deletion phrase {phrase!r}"
    for phrase in EXPECT["forbidden_public"]:
        assert normalize_text(phrase) not in normalize_text(flat)


def test_workbook_structure_formulas_controls():
    wb = load_workbook(_path("xlsx"), data_only=False)
    assert wb.sheetnames == EXPECT["sheet_order"]
    for sheet, (table_name, ref) in EXPECT["tables"].items():
        ws = _sheet(wb, sheet)
        expected_min_cols = range_boundaries(ref)[2] - range_boundaries(ref)[0] + 1
        expected_min_rows = range_boundaries(ref)[3] - range_boundaries(ref)[1] + 1
        _table_has_size(ws, table_name, expected_min_rows, expected_min_cols)

    expected_headers = {
        "Package Summary": ["Field", "Value", "Source", "Review"],
        "Change Log": ["Change ID", "Clause", "Original Text", "Final Text", "Source", "Risk", "Owner", "Status"],
        "Obligations Register": ["Obligation", "Frequency/Trigger", "Due/Target", "Evidence", "Owner", "Monitor?"],
        "Source Map": ["Source", "Used For", "Public Output Treatment"],
        "Approval Checklist": ["Reviewer", "Area", "Required Check", "Status"],
        "Private Fallback": ["Item", "Treatment"],
    }
    for sheet, headers in expected_headers.items():
        ws = _sheet(wb, sheet)
        table = _table(ws, EXPECT["tables"][sheet][0])
        _require_headers(ws, table, headers)

    for sheet in EXPECT["hidden_sheets"]:
        assert _sheet(wb, sheet).sheet_state in ("hidden", "veryHidden")
    names = _defined_name_names(wb)
    for name in EXPECT["defined_names"]:
        assert name in names

    summary = _sheet(wb, "Package Summary")
    summary_table = _table(summary, "tblPackageSummary")
    for label, tokens in {
        "Accepted changes": ["=", "COUNTIF"],
        "Monitored obligations": ["=", "COUNTIF"],
        "Package ready": ["=", "IF"],
    }.items():
        row, _ = _find_sheet_row(summary, label, summary_table)
        _formula_has_tokens(_cell_by_header(summary, summary_table, row, "Value").value, tokens, f"Package Summary {label}")

    for sheet, ranges in EXPECT["data_validation"].items():
        ws = _sheet(wb, sheet)
        for rng in ranges:
            assert _range_covered(ws, rng), f"{sheet}: missing validation over {rng}"
    for sheet, area in EXPECT["print_areas"].items():
        actual = _clean_area(_sheet(wb, sheet).print_area)
        table = _table(_sheet(wb, sheet), EXPECT["tables"][sheet][0])
        target = _clean_area(table.ref)
        assert target in actual or actual in target


def test_workbook_values_and_privacy():
    wb = load_workbook(_path("xlsx"), data_only=False)

    summary = _sheet(wb, "Package Summary")
    summary_table = _table(summary, "tblPackageSummary")
    for label, tokens in {
        "Supplier": ["Cloud Harbor Systems, Inc."],
        "Customer": ["Northwind Health Plan"],
        "Effective date": ["2026-09-01"],
    }.items():
        _, row_values = _find_sheet_row(summary, label, summary_table)
        _assert_row_contains(row_values, tokens, f"Package Summary {label}")
    summary_rows = _table_data_rows(summary, summary_table)
    assert _has_row_containing(summary_rows, ["Service Level Amendment No. 3"]), (
        "Package Summary: missing amendment identity"
    )
    assert (
        _has_row_containing(summary_rows, ["Private", "excluded"])
        or _has_row_containing(summary_rows, ["internal_fallback_positions.docx", "excluded"])
    ), "Package Summary: missing private-source exclusion"

    changes = _sheet(wb, "Change Log")
    changes_table = _table(changes, "tblChangeLog")
    for change_id, phrases in {
        "CH-01": ["Priority 1", "four business hours", "two business hours", "Accepted"],
        "CH-02": ["Security Evidence", "Quarterly", "Monthly", "fifth business day", "Accepted"],
        "CH-03": ["Service Credit", "3%", "6%", "Accepted"],
        "CH-04": ["Excluded Requests", "Ad hoc data exports", "Custom data extract", "Accepted"],
        "CH-05": ["Operational Contacts", "contracts@northwind.example", "vendor-notices@northwind.example", "Priya Nair", "Accepted"],
    }.items():
        _, row_values = _find_sheet_row(changes, change_id, changes_table)
        row_text = " | ".join(row_values)
        missing = [phrase for phrase in phrases if not (_tokens_in_order(row_text, phrase) or _norm_key(phrase) in _norm_key(row_text))]
        assert not missing, f"Change Log {change_id}: missing {missing!r} in {row_values!r}"

    obligations = _sheet(wb, "Obligations Register")
    obligations_table = _table(obligations, "tblObligations")
    for obligation, tokens in {
        "Priority 1 response": ["business hours", "Yes"],
        "Priority 2 response": ["business day", "Yes"],
        "Security evidence package": ["Monthly", "business day", "SOC 2", "Yes"],
        "Priority 1 service credit": ["6", "Yes"],
        "credit cap": ["12", "Yes"],
        "custom data extract": ["exclu"],
    }.items():
        _, row_values = _find_sheet_row_any(obligations, [obligation, obligation.replace("custom data extract", "Excluded custom data extracts")], obligations_table)
        _assert_row_contains(row_values, tokens, f"Obligations Register {obligation}")

    source = _sheet(wb, "Source Map")
    source_table = _table(source, "tblSourceMap")
    for item, token_options in {
        "current_sla_amendment.docx": [["Original", "clause"], ["Base", "contract"]],
        "counterparty_comment_digest.pdf": [["redlines"], ["change"], ["comment"]],
        "security_evidence_schedule.xlsx": [["Evidence"]],
        "precedent_clause_library.docx": [["Exclusion"], ["clause"]],
        "deal_summary.pdf": [["Credit"], ["Effective date"]],
        "internal_fallback_positions.docx": [["internal", "Exclude"], ["private", "Exclude"], ["strategy", "Exclude"]],
    }.items():
        _, row_values = _find_sheet_row(source, item, source_table)
        assert any(all(_norm_key(token) in _norm_key(" | ".join(row_values)) for token in tokens) for tokens in token_options), (
            f"Source Map {item}: missing accepted source semantics in {row_values!r}"
        )

    approval = _sheet(wb, "Approval Checklist")
    approval_table = _table(approval, "tblApproval")
    for reviewer, tokens in {
        "Legal Ops": [],
        "Security GRC": ["Evidence"],
        "Finance": ["Service Credit"],
        "Vendor Manager": ["Contacts"],
    }.items():
        _, row_values = _find_sheet_row(approval, reviewer, approval_table)
        if tokens:
            _assert_row_contains(row_values, tokens, f"Approval Checklist {reviewer}")
        assert _norm(row_values[-1]), f"Approval Checklist {reviewer}: missing status in {row_values!r}"

    public = _sheet_text(wb, ["Package Summary", "Change Log", "Obligations Register", "Source Map", "Approval Checklist"])
    forbid_any(public, EXPECT["forbidden_public"], "public workbook sheets")


def test_cross_output_consistency():
    clean = _doc_text(Document(_path("clean")))
    wb = load_workbook(_path("xlsx"), data_only=False)
    workbook_public = _sheet_text(wb, ["Change Log", "Obligations Register", "Approval Checklist"])
    for anchor in ["two business hours", "monthly security evidence package", "6% of monthly SaaS fees", "12%", "vendor-notices@northwind.example", "Priya Nair"]:
        assert _clean_required_ok(clean, anchor) or normalize_text(anchor) in normalize_text(clean)
        assert _tokens_in_order(workbook_public, anchor) or normalize_text(anchor) in normalize_text(workbook_public)

    summary = _sheet(wb, "Package Summary")
    summary_table = _table(summary, "tblPackageSummary")
    row, _ = _find_sheet_row(summary, "Package ready", summary_table)
    _formula_has_tokens(_cell_by_header(summary, summary_table, row, "Value").value, ["IF"], "package-ready status")
    changes = _sheet(wb, "Change Log")
    changes_table = _table(changes, "tblChangeLog")
    for change_id in ["CH-01", "CH-02", "CH-03", "CH-04", "CH-05"]:
        row, _ = _find_sheet_row(changes, change_id, changes_table)
        assert _norm_key(_cell_by_header(changes, changes_table, row, "Status").value) == "accepted"
