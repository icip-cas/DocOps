import json
import os
import re
import sys
import zipfile
from pathlib import Path
from datetime import date, datetime

import pdfplumber
from docx import Document
from pypdf import PdfReader

sys.path.insert(0, str(Path(__file__).parent))
from verifier_utils import *  # noqa: F401,F403

META_PATH = Path(os.environ.get('TASK_METADATA_PATH', '/tests/task_metadata.json'))
META = json.loads(META_PATH.read_text(encoding='utf-8'))
EXPECT = META['verifier_expectations']


def _vopt_norm_text(value):
    if value is None:
        return ''
    if isinstance(value, datetime):
        if value.hour or value.minute or value.second:
            return value.strftime('%Y-%m-%d %H:%M')
        return value.strftime('%Y-%m-%d')
    if isinstance(value, date):
        return value.strftime('%Y-%m-%d')
    if isinstance(value, float) and value.is_integer():
        value = int(value)
    text = str(value)
    text = text.replace('\u2013', '-').replace('\u2014', '-').replace('\u2212', '-')
    text = text.replace('\xa0', ' ')
    text = re.sub(r'\s+', ' ', text).strip()
    return text


def _vopt_norm_formula(value):
    return re.sub(r'\s+', '', _vopt_norm_text(value)).upper()


def _vopt_assert_formula(actual, expected, label='formula'):
    assert _vopt_norm_formula(actual) == _vopt_norm_formula(expected), f"{label}: expected {expected!r}, found {actual!r}"


def _vopt_norm_rows(rows):
    return [[_vopt_norm_text(cell) for cell in row] for row in rows]


def _vopt_rows_equal(actual, expected):
    return _vopt_norm_rows(actual) == _vopt_norm_rows(expected)


def _vopt_norm_key(value):
    return re.sub(r'[^a-z0-9]+', '', _vopt_norm_text(value).lower())


def _vopt_rows_text(rows):
    return '\n'.join('|'.join(_vopt_norm_text(cell) for cell in row) for row in rows)


def _vopt_header_has(header, required):
    actual = {_vopt_norm_text(cell).lower() for cell in header}
    missing = [cell for cell in required if _vopt_norm_text(cell).lower() not in actual]
    assert not missing, f"Missing expected header columns {missing!r}; actual={header!r}"


def _vopt_section_text(section_part):
    return '\n'.join(p.text for p in getattr(section_part, 'paragraphs', []) if p.text is not None)


def _vopt_header_text(doc):
    return '\n'.join(_vopt_section_text(section.header) for section in doc.sections)


def _vopt_footer_text(doc):
    return '\n'.join(_vopt_section_text(section.footer) for section in doc.sections)


def _vopt_table_ref(ws, expected_name=None):
    tables = ws.tables
    if expected_name:
        for name in tables.keys():
            if str(name).lower() == str(expected_name).lower():
                return tables[name].ref
    vals = list(tables.values())
    assert vals, f"No Excel native table found on sheet {ws.title!r}"
    return vals[0].ref


def _vopt_clean_area(value):
    text = _vopt_norm_text(value).replace("'", '').replace('$', '').replace(' ', '')
    if '!' in text:
        text = text.split('!', 1)[1]
    return text.upper()


def _vopt_assert_print_area(ws, expected):
    actual = _vopt_clean_area(ws.print_area)
    target = _vopt_clean_area(expected)
    assert target in actual or actual in target, f"{ws.title}: expected print area {expected!r}, found {ws.print_area!r}"


def _vopt_freeze_pane(value):
    return getattr(value, 'coordinate', value)


def _vopt_validated_cells(ws):
    cells = set()
    for dv in ws.data_validations.dataValidation:
        for rng in dv.cells.ranges:
            text = str(rng).replace('$', '')
            try:
                min_col, min_row, max_col, max_row = range_boundaries(text)
            except ValueError:
                cells.add(text)
                continue
            if max_row > 10000:
                max_row = max(min_row, 200)
            for row in range(min_row, max_row + 1):
                for col in range(min_col, max_col + 1):
                    cells.add(f"{get_column_letter(col)}{row}")
    return cells


class _VoptValidationRanges(list):
    def __init__(self, ws, ranges):
        super().__init__(ranges)
        self.ws = ws

    def __contains__(self, expected):
        expected = str(expected).replace('$', '')
        if super().__contains__(expected):
            return True
        try:
            min_col, min_row, max_col, max_row = range_boundaries(expected)
        except ValueError:
            return super().__contains__(expected)
        target = {f"{get_column_letter(col)}{row}" for row in range(min_row, max_row + 1) for col in range(min_col, max_col + 1)}
        return target.issubset(_vopt_validated_cells(self.ws))


def _vopt_dv_ranges(ws):
    ranges = []
    for dv in ws.data_validations.dataValidation:
        ranges.extend(str(rng).replace('$', '') for rng in dv.cells.ranges)
    return _VoptValidationRanges(ws, ranges)


def _vopt_ordered_subset(expected, actual):
    actual_iter = iter([_vopt_norm_text(item) for item in actual])
    for item in [_vopt_norm_text(value) for value in expected]:
        for candidate in actual_iter:
            if item == candidate:
                break
        else:
            return False
    return True



def _resolve(kind):
    if kind == 'document':
        return Path(os.environ.get('DOCX_OUTPUT_PATH', EXPECT['document_output']))
    if kind == 'packet':
        return Path(os.environ.get('PDF_OUTPUT_PATH', EXPECT['packet_output']))
    raise KeyError(kind)


def _doc_text(doc):
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append('|'.join(cell.text.strip() for cell in row.cells))
    return '\n'.join(parts)


def _headings(doc):
    out = []
    for p in doc.paragraphs:
        style = p.style.name if p.style else ''
        if p.text.strip() and (style.startswith('Heading') or style == 'Title'):
            out.append(p.text.strip())
    return out


def _has_toc(path):
    with zipfile.ZipFile(path) as zf:
        xml = zf.read('word/document.xml').decode('utf-8', errors='ignore')
    return bool(re.search(r'TOC\s+(?:\\)?o|TOC\s*(?:&quot;|")', xml))


def _table_rows(table):
    return [[cell.text.strip() for cell in row.cells] for row in table.rows]


def _highlighted(doc, text):
    for p in doc.paragraphs:
        if _vopt_norm_text(text) in _vopt_norm_text(p.text):
            return any(run.font.highlight_color is not None for run in p.runs) or '<w:highlight' in p._p.xml
    return False


def _has_highlighted_terms(doc, required_terms):
    for p in doc.paragraphs:
        para_text = _vopt_norm_text(p.text).lower()
        if not para_text:
            continue
        has_highlight = any(run.font.highlight_color is not None for run in p.runs) or '<w:highlight' in p._p.xml
        if has_highlight and all(term.lower() in para_text for term in required_terms):
            return True
    return False


def _assert_heading_sequence(doc):
    actual = [_vopt_norm_key(item) for item in _headings(doc)]
    expected = [_vopt_norm_key(item) for item in EXPECT['heading_order']]
    pos = 0
    for item in actual:
        if pos < len(expected) and item == expected[pos]:
            pos += 1
    assert pos == len(expected), f"Missing or out-of-order required headings: {_headings(doc)!r}"


def _assert_clause_matrix(doc):
    expected = EXPECT['clause_rows']
    for table in doc.tables:
        rows = _vopt_norm_rows(_table_rows(table))
        header = [_vopt_norm_key(cell) for cell in rows[0]] if rows else []
        if {'clause', 'title', 'finalplacement', 'requiredstyle'}.issubset(set(header)):
            body = rows[1:]
            for exp in expected:
                matches = [row for row in body if row and _vopt_norm_text(row[0]) == exp[0]]
                assert matches, f'Missing clause matrix row for {exp[0]}'
                row_text = _vopt_norm_text(' | '.join(matches[0])).lower()
                for token in exp[1:]:
                    assert _vopt_norm_text(token).lower() in row_text, f'Clause {exp[0]} missing {token!r}'
            return
    raise AssertionError('Missing clause formatting matrix table')


def _pdf_title_matches(expected, page_text):
    norm_page = _vopt_norm_key(page_text)
    norm_expected = _vopt_norm_key(expected)
    if norm_expected in norm_page:
        return True
    aliases = {
        'exhibitatadataprocessingaddendum': ['exhibitadataprocessingaddendum'],
        'exhibitbservicelevelcredits': ['exhibitbservicelevelcredits'],
    }
    return any(alias in norm_page for alias in aliases.get(norm_expected, []))


def _pdf_texts(path):
    with pdfplumber.open(str(path)) as pdf:
        return [page.extract_text() or '' for page in pdf.pages]


def _contains_all(text, terms):
    norm = _vopt_norm_text(text).lower()
    return all(_vopt_norm_text(term).lower() in norm for term in terms)


def _assert_docx_semantic_content(text):
    assert _contains_all(text, ['1.1', 'Service Catalog']), 'docx: missing Service Catalog clause'
    assert _contains_all(text, ['Clause 2.1', 'Data Processing', 'Exhibit A']), 'docx: missing Clause 2.1 / Exhibit A promotion'
    assert _contains_all(text, ['Clause 3.1', 'Service Level', 'Exhibit B']), 'docx: missing Clause 3.1 / Exhibit B promotion'
    assert (
        _contains_all(text, ['Exhibits A', 'B', 'only'])
        or _contains_all(text, ['Exhibits: A-B'])
        or _contains_all(text, ['Exhibits A-B'])
    ), 'docx: missing A-B-only exhibit boundary'


def _assert_pdf_semantic_content(text):
    assert _contains_all(text, ['Clause 2.1']), 'pdf: missing Clause 2.1'
    assert _contains_all(text, ['Clause 3.1']), 'pdf: missing Clause 3.1'
    assert EXPECT['header_contains'] in _vopt_norm_text(text), 'pdf: missing final header text'
    assert (
        _contains_all(text, ['Exhibits A', 'B', 'only'])
        or _contains_all(text, ['Exhibits Included', 'A', 'B'])
        or _contains_all(text, ['Exhibits: A-B'])
        or _contains_all(text, ['Exhibits A-B'])
    ), 'pdf: missing A-B-only exhibit statement'


def _has_obsolete_exhibit_c(text):
    return bool(re.search(r'\bExhibit\s+C\b', _vopt_norm_text(text), flags=re.IGNORECASE))


def test_outputs_exist():
    assert _resolve('document').exists()
    assert _resolve('packet').exists()
    assert _resolve('document').suffix.lower() == '.docx'
    assert _resolve('packet').suffix.lower() == '.pdf'


def test_docx_heading_toc_and_style_content():
    path = _resolve('document')
    doc = Document(path)
    _assert_heading_sequence(doc)
    text = _doc_text(doc)
    _assert_docx_semantic_content(text)
    forbid_any(text, EXPECT['forbidden_text'], 'docx')
    assert _has_toc(path), 'Expected real Word TOC field'
    assert EXPECT['header_contains'] in _vopt_header_text(doc)
    assert EXPECT['footer_contains'] in _vopt_footer_text(doc)
    _assert_clause_matrix(doc)
    assert _has_highlighted_terms(doc, ['Clause 2.1', 'Exhibit A']), 'Missing highlighted Clause 2.1 / Exhibit A exception'
    assert _has_highlighted_terms(doc, ['Clause 3.1', 'Exhibit B']), 'Missing highlighted Clause 3.1 / Exhibit B exception'


def test_pdf_packet_pages_bookmarks_and_boundary():
    path = _resolve('packet')
    reader = PdfReader(str(path))
    assert len(reader.pages) == EXPECT['pdf_page_count']
    texts = _pdf_texts(path)
    for expected_title, page_text in zip(EXPECT['pdf_titles'], texts):
        assert _pdf_title_matches(expected_title, page_text), f"Missing PDF title {expected_title!r} on its expected page"
    full = '\n'.join(texts)
    _assert_pdf_semantic_content(full)
    forbid_any(full, EXPECT['pdf_forbidden'], 'pdf')
    outline = [title for _level, title in flatten_outline(reader.outline)]
    for title in EXPECT['outline_titles']:
        assert any(_vopt_norm_key(title) == _vopt_norm_key(item) for item in outline), f'Missing PDF bookmark: {title}'


def test_cross_output_exhibit_consistency():
    doc_text = _doc_text(Document(_resolve('document')))
    pdf_text = '\n'.join(_pdf_texts(_resolve('packet')))
    for exhibit in EXPECT['public_exhibits']:
        assert exhibit in doc_text, f'{exhibit} missing from Word output'
        assert exhibit in pdf_text, f'{exhibit} missing from PDF output'
    assert not _has_obsolete_exhibit_c(doc_text)
    assert not _has_obsolete_exhibit_c(pdf_text)
